#!/usr/bin/env python3
"""Behaviour tests for the docx skill's scripts (discussions/059 S4).

L1 proves each declared capability runs and produces an artifact; L2 proves the
artifact is a legal, non-lossy Word document. Neither can say whether the RESULT is
right — that a phrase Word split across three runs was actually found, that the
appended paragraph landed before `<w:sectPr>` and not after it, that an unpack and
repack gave back the same 17 parts in the same order. Those claims are this file's
job.

    python3 scripts/test-docx-skill.py
    python3 scripts/test-docx-skill.py --json

Every assertion runs twice: once against the real output of the real scripts (must
stay silent) and once against output carrying exactly the defect it hunts (must
fire). The flaws are not invented damage — each is the implementation somebody
reaches for first. `replace-run-by-run` above all: iterating `paragraph.runs` and
calling `str.replace` on each is what every example on the internet does, and on
this skill's own sample it finds **1 of the 2** occurrences — the one that happens
to sit inside a single run — and silently misses the one that spans two. Partial
success is why nobody notices.

The run prints a flaw -> fired-checks matrix. "All the negative controls went red"
is not the claim; "the RIGHT check went red" is, and only the matrix shows the
difference.

Exit 0 = every assertion behaved, 1 = something did not.
"""
from __future__ import annotations

import argparse
import copy
import json
import os
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

# ── stdout must be UTF-8 on every platform, and on Windows it is not ──────────
# This gate prints ✅/❌ and Chinese. Windows encodes a CAPTURED stdout in the
# machine's ANSI code page and Python only defaults to UTF-8 from 3.15 (PEP 686);
# CI pins 3.11. Measured on CI: this script died with
# `UnicodeEncodeError: 'charmap' codec can't encode character '\u2705'` inside its
# own `print(json.dumps(...))`. The skills were fixed for this first and the GATES
# were missed — the same defect has two homes, and only one of them was product.
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        try:
            _stream.reconfigure(encoding="utf-8")
        except (ValueError, OSError):        # already detached / not reconfigurable
            pass


REPO = Path(__file__).resolve().parent.parent
SKILL = REPO / "skills" / "builtin" / "docx"
FIXTURES = SKILL / "fixtures"
REPORT = FIXTURES / "report.docx"
REVISED = FIXTURES / "revised.docx"      # report.docx after a round of edits (W18)
UNORDERED = FIXTURES / "unordered.docx"
PY = sys.executable

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# What report.docx is known to hold. Spelled out rather than read back from the
# file: an expectation derived from the artifact it checks agrees with itself no
# matter what the artifact says.
TITLE = "二零二六年第三季度经营分析报告"
SPLIT_PARAGRAPH = "2026 年第三季度营业收入同比增长 12%，毛利率保持稳定。"
# Stored as "2026 年第" | "三季度" | "营业收入…" — this phrase spans the first two.
CROSS_RUN = "第三季度"
CROSS_RUN_NEW = "第四季度"
# The same phrase also occurs inside the TITLE, in a single run. Both in one call is
# deliberate: a replacement that only handles one of the two cases is a partial
# implementation that looks complete from the outside.
CROSS_RUN_TOTAL = 2
CROSS_RUN_SPANNING = 1
# What the per-run implementation finds, measured by running it: exactly the
# occurrences that do NOT span a run. It is not zero, and that matters — a tool that
# found nothing would be reported as broken within a minute.
NAIVE_HITS = CROSS_RUN_TOTAL - CROSS_RUN_SPANNING
PLACEHOLDER = "{{客户名称}}"
PLACEHOLDER_NEW = "示例科技有限公司"
INSERTED = "净利润"          # inside <w:ins>
# Inside <w:del>, therefore NOT document text. A phrase that occurs nowhere else:
# an earlier draft used "毛利", a substring of the ordinary body text
# "毛利率保持稳定。", so R2 fired on a correct implementation.
DELETED = "扣非净利"
COMMENT_TEXT = "请与银行流水核对后再定稿。"
HEADER_TEXT = "内部资料 · 请勿外传 · 密级：{{密级}}"
LIST_TEXT = "费用结构持续优化，管理费用同比下降。"
TABLE_CELL = "营业收入"
PARTS_TOTAL = 17
APPEND_TEXT = "结论：维持全年增长预期。"

# The three order defects unordered.docx carries, and nothing else.
ORDER_DEFECTS = 3

# W5's placeholders. 客户名称 is split across runs in the fixture, 密级 lives in the
# header — a filler that only walks word/document.xml leaves the letterhead unfilled.
FILL_SPLIT = "客户名称"
FILL_HEADER = "密级"
FILL_LEFT = "日期"          # deliberately NOT supplied, so `unfilled` has a subject
FILL_TYPO = "客户名"        # a key that matches no placeholder

# Assertions that need LibreOffice. On a host without it they are SKIPPED and named,
# never folded into the pass count — and so are the negative controls that could not
# fire, because a control nobody ran is not a control (059 §7, the recalc-drift雷).
# Q3 and Q5 are here because W19's two most load-bearing claims are only true if a
# layout engine agrees: a header row either does or does not reappear on page 2, and
# a cell's text either does or does not survive on one line.
SOFFICE_CHECKS = {"Y1", "Y2", "Y3", "Q3", "Q5"}

STDOUT_BUDGET = 6000        # bytes one call may print for a long document
SCALE_PARAGRAPHS = 2000     # comfortably past docxcommon.STDOUT_ITEM_LIMIT
# The other shape of oversized report, and the one the item-count trimmer cannot
# see: FEW entries, each enormous. One table this tall printed 130,602 bytes before
# the byte budget existed. Found by asking "is this compatible with Team mode?",
# not by any assertion — C1 only ever exercised a long list.
SCALE_TABLE_ROWS = 800

# ── W18: the A/B pair the diff is measured on ─────────────────────────────────
DOC_PART = "word/document.xml"
DIFF_HEADER_PART = "word/header1.xml"
# One of every whitelisted change, in one document, so "it found the text edit" can
# never be mistaken for "it found everything".
DIFF_NEW_STYLE = "Heading1"
DIFF_ADDED = "新增结论：全年目标维持不变。"
DIFF_REMOVED = "经营性现金流保持健康水平，Q3 无异常。"
DIFF_CELL_NEW = "1,860"     # a table cell's text
DIFF_HEADER_NEW = "机密资料 · 请勿外传 · 密级：{{密级}}"
# text · style · removed · added · table cell · header = six, and the count is
# asserted so that a seventh finding (noise leaking in) is as loud as a missing one.
DIFF_EXPECTED = 6
DIFF_MOVED_FROM, DIFF_MOVED_TO = 6, 2
# The categories the report must always name, with a count for each document — even
# when every count is zero. A report that mentions them only when it found some
# cannot be told from one that never looked.
DIFF_IGNORED_KEYS = ("rsid", "proofErr", "bookmark", "lang", "empty_runs",
                     "attribute_order", "zip_entry_order")

# ── W19: what the table presets are measured against ──────────────────────────
TABLE_PRESETS = ("grid", "finance", "banded")
# report.docx's table, its cells rewritten so that column width has something to bite
# on: one column of Chinese, two of Latin, with the SAME character count. That is the
# whole experiment — `len()` calls those columns equal and the page does not.
#
# ⚠️ Every string here occurs nowhere else in report.docx, and V0 checks that it stays
# that way. The first draft used 二零二六年第三季度经营分析, which is a PREFIX OF THE
# DOCUMENT'S TITLE, so "the cell's text survived on one line" read as true off the
# title line whether the cell had wrapped or not — the assertion passed for both
# implementations. Same shape as the 毛利 trap further up this file.
WIDE_ROWS = [
    ["费用项目", "REVENUE-2026", "GROWTH-RATE%"],
    ["华南大区渠道推广费用明细表", "1,240,000", "+12.0%"],
    ["西北大区营销费用合计数", "986,500", "+1.2pt"],
]
WIDE_LONGEST = "华南大区渠道推广费用明细表"
# Display width per column (a CJK character counts two) versus len(). The two agree
# on columns 2 and 3 and differ by a factor of two on column 1, which is the only
# reason a naive implementation and a correct one can be told apart at all.
WIDE_DISPLAY = [26, 12, 12]
WIDE_LENS = [13, 12, 12]
# dxa one display-width cell needs, MEASURED off a render (CJK came back at exactly
# 105, Latin between 94 and 122). The skill allocates 130 — the widest measured,
# rounded up. Repeated here rather than imported: an expectation that reads the
# implementation's own constant agrees with it whatever that constant becomes.
WIDE_DXA_PER_CELL = 130
WIDE_MARGIN = {"grid": 108, "finance": 144, "banded": 108}
# Rows that must carry w:tblHeader, by preset: the header row, or nothing at all.
PRESET_REPEATS = {"grid": [], "finance": [0], "banded": [0]}
# Border weights in EIGHTHS of a point, per edge, by preset. 0 means an explicit
# `w:val="none"` — not an omitted element, which would let a table style show through.
PRESET_BORDERS = {
    "grid": {"top": 4, "left": 4, "bottom": 4, "right": 4, "insideH": 4, "insideV": 4},
    "finance": {"top": 12, "left": 0, "bottom": 12, "right": 0,
                "insideH": 0, "insideV": 0},
    "banded": {"top": 4, "left": 4, "bottom": 4, "right": 4, "insideH": 4, "insideV": 4},
}
# Rows enough to push the table onto a third page, so "the header came back" is asked
# of more than one page break.
TALL_PRESET_ROWS = 60

# W6/W7. The fixture already carries ONE revision by 张审阅 (净利润 inserted, 扣非净利
# deleted) — that second author is what makes per-author filtering testable at all.
REVISER = "张三"
FIXTURE_AUTHOR = "张审阅"
TRACKED_PARAGRAPH = 1           # the one holding the cross-run phrase
REVISION_PARAGRAPH = 5          # the one the fixture already had a revision in
ACCEPTED_REVISION = "本季度净利润同比增长。"
REJECTED_REVISION = "本季度扣非净利同比增长。"
INSERTED_PARAGRAPH = "新增结论段落。"

# W8.
COMMENT_ANCHOR_SPLIT = "三季度营业收入"    # its first occurrence spans two runs
NEW_COMMENT = "季度口径需与年报一致。"

# ── W9 / W10 / W11 / W16: the layout family, built on outline.docx ────────────
# outline.docx exists because report.docx cannot test "create": it already HAS a
# header, a footer, a picture and a png Default, so every create path would run as a
# replace and the package wiring would never be exercised.
OUTLINE = FIXTURES / "outline.docx"
FONTLESS = FIXTURES / "fontless.docx"
CHART = FIXTURES / "chart.png"

OUTLINE_PARAGRAPHS = 10
OUTLINE_HEADINGS = ("经营概况", "收入分析", "分产品收入", "成本与费用", "风险提示")
OUTLINE_LEVELS = (1, 2, 3, 2, 1)
OUTLINE_BODY = "营业收入同比增长，主要来自华东区域。"

# W9. A header is four package pieces; the first-page and even-page variants are
# five, and the fifth is the one that decides whether the other four do anything.
HEADER_NEW = "示例科技有限公司 · 季度经营分析"
FOOTER_NEW = "内部资料，请勿外传"
FIRST_HEADER = "封面页眉"
EVEN_HEADER = "偶数页眉"

# W10. Six paragraphs: one title plus one per heading. A field with no cached result
# would add two, which is why the count is asserted and not just "more than before".
TOC_TITLE = "目录"
TOC_PARAGRAPHS_ADDED = 6
TOC_PLACEHOLDER = "—"
# Measured, in a PDF LibreOffice made from the artifact: with TOCHeading based on
# Heading1 and outline numbering attached to the heading styles, the contents page
# took number 1 and the real first chapter became 2 — "1. 目录 / 2. 经营概况 /
# 2.1 收入分析 / 3. 风险提示". Nothing in the package was invalid.
TOC_HEADING_STYLE = "TOCHeading"

# W11. Measured from the fixture rather than assumed, and the two numbers are far
# enough apart that an implementation reading the file's own pHYs density and one
# assuming the web's 96 dpi cannot both be called right.
CHART_PX = (240, 120)
CHART_INTRINSIC_EMU = (1462919, 731460)      # 240x120 at the declared 150 dpi
CHART_96DPI_EMU = (2286000, 1143000)         # what "everyone uses 96" produces
IMAGE_WIDTH_CM = 8
IMAGE_WIDTH_EMU = (2880000, 1440000)         # 8cm, aspect ratio kept
# A 1x1 GIF89a: the shortest valid picture in a DIFFERENT format, so replacing with
# it exercises the path where the old media part has to go and a new content-type
# Default has to be declared.
TINY_GIF = bytes.fromhex(
    "47494638396101000100800000000000ffffff21f90401000000002c00000000"
    "010001000002024401003b")

# W16. fontless.docx carries three DIFFERENT ways of losing a font binding, and they
# do not have the same right answer — which is the whole reason there are three.
FONTLESS_UNBOUND_RUNS = 3
FONTLESS_STYLE_EA = "黑体"          # what Heading2 already says; a repair must keep it
FONTLESS_KEPT_ASCII = "Times New Roman"   # a deliberate latin face, must survive
FONTLESS_FALLBACK_EA = "宋体"

# W12. A style id that is not one of Word's built-ins, so `customStyle` has a
# subject; Chinese on purpose, because a style id is a string and the paths that
# carry one are the same paths the Windows encoding defect lived in.
NEW_STYLE = "正文小字"
# Heading1 is used by exactly ONE paragraph in report.docx. That number is what turns
# "modifying a style repaints everything using it" into a checkable claim rather than
# a warning nobody can verify.
HEADING_USERS = 1

# W4. sample.md carries every construct the generator maps AND three it does not —
# a footnote and two raw HTML tags — because "what it cannot map is NAMED, never
# dropped" needs a subject. A fixture with only supported syntax would let a
# generator that silently discarded footnotes pass every check there is.
SAMPLE_MD = FIXTURES / "sample.md"
# ⚠️ TWO different paragraph counts, and they are both right. `paragraph_texts()`
# here walks EVERY <w:p> (24: fifteen in the body plus nine inside the 3x3 table's
# cells), while python-docx's `doc.paragraphs` — which L2's D3 uses, and which
# capabilities.json therefore states as 15 — walks only the body's direct children.
# Measured rather than assumed: the first version of this constant took the number
# from the wrong ruler and fired on a correct document.
MD_PARAGRAPHS = 24
MD_PARAGRAPHS_PYDOCX = 15        # what capabilities.json's `expect` must say
MD_TABLES = 1
MD_UNSUPPORTED = 3
MD_STYLES = ("Heading1", "Heading2", "SourceCode", "Quote", "Hyperlink", "CodeChar")
# The deepest list nesting in the fixture: "- 华东…" then "  - 其中线上渠道…".
MD_LIST_DEPTH = 1

# W14. The schemas ship with the skill (13 files, the transitive closure of wml.xsd).
# Measured packaging cost: +64,761 bytes / +1.8% of the built-in skills archive —
# NOT the "+1/3" the decision was nearly made on, which compared an uncompressed
# 984 KB against a compressed 3.4 MB.
SCHEMA_DIR = SKILL / "schemas"
SCHEMA_FILES = 13
# Parts report.docx has a grammar for. `word/document.xml` and friends; customXml and
# docProps have none, and saying so is what stops "valid" reading as "all checked".
VALIDATED_PARTS = 8



# Claims this host could not exercise. Reported separately and never folded into the
# pass count — a skip and a pass look identical at a glance, which is exactly how a
# wrong expectation once sat unnoticed for a month.
SKIPS: list[str] = []


def run_script(name: str, *args: str) -> subprocess.CompletedProcess:
    return subprocess.run([PY, str(SKILL / "scripts" / name), *map(str, args)],
                          capture_output=True, text=True, encoding="utf-8",
                          errors="replace", timeout=300)


# Windows encodes a captured stdout in the machine's ANSI code page, and Python only
# defaults to UTF-8 from 3.15 (PEP 686). Every report this skill prints carries
# Chinese, so on Windows every entry point died with UnicodeEncodeError the moment an
# agent captured its output. Reproduced here by forcing the code page, because the
# defect is invisible on macOS and Linux — it was found by a CI run, not by any gate.
WINDOWS_CODE_PAGES = ("cp1252", "cp936")     # western install, Chinese install


def run_script_encoded(name: str, code_page: str, *args: str):
    env = {**os.environ, "PYTHONIOENCODING": code_page}
    return subprocess.run([PY, str(SKILL / "scripts" / name), *map(str, args)],
                          capture_output=True, env=env, timeout=300)


def parts_of(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as z:
        return {n: z.read(n) for n in z.namelist() if not n.endswith("/")}


def part_order(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as z:
        return [i.filename for i in z.infolist() if not i.filename.endswith("/")]


def rewrite_zip(src: Path, dst: Path, mutate) -> None:
    """Copy a zip, passing each (name, bytes) through mutate; None drops the entry."""
    with zipfile.ZipFile(src) as zin, \
            zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = mutate(item.filename, zin.read(item.filename))
            if data is not None:
                zout.writestr(item, data)


def tree_of(path: Path, part: str = "word/document.xml"):
    from lxml import etree
    return etree.fromstring(parts_of(path)[part])


def paragraph_texts(path: Path) -> list[str]:
    """Paragraph text the way a reader that understands revisions sees it.

    Written here rather than imported from the skill: an assertion that measures the
    implementation with the implementation's own ruler agrees with it by
    construction. `<w:delText>` is excluded because deleted text is not in the
    document; text inside `<w:ins>` is included because inserted text is.
    """
    out = []
    for para in tree_of(path).iter(W + "p"):
        chunks = []
        for node in para.iter():
            tag = str(node.tag)
            if tag == W + "t" and not any(
                    str(a.tag) == W + "del" for a in node.iterancestors()):
                chunks.append(node.text or "")
            elif tag == W + "tab":
                chunks.append("\t")
            elif tag == W + "br":
                chunks.append("\n")
        out.append("".join(chunks))
    return out


def body_child_names(path: Path) -> list[str]:
    body = tree_of(path).find(W + "body")
    return [str(c.tag).rsplit("}", 1)[-1] for c in body]


def naive_run_replace(path: Path, needle: str) -> int:
    """The implementation everybody writes first, run for real on the same file.

    Not a description of a defect — the defect itself, executed, so the number it
    produces is measured rather than asserted. python-docx is used because that is
    what the reference implementations use; the result (0) is a property of the
    document, not of the library.
    """
    import docx
    hits = 0
    doc = docx.Document(str(path))
    for para in doc.paragraphs:
        for r in para.runs:
            hits += r.text.count(needle)
    return hits


def with_line_break(src: Path, dst: Path) -> None:
    """Put a `<w:br/>` in the middle of a phrase, so a search for it cannot match.

    This is not damage: a line break inside a sentence is ordinary. It is here
    because "the phrase is not found and the tool does not say why" is the single
    most common way a replace call wastes someone's afternoon.
    """
    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/document.xml":
            return data
        anchor = "<w:t>三季度</w:t></w:r>".encode()
        assert anchor in data, "the split run is not what this mutation expects"
        return data.replace(anchor, anchor + '<w:r><w:br/></w:r>'.encode(), 1)
    rewrite_zip(src, dst, mutate)


def without_numbering(src: Path, dst: Path) -> None:
    """Drop word/numbering.xml the way a package must be trimmed: all three things."""
    def mutate(name: str, data: bytes) -> bytes | None:
        if name == "word/numbering.xml":
            return None
        if name == "[Content_Types].xml":
            return data.replace(
                b'<Override PartName="/word/numbering.xml" ContentType="application/'
                b'vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>',
                b"")
        if name == "word/_rels/document.xml.rels":
            return data.replace(
                b'<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/'
                b'officeDocument/2006/relationships/numbering" Target="numbering.xml"/>',
                b"")
        if name == "word/document.xml":
            return data.replace(b'<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/>'
                                b"</w:numPr>", b"")
        return data
    rewrite_zip(src, dst, mutate)


def tall_table(src: Path, dst: Path, rows: int) -> None:
    """One table with `rows` rows: a report list of length ONE, holding megabytes."""
    cell = ('<w:tc><w:tcPr><w:tcW w:w="2500" w:type="dxa"/></w:tcPr><w:p><w:r><w:rPr>'
            '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="宋体" '
            'w:cs="Calibri"/></w:rPr><w:t>某个中文单元格内容占位</w:t></w:r></w:p></w:tc>')
    row = "<w:tr>" + cell * 3 + "</w:tr>"

    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/document.xml":
            return data
        return data.replace(b"</w:tbl>", (row * rows).encode() + b"</w:tbl>", 1)
    rewrite_zip(src, dst, mutate)


def long_document(src: Path, dst: Path, paragraphs: int) -> None:
    """A document with `paragraphs` paragraphs, for the stdout budget."""
    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/document.xml":
            return data
        one = ("<w:p><w:r><w:rPr><w:rFonts w:ascii=\"Calibri\" w:hAnsi=\"Calibri\" "
               "w:eastAsia=\"宋体\" w:cs=\"Calibri\"/></w:rPr>"
               "<w:t>批量段落用于测量输出预算</w:t></w:r></w:p>")
        return data.replace(b"<w:sectPr>", (one * paragraphs).encode() + b"<w:sectPr>", 1)
    rewrite_zip(src, dst, mutate)


# ── package plumbing, read straight from the zip ──────────────────────────────
CT_NS = "{http://schemas.openxmlformats.org/package/2006/content-types}"
PR_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"


def overrides_of(path: Path) -> set[str]:
    root = tree_of(path, "[Content_Types].xml")
    return {el.get("PartName").lstrip("/") for el in root
            if el.tag == CT_NS + "Override"}


def defaults_of(path: Path) -> dict[str, str]:
    root = tree_of(path, "[Content_Types].xml")
    return {(el.get("Extension") or "").lower(): el.get("ContentType")
            for el in root if el.tag == CT_NS + "Default"}


def rels_of(path: Path, owner: str = "word/_rels/document.xml.rels") -> dict[str, str]:
    if owner not in parts_of(path):
        return {}
    return {el.get("Id"): (el.get("Target") or "").lstrip("./")
            for el in tree_of(path, owner) if el.tag == PR_NS + "Relationship"}


def sect_refs(path: Path) -> list[dict]:
    """Every header/footer binding in the body's `w:sectPr`, with what it resolves to."""
    root = tree_of(path)
    body = root.find(W + "body")
    sect = next((el for el in body if el.tag == W + "sectPr"), None)
    if sect is None:
        return []
    rels = rels_of(path)
    out = []
    for el in sect:
        if el.tag not in (W + "headerReference", W + "footerReference"):
            continue
        rid = el.get(R_NS + "id")
        target = rels.get(rid)
        out.append({"kind": "header" if el.tag.endswith("headerReference") else "footer",
                    "type": el.get(W + "type") or "default", "rid": rid,
                    "part": f"word/{target}" if target else None})
    return out


def sect_tags(path: Path) -> list[str]:
    root = tree_of(path)
    body = root.find(W + "body")
    sect = next((el for el in body if el.tag == W + "sectPr"), None)
    return [str(el.tag).rsplit("}", 1)[-1] for el in sect] if sect is not None else []


def settings_tags(path: Path) -> list[str]:
    if "word/settings.xml" not in parts_of(path):
        return []
    return [str(el.tag).rsplit("}", 1)[-1] for el in tree_of(path, "word/settings.xml")]


def part_text(path: Path, part: str) -> str:
    return parts_of(path).get(part, b"").decode("utf-8", "replace")


def toc_entries(path: Path) -> list[dict]:
    """The cached TOC result, read out of the paragraphs that carry the field."""
    root = tree_of(path)
    body = root.find(W + "body")
    out = []
    for para in body.findall(W + "p"):
        ppr = para.find(W + "pPr")
        pstyle = ppr.find(W + "pStyle") if ppr is not None else None
        style = pstyle.get(W + "val") if pstyle is not None else None
        if not style or not style.startswith("TOC") or style == TOC_HEADING_STYLE:
            continue
        links = [h for h in para.iter(W + "hyperlink")]
        out.append({
            "style": style,
            "text": "".join(t.text or "" for t in para.iter(W + "t")),
            "anchors": [h.get(W + "anchor") for h in links],
            "link_text": "".join(t.text or "" for h in links for t in h.iter(W + "t")),
        })
    return out


def bookmark_names(path: Path) -> set[str]:
    return {b.get(W + "name") for b in tree_of(path).iter(W + "bookmarkStart")}


def style_numbering(path: Path) -> dict[str, str | None]:
    """`{styleId: numId}` for every style whose pPr carries a `w:numPr`."""
    out = {}
    for style in tree_of(path, "word/styles.xml").findall(W + "style"):
        ppr = style.find(W + "pPr")
        numpr = ppr.find(W + "numPr") if ppr is not None else None
        if numpr is None:
            continue
        num_id = numpr.find(W + "numId")
        out[style.get(W + "styleId")] = num_id.get(W + "val") if num_id is not None \
            else None
    return out


def style_based_on(path: Path) -> dict[str, str | None]:
    out = {}
    for style in tree_of(path, "word/styles.xml").findall(W + "style"):
        based = style.find(W + "basedOn")
        out[style.get(W + "styleId")] = based.get(W + "val") if based is not None else None
    return out


def abstract_levels(path: Path) -> list[dict]:
    """The levels of the LAST abstractNum, which is the one just authored."""
    root = tree_of(path, "word/numbering.xml")
    abstracts = root.findall(W + "abstractNum")
    if not abstracts:
        return []
    out = []
    for lvl in abstracts[-1].findall(W + "lvl"):
        pstyle = lvl.find(W + "pStyle")
        text = lvl.find(W + "lvlText")
        out.append({"ilvl": lvl.get(W + "ilvl"),
                    "style": pstyle.get(W + "val") if pstyle is not None else None,
                    "text": text.get(W + "val") if text is not None else None})
    return out


def drawings_of(path: Path) -> list[dict]:
    root = tree_of(path)
    rels = rels_of(path)
    out = []
    for drawing in root.iter(W + "drawing"):
        extent = next(iter(drawing.iter(WP_NS + "extent")), None)
        inner = next(iter(drawing.iter(A_NS + "ext")), None)
        blip = next(iter(drawing.iter(A_NS + "blip")), None)
        doc_pr = next(iter(drawing.iter(WP_NS + "docPr")), None)
        rid = blip.get(R_NS + "embed") if blip is not None else None
        target = rels.get(rid)
        out.append({
            "rid": rid, "part": f"word/{target}" if target else None,
            "alt": doc_pr.get("descr") if doc_pr is not None else None,
            "extent": [int(extent.get("cx")), int(extent.get("cy"))]
            if extent is not None else None,
            "inner": [int(inner.get("cx")), int(inner.get("cy"))]
            if inner is not None else None,
        })
    return out


def unbound_cjk_runs(path: Path) -> list[dict]:
    """CJK runs whose OWN `w:rFonts` does not state both faces, read from the XML.

    ⚠️ Computed here rather than taken from the skill's report, for the same reason
    V0 is: a fixture fact derived from the tool that is being judged by it agrees
    with that tool no matter what either of them does. This is the fourth time in
    this task that mistake had to be undone (059 §六·补二 / §六·补五).
    """
    def cjk(s: str) -> bool:
        return any("⺀" <= c <= "鿿" or "豈" <= c <= "﫿" for c in s)

    out = []
    root = tree_of(path)
    for p_index, para in enumerate(root.iter(W + "p")):
        for run in para.iter(W + "r"):
            text = "".join(t.text or "" for t in run.iter(W + "t"))
            if not cjk(text):
                continue
            rpr = run.find(W + "rPr")
            rfonts = rpr.find(W + "rFonts") if rpr is not None else None
            missing = [slot for slot in ("ascii", "eastAsia")
                       if rfonts is None or not (rfonts.get(W + slot)
                                                 or rfonts.get(W + slot + "Theme"))]
            if missing:
                out.append({"paragraph": p_index, "text": text[:12],
                            "missing": missing})
    return out


def fixture_facts() -> dict:
    """Facts about the fixtures, read from the FILES and from nothing else.

    ⚠️ The vacuity check must not consult any of the skill's own reports. A V0 that
    reads what a script said cannot tell "the fixture stopped exercising this" from
    "the implementation stopped doing it", and then it fires on defects that belong
    to other assertions — a mistake this repo has now made three times in one task
    (059 §六·补二). Everything below is computed by walking the XML here.
    """
    root = tree_of(REPORT)
    split = [p for p in root.iter(W + "p")
             if "".join(t.text or "" for t in p.iter(W + "t")) == SPLIT_PARAGRAPH]
    spans = 0
    if split:
        offset, hit = 0, set()
        for node in split[0].iter(W + "t"):
            text = node.text or ""
            start, end = offset, offset + len(text)
            at = SPLIT_PARAGRAPH.find(CROSS_RUN)
            if start < at + len(CROSS_RUN) and end > at:
                hit.add(id(node))
            offset = end
        spans = len(hit)

    unordered = tree_of(UNORDERED)
    body = unordered.find(W + "body")
    defects = 0
    kids = list(body)
    if kids and str(kids[-1].tag) != W + "sectPr":
        defects += 1
    for para in unordered.iter(W + "p"):
        children = [str(c.tag) for c in para]
        if W + "pPr" in children and children.index(W + "pPr") != 0:
            defects += 1
    for r in unordered.iter(W + "r"):
        children = [str(c.tag) for c in r]
        if W + "rPr" in children and children.index(W + "rPr") != 0:
            defects += 1

    outline_root = tree_of(OUTLINE)
    outline_styles = tree_of(OUTLINE, "word/styles.xml")
    heading_levels = []
    for para in outline_root.iter(W + "p"):
        ppr = para.find(W + "pPr")
        pstyle = ppr.find(W + "pStyle") if ppr is not None else None
        val = pstyle.get(W + "val") if pstyle is not None else None
        if val and val.startswith("Heading") and val[7:].isdigit():
            heading_levels.append(int(val[7:]))
    chart = CHART.read_bytes()

    return {
        "phrase_runs": spans,
        "naive_hits": naive_run_replace(REPORT, CROSS_RUN),
        # outline.docx: what it must NOT have is the point of it.
        "outline_paragraphs": len(list(outline_root.iter(W + "p"))),
        "outline_heading_levels": heading_levels,
        "outline_header_footer_parts": sum(1 for n in parts_of(OUTLINE)
                                           if n.startswith(("word/header",
                                                            "word/footer"))),
        "outline_sect_refs": len(sect_refs(OUTLINE)),
        "outline_png_default": "png" in defaults_of(OUTLINE),
        "outline_media_parts": sum(1 for n in parts_of(OUTLINE)
                                   if n.startswith("word/media/")),
        "outline_heading_styles": sum(
            1 for s in outline_styles.findall(W + "style")
            if (s.get(W + "styleId") or "").startswith("Heading")),
        # fontless.docx: three shapes, and one of them must resolve through a style
        # that already names a face — otherwise "keep what the document said" is
        # untestable and writing the fallback everywhere would look correct.
        "fontless_unbound": len(unbound_cjk_runs(FONTLESS)),
        "fontless_style_says": FONTLESS_STYLE_EA in part_text(FONTLESS,
                                                              "word/styles.xml"),
        "fontless_docdefaults_say": "rFonts" in part_text(
            FONTLESS, "word/styles.xml").split("</w:docDefaults>")[0],
        "fontless_kept_ascii": FONTLESS_KEPT_ASCII in part_text(FONTLESS,
                                                                "word/document.xml"),
        "report_unbound": len(unbound_cjk_runs(REPORT)),
        # chart.png: not square, and it states a density that is not 96 dpi.
        "chart_square": CHART_PX[0] == CHART_PX[1],
        "chart_has_phys": b"pHYs" in chart,
        "tables": len(list(root.iter(W + "tbl"))),
        "insertions": len(list(root.iter(W + "ins"))),
        "deletions": len(list(root.iter(W + "del"))),
        "comments": len(list(tree_of(REPORT, "word/comments.xml").iter(W + "comment"))),
        "custom_xml": any(n.startswith("customXml/") for n in parts_of(REPORT)),
        "header_footer_parts": sum(1 for n in parts_of(REPORT)
                                   if n.startswith(("word/header", "word/footer"))),
        "footer_fields": sum(1 for f in tree_of(REPORT, "word/footer1.xml")
                             .iter(W + "fldChar")
                             if f.get(W + "fldCharType") == "begin"),
        "list_paragraphs": len(list(root.iter(W + "numPr"))),
        "parts": len(parts_of(REPORT)),
        "order_defects": defects,
    }


# ── collect: run the real scripts once ────────────────────────────────────────
def collect(work: Path) -> dict:
    ctx: dict = {"fixture": fixture_facts()}

    # --- W1 read --------------------------------------------------------------
    r = run_script("docx_read.py", "--in", REPORT, "--outline", "--tables", "--text",
                   "--out", work / "document.json")
    read_report = json.loads((work / "document.json").read_text(encoding="utf-8"))
    ctx["read"] = {"exit": r.returncode, "report": read_report}

    # python-docx's own answer for the revision paragraph, measured rather than
    # described: it walks direct <w:r> children only, so a tracked insertion is
    # missing from what it returns. The skill must not agree with it.
    import docx as pydocx
    pd = pydocx.Document(str(REPORT))
    ctx["read"]["python_docx_paragraph_texts"] = [p.text for p in pd.paragraphs]

    # --- W2 replace, and the control everyone writes instead -------------------
    replaced = work / "replaced.docx"
    e = run_script("docx_edit.py", "--in", REPORT, "--out", replaced,
                   "--replace", f"{CROSS_RUN}={CROSS_RUN_NEW}",
                   "--replace", f"{PLACEHOLDER}={PLACEHOLDER_NEW}",
                   "--report", work / "replace.json")
    replace_report = json.loads((work / "replace.json").read_text(encoding="utf-8"))
    before, after = parts_of(REPORT), parts_of(replaced)
    ctx["replace"] = {
        "exit": e.returncode,
        "report": replace_report,
        "texts": paragraph_texts(replaced),
        "header_text": header_text(replaced),
        "parts_before": sorted(before),
        "parts_after": sorted(after),
        "identical": sorted(n for n in before if after.get(n) == before[n]),
        # The rPr of the run that received the replacement, before and after: an
        # edit that rebuilds the run loses its font binding and size.
        "rpr_before": rpr_of_first_run(REPORT, 1),
        "rpr_after": rpr_of_first_run(replaced, 1),
    }

    # --- a phrase broken by a line break --------------------------------------
    broken = work / "with-break.docx"
    with_line_break(REPORT, broken)
    nm = run_script("docx_edit.py", "--in", broken, "--out", work / "nm.docx",
                    "--replace", "三季度营业收入=第四季度营业收入",
                    "--report", work / "nearmiss.json")
    ctx["near_miss"] = {"exit": nm.returncode,
                        "report": json.loads((work / "nearmiss.json")
                                             .read_text(encoding="utf-8"))}

    # --- replacing inside tracked content -------------------------------------
    rv = run_script("docx_edit.py", "--in", REPORT, "--out", work / "rev.docx",
                    "--replace", f"{INSERTED}=归母净利润",
                    "--report", work / "rev.json")
    ctx["revision_edit"] = {"exit": rv.returncode,
                            "report": json.loads((work / "rev.json")
                                                 .read_text(encoding="utf-8"))}

    # --- headers ---------------------------------------------------------------
    hd = run_script("docx_edit.py", "--in", REPORT, "--out", work / "hdr.docx",
                    "--replace", "内部资料=公开资料", "--in-headers",
                    "--report", work / "hdr.json")
    ctx["headers"] = {"exit": hd.returncode,
                      "report": json.loads((work / "hdr.json")
                                           .read_text(encoding="utf-8")),
                      "header_text": header_text(work / "hdr.docx")}

    # --- W3 append -------------------------------------------------------------
    appended = work / "appended.docx"
    a = run_script("docx_edit.py", "--in", REPORT, "--out", appended,
                   "--append-paragraph", APPEND_TEXT, "--report", work / "append.json")
    ctx["append"] = {
        "exit": a.returncode,
        "report": json.loads((work / "append.json").read_text(encoding="utf-8")),
        "body_children": body_child_names(appended),
        "texts": paragraph_texts(appended),
        "fonts": appended_run_fonts(appended),
        "section_survives": bool(tree_of(appended).find(f"{W}body/{W}sectPr") is not None),
    }

    plain = work / "no-numbering.docx"
    without_numbering(REPORT, plain)
    nl = run_script("docx_edit.py", "--in", plain, "--out", work / "listed.docx",
                    "--append-paragraph", "列表项", "--list", "1")
    ctx["append"]["list_without_numbering"] = {
        "exit": nl.returncode, "stderr": nl.stderr.strip(),
        "wrote": (work / "listed.docx").exists()}
    ok = run_script("docx_edit.py", "--in", REPORT, "--out", work / "listed-ok.docx",
                    "--append-paragraph", "列表项", "--list", "1")
    ctx["append"]["list_with_numbering"] = {
        "exit": ok.returncode,
        "num_ids": num_ids_of(work / "listed-ok.docx")}

    # --- W13 unpack / pack -----------------------------------------------------
    p = run_script("docx_package.py", "--in", REPORT, "--unpack", work / "unpacked",
                   "--out", work / "rebuilt.docx", "--report", work / "package.json")
    pack_report = json.loads((work / "package.json").read_text(encoding="utf-8"))
    rebuilt = parts_of(work / "rebuilt.docx")
    ctx["package"] = {
        "exit": p.returncode,
        "report": pack_report,
        "identical": sum(1 for n in before if rebuilt.get(n) == before[n]),
        "order_before": part_order(REPORT),
        "order_after": part_order(work / "rebuilt.docx"),
    }

    # A part added by hand between unpack and pack must survive the round trip —
    # editing a part is the whole reason to unpack one.
    (work / "unpacked" / "word" / "added.xml").write_text(
        '<?xml version="1.0"?><added/>', encoding="utf-8")
    ap = run_script("docx_package.py", "--pack", work / "unpacked",
                    "--out", work / "with-added.docx")
    ctx["package"]["added_part"] = {
        "exit": ap.returncode,
        "present": "word/added.xml" in parts_of(work / "with-added.docx")}

    # Without the manifest the order is whatever the filesystem hands back, which is
    # the control that proves the manifest is what preserves it.
    (work / "unpacked" / "_manifest.json").unlink()
    (work / "unpacked" / "word" / "added.xml").unlink()
    nomf = run_script("docx_package.py", "--pack", work / "unpacked",
                      "--out", work / "no-manifest.docx")
    ctx["package"]["without_manifest_order"] = part_order(work / "no-manifest.docx") \
        if nomf.returncode == 0 else []

    ctx["package"]["traversal"] = traversal_refused(work)

    # --- W15 element order -----------------------------------------------------
    chk = run_script("docx_package.py", "--in", UNORDERED, "--check",
                     "--report", work / "check.json")
    check_report = json.loads((work / "check.json").read_text(encoding="utf-8"))
    fixed = work / "ordered.docx"
    fx = run_script("docx_package.py", "--in", UNORDERED, "--fix-order", "--out", fixed,
                    "--report", work / "order.json")
    fix_report = json.loads((work / "order.json").read_text(encoding="utf-8"))
    after_chk = run_script("docx_package.py", "--in", fixed, "--check",
                           "--report", work / "check2.json")
    unordered_parts = parts_of(UNORDERED)
    fixed_parts = parts_of(fixed)
    ctx["order"] = {
        "check_exit": chk.returncode,
        "check_findings": check_report["findings"],
        "fix_exit": fx.returncode,
        "fix_report": fix_report,
        "after_exit": after_chk.returncode,
        "after_findings": json.loads((work / "check2.json")
                                     .read_text(encoding="utf-8"))["findings"],
        "text_before": paragraph_texts(UNORDERED),
        "text_after": paragraph_texts(fixed),
        "other_parts_identical": sum(1 for n in unordered_parts
                                     if n != "word/document.xml"
                                     and fixed_parts.get(n) == unordered_parts[n]),
        "other_parts_total": len(unordered_parts) - 1,
    }

    # --- W5 template fill ------------------------------------------------------
    filled = work / "filled.docx"
    tf = run_script("docx_template.py", "--in", REPORT, "--out", filled,
                    "--set", f"{FILL_SPLIT}=示例科技有限公司",
                    "--set", f"{FILL_HEADER}=内部",
                    "--set", f"{FILL_TYPO}=错的键",
                    "--report", work / "fill.json")
    strict = run_script("docx_template.py", "--in", REPORT,
                        "--out", work / "strict.docx",
                        "--set", f"{FILL_SPLIT}=示例科技有限公司", "--strict")
    ctx["template"] = {
        "exit": tf.returncode,
        "report": json.loads((work / "fill.json").read_text(encoding="utf-8")),
        "header_text": header_text(filled),
        "texts": paragraph_texts(filled),
        "strict": {"exit": strict.returncode, "stderr": strict.stderr.strip(),
                   "wrote": (work / "strict.docx").exists()},
    }

    # --- W17 render ------------------------------------------------------------
    ctx["pdf"] = collect_pdf(work)

    # --- W6 / W7 tracked changes ------------------------------------------------
    ctx["revise"] = collect_revisions(work)

    # --- W8 comments -------------------------------------------------------------
    ctx["comment"] = collect_comments(work)

    # --- W9 / W10 / W11 / W16 the layout family ----------------------------------
    ctx["headerfooter"] = collect_headers(work)
    ctx["toc"] = collect_toc(work)
    ctx["image"] = collect_images(work)
    ctx["fonts"] = collect_fonts(work)
    ctx["styles"] = collect_styles(work)
    ctx["markdown"] = collect_markdown(work)
    ctx["validate"] = collect_validate(work)

    # --- W19 table presets --------------------------------------------------------
    ctx["tables"] = collect_tables(work)

    # --- W18 document diff ---------------------------------------------------------
    ctx["diff"] = collect_diff(work)

    # --- contracts -------------------------------------------------------------
    big = work / "big.docx"
    long_document(REPORT, big, SCALE_PARAGRAPHS)
    out = run_script("docx_read.py", "--in", big, "--outline")
    tall = work / "tall.docx"
    tall_table(REPORT, tall, SCALE_TABLE_ROWS)
    wide = run_script("docx_read.py", "--in", tall, "--tables")
    in_place = run_script("docx_edit.py", "--in", REPORT, "--out", REPORT,
                          "--replace", "a=b")
    missing = run_script("docx_read.py", "--in", work / "nope.docx")
    not_word = work / "notword.xlsx"
    make_xlsx_like(not_word)
    wrong_kind = run_script("docx_read.py", "--in", not_word)
    ctx["contracts"] = {
        "stdout_bytes": len(out.stdout.encode()),
        "stdout_exit": out.returncode,
        "tall_stdout_bytes": len(wide.stdout.encode()),
        "tall_exit": wide.returncode,
        "tall_stdout": wide.stdout,
        "in_place": {"exit": in_place.returncode, "stderr": in_place.stderr.strip()},
        "missing": {"exit": missing.returncode, "stderr": missing.stderr.strip(),
                    "traceback": "Traceback" in missing.stderr},
        "wrong_kind": {"exit": wrong_kind.returncode,
                       "stderr": wrong_kind.stderr.strip(),
                       "traceback": "Traceback" in wrong_kind.stderr},
    }

    # --- every entry point survives a Windows code page -------------------------
    entries = sorted(p.name for p in (SKILL / "scripts").glob("docx_*.py"))
    probes = {
        "docx_read.py": ["--in", REPORT],
        "docx_header.py": ["--in", REPORT, "--list"],
        "docx_toc.py": ["--in", OUTLINE, "--list"],
        "docx_image.py": ["--in", REPORT, "--list"],
        "docx_fonts.py": ["--in", REPORT, "--check"],
        "docx_package.py": ["--in", REPORT, "--check"],
        "docx_revise.py": ["--in", REPORT, "--list"],
        "docx_comment.py": ["--in", REPORT, "--list"],
        "docx_template.py": ["--in", REPORT, "--list"],
        "docx_style.py": ["--in", REPORT, "--list"],
        # --help: this one's inputs are a Markdown file rather than a fixture, and
        # the encoding path is crossed at import of docxcommon either way.
        "docx_from_md.py": ["--help"],
        "docx_validate.py": ["--in", REPORT],
        "docx_table.py": ["--in", REPORT, "--measure"],
        "docx_diff.py": ["--a", REPORT, "--b", REPORT],
        "docx_edit.py": ["--in", REPORT, "--out", work / "cp.docx", "--replace", "a=b"],
        # --help rather than a conversion: it still crosses the encoding path (the
        # reconfigure happens when docxcommon is imported, before argparse runs) and
        # its help text carries the module docstring's Chinese, without needing
        # LibreOffice or spending seconds on a render.
        "docx_pdf.py": ["--help"],
    }
    runs = {}
    for page in WINDOWS_CODE_PAGES:
        for name in entries:
            args = probes.get(name)
            if args is None:
                continue
            r = run_script_encoded(name, page, *args)
            runs[f"{page}:{name}"] = {
                "exit": r.returncode,
                "cjk_in_stdout": any("一" <= c <= "鿿"
                                     for c in r.stdout.decode("utf-8", "replace")),
                "stderr": r.stderr.decode("utf-8", "replace")[:200],
            }
    # The vacuity guard: if a bare print of Chinese under cp1252 does NOT fail, this
    # host is not reproducing the condition and every result above proves nothing.
    bare = subprocess.run([PY, "-c", "print('中文')"], capture_output=True,
                          env={**os.environ, "PYTHONIOENCODING": "cp1252"}, timeout=60)
    ctx["encoding"] = {
        "entry_points": entries,
        "probed": sorted(probes),
        "runs": runs,
        "bare_print_exit": bare.returncode,
    }

    # --- the fixtures are byte-reproducible ------------------------------------
    regen = work / "regen"
    rg = subprocess.run([PY, str(FIXTURES / "make_fixtures.py"), "--out-dir", str(regen)],
                        capture_output=True, text=True, timeout=120)
    ctx["fixtures"] = {
        "exit": rg.returncode,
        "identical": {name: (regen / name).is_file()
                      and (regen / name).read_bytes() == (FIXTURES / name).read_bytes()
                      for name in ("report.docx", "revised.docx",
                                   "unordered.docx")},
    }
    return ctx


# ── W9: headers and footers ───────────────────────────────────────────────────
def collect_headers(work: Path) -> dict:
    out: dict = {}
    letterhead = work / "letterhead.docx"
    r = run_script("docx_header.py", "--in", OUTLINE, "--out", letterhead,
                   "--header", HEADER_NEW, "--footer", FOOTER_NEW, "--page-number",
                   "--report", work / "header.json")
    out["create"] = {
        "exit": r.returncode,
        "report": json.loads((work / "header.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(letterhead)),
        "overrides": sorted(overrides_of(letterhead)),
        "rel_targets": sorted(rels_of(letterhead).values()),
        "refs": sect_refs(letterhead),
        "footer_xml": part_text(letterhead, "word/footer1.xml"),
    }

    first = work / "first.docx"
    fr = run_script("docx_header.py", "--in", letterhead, "--out", first,
                    "--type", "first", "--header", FIRST_HEADER,
                    "--report", work / "first.json")
    out["first"] = {
        "exit": fr.returncode,
        "report": json.loads((work / "first.json").read_text(encoding="utf-8")),
        "sect_tags": sect_tags(first),
        "refs": sect_refs(first),
    }

    even = work / "even.docx"
    er = run_script("docx_header.py", "--in", first, "--out", even,
                    "--type", "even", "--header", EVEN_HEADER,
                    "--report", work / "even.json")
    out["even"] = {
        "exit": er.returncode,
        "report": json.loads((work / "even.json").read_text(encoding="utf-8")),
        "settings_tags": settings_tags(even),
        "sect_tags": sect_tags(even),
        "refs": sect_refs(even),
    }

    removed = work / "removed.docx"
    rr = run_script("docx_header.py", "--in", even, "--out", removed,
                    "--type", "first", "--remove", "header",
                    "--report", work / "removed.json")
    out["remove"] = {
        "exit": rr.returncode,
        "report": json.loads((work / "removed.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(removed)),
        "overrides": sorted(overrides_of(removed)),
        "rel_targets": sorted(rels_of(removed).values()),
        "sect_tags": sect_tags(removed),
        "refs": sect_refs(removed),
    }
    return out


# ── W10: contents and outline numbering ───────────────────────────────────────
def collect_toc(work: Path) -> dict:
    contents = work / "contents.docx"
    r = run_script("docx_toc.py", "--in", OUTLINE, "--out", contents,
                   "--toc", "--levels", "3", "--title", TOC_TITLE,
                   "--outline-numbering", "--report", work / "toc.json")
    out = {
        "exit": r.returncode,
        "report": json.loads((work / "toc.json").read_text(encoding="utf-8")),
        "paragraphs": len(paragraph_texts(contents)),
        "entries": toc_entries(contents),
        "bookmarks": sorted(bookmark_names(contents)),
        "document_xml": part_text(contents, "word/document.xml"),
        "settings_tags": settings_tags(contents),
        "style_numbering": style_numbering(contents),
        "style_based_on": style_based_on(contents),
        "abstract_levels": abstract_levels(contents),
    }
    # The other half of the product decision: --no-cache writes the field and no
    # result, which is a legitimate mode and must be distinguishable from a bug.
    bare = work / "bare-toc.docx"
    nr = run_script("docx_toc.py", "--in", OUTLINE, "--out", bare, "--toc",
                    "--no-cache", "--report", work / "bare-toc.json")
    # `entries` counts paragraphs carrying a TOC style — and --no-cache writes ONE
    # of those to hold the field markers, so a count of zero was never the right
    # expectation. What "no result" actually means is that no entry carries a
    # heading: the cached result is exactly the hyperlinked entries.
    out["no_cache"] = {"exit": nr.returncode,
                       "paragraphs": len(paragraph_texts(bare)),
                       "cached_entries": sum(1 for e in toc_entries(bare)
                                             if e["anchors"])}
    # Running it twice must refuse: two contents pages look identical until someone
    # updates them, so the second one is the one nobody notices.
    again = run_script("docx_toc.py", "--in", contents, "--out", work / "twice.docx",
                       "--toc")
    out["twice"] = {"exit": again.returncode, "stderr": again.stderr.strip(),
                    "wrote": (work / "twice.docx").exists()}
    return out


# ── W11: pictures ─────────────────────────────────────────────────────────────
def collect_images(work: Path) -> dict:
    sized = work / "illustrated.docx"
    r = run_script("docx_image.py", "--in", OUTLINE, "--out", sized,
                   "--insert", CHART, "--after", OUTLINE_BODY,
                   "--width-cm", str(IMAGE_WIDTH_CM), "--alt", "季度收入趋势图",
                   "--report", work / "image.json")
    out = {
        "exit": r.returncode,
        "report": json.loads((work / "image.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(sized)),
        "defaults": defaults_of(sized),
        "drawings": drawings_of(sized),
        "paragraphs": len(paragraph_texts(sized)),
        "body_children": body_child_names(sized),
    }
    # No --width-cm: the size then comes from the file's own declared density, and
    # that is the number a 96-dpi assumption gets wrong.
    intrinsic = work / "intrinsic.docx"
    ir = run_script("docx_image.py", "--in", OUTLINE, "--out", intrinsic,
                    "--insert", CHART, "--report", work / "intrinsic.json")
    out["intrinsic"] = {
        "exit": ir.returncode,
        "drawings": drawings_of(intrinsic),
        "report": json.loads((work / "intrinsic.json").read_text(encoding="utf-8")),
    }
    # Replace with the SAME format: the relationship and content type already say the
    # right thing, so only bytes may change.
    same = work / "replaced-same.docx"
    bigger = work / "bigger.png"
    bigger.write_bytes(CHART.read_bytes() + b"\x00" * 8)   # different bytes, same format
    sr = run_script("docx_image.py", "--in", sized, "--out", same,
                    "--replace", "0", "--with", bigger, "--report", work / "same.json")
    out["replace_same"] = {
        "exit": sr.returncode,
        "report": json.loads((work / "same.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(same)),
        "drawings": drawings_of(same),
        "media_bytes": {n: len(d) for n, d in parts_of(same).items()
                        if n.startswith("word/media/")},
    }
    # Replace with a DIFFERENT format: the old media part has to go — all three of
    # bytes, Override/Default and Relationship — and a new Default declared.
    other = work / "replaced-gif.docx"
    gif = work / "tiny.gif"
    gif.write_bytes(TINY_GIF)
    gr = run_script("docx_image.py", "--in", sized, "--out", other,
                    "--replace", "0", "--with", gif, "--report", work / "gif.json")
    out["replace_other"] = {
        "exit": gr.returncode,
        "report": json.loads((work / "gif.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(other)),
        "defaults": defaults_of(other),
        "drawings": drawings_of(other),
        "rel_targets": sorted(rels_of(other).values()),
    }
    return out


# ── W16: font bindings ────────────────────────────────────────────────────────
def collect_fonts(work: Path) -> dict:
    bound = work / "bound.docx"
    r = run_script("docx_fonts.py", "--in", FONTLESS, "--out", bound, "--fix",
                   "--east-asia", FONTLESS_FALLBACK_EA, "--ascii", "Calibri",
                   "--report", work / "fonts.json")
    out = {
        "exit": r.returncode,
        "report": json.loads((work / "fonts.json").read_text(encoding="utf-8")),
        "document_xml": part_text(bound, "word/document.xml"),
        "still_unbound": unbound_cjk_runs(bound),
        "parts_changed": sorted(n for n, d in parts_of(bound).items()
                                if parts_of(FONTLESS).get(n) != d),
    }
    clean = work / "clean.json"
    cr = run_script("docx_fonts.py", "--in", REPORT, "--check", "--report", clean)
    out["check_clean"] = {
        "exit": cr.returncode,
        "report": json.loads(clean.read_text(encoding="utf-8")),
        "wrote_anything": sorted(p.name for p in work.glob("*.docx")
                                 if p.name.startswith("check")),
    }
    sr = run_script("docx_fonts.py", "--in", FONTLESS, "--out", work / "strict.docx",
                    "--fix", "--strict")
    out["strict"] = {"exit": sr.returncode, "stderr": sr.stderr.strip(),
                     "wrote": (work / "strict.docx").exists()}
    return out


# ── W12: styles ───────────────────────────────────────────────────────────────
def styles_of(path: Path) -> dict[str, dict]:
    """Every `w:style`, read from the file: id -> {name, basedOn, children, rPr}."""
    out = {}
    for style in tree_of(path, "word/styles.xml").findall(W + "style"):
        kids = [str(c.tag).rsplit("}", 1)[-1] for c in style]
        rpr = style.find(W + "rPr")
        out[style.get(W + "styleId")] = {
            "children": kids,
            "custom": style.get(W + "customStyle") == "1",
            "based_on": (style.find(W + "basedOn").get(W + "val")
                         if style.find(W + "basedOn") is not None else None),
            "rpr": [str(c.tag).rsplit("}", 1)[-1] for c in rpr] if rpr is not None else [],
            "sz": (rpr.find(W + "sz").get(W + "val")
                   if rpr is not None and rpr.find(W + "sz") is not None else None),
        }
    return out


def paragraph_styles(path: Path) -> list[str | None]:
    out = []
    for para in tree_of(path).find(W + "body").findall(W + "p"):
        ppr = para.find(W + "pPr")
        ps = ppr.find(W + "pStyle") if ppr is not None else None
        out.append(ps.get(W + "val") if ps is not None else None)
    return out


def collect_styles(work: Path) -> dict:
    out: dict = {}
    made = work / "restyled.docx"
    r = run_script("docx_style.py", "--in", REPORT, "--out", made,
                   "--set", NEW_STYLE, "--name", "Body Small", "--based-on", "Normal",
                   "--size", "9", "--east-asia", "宋体", "--font", "Calibri",
                   "--report", work / "style.json")
    out["create"] = {
        "exit": r.returncode,
        "report": json.loads((work / "style.json").read_text(encoding="utf-8")),
        "styles": styles_of(made),
        "parts_changed": sorted(n for n, d in parts_of(made).items()
                                if parts_of(REPORT).get(n) != d),
    }
    # Modifying one that exists: refused without --overwrite, and the refusal has to
    # say how many paragraphs it would have repainted.
    bare = run_script("docx_style.py", "--in", REPORT, "--out", work / "no.docx",
                      "--set", "Heading1", "--color", "1F5CA8")
    out["no_overwrite"] = {"exit": bare.returncode, "stderr": bare.stderr.strip(),
                           "wrote": (work / "no.docx").exists()}
    ow = work / "overwritten.docx"
    o = run_script("docx_style.py", "--in", REPORT, "--out", ow, "--set", "Heading1",
                   "--overwrite", "--color", "FF0000", "--underline", "--italic",
                   "--size", "20", "--report", work / "ow.json")
    out["overwrite"] = {
        "exit": o.returncode,
        "report": json.loads((work / "ow.json").read_text(encoding="utf-8")),
        "styles": styles_of(ow),
    }
    # Deleting one still in use, with and without --reassign.
    d = run_script("docx_style.py", "--in", REPORT, "--out", work / "del.docx",
                   "--delete", "Heading1")
    out["delete_in_use"] = {"exit": d.returncode, "stderr": d.stderr.strip(),
                            "wrote": (work / "del.docx").exists()}
    re_ = work / "reassigned.docx"
    rr = run_script("docx_style.py", "--in", REPORT, "--out", re_, "--delete",
                    "Heading1", "--reassign", "Normal", "--report", work / "re.json")
    out["reassign"] = {
        "exit": rr.returncode,
        "report": json.loads((work / "re.json").read_text(encoding="utf-8")),
        "styles": styles_of(re_),
        "paragraph_styles": paragraph_styles(re_),
    }
    cyc = run_script("docx_style.py", "--in", REPORT, "--out", work / "cyc.docx",
                     "--set", "Heading1", "--overwrite", "--based-on", "Heading1")
    out["cycle"] = {"exit": cyc.returncode, "stderr": cyc.stderr.strip(),
                    "wrote": (work / "cyc.docx").exists()}
    return out



# ── W18: diffing two documents ────────────────────────────────────────────────
def doc_signature(path: Path) -> list[tuple]:
    """(part, style, text) for every paragraph, headers and footers included.

    This file's OWN reading of what a document says, written here rather than taken
    from the script's report, because the round trip is the whole claim and a claim
    checked with the implementation's own ruler is not checked.
    """
    out = []
    parts = parts_of(path)
    names = [DOC_PART] + sorted(n for n in parts
                                if n.startswith(("word/header", "word/footer"))
                                and n.endswith(".xml"))
    for name in names:
        root = tree_of(path, name)
        for para in root.iter(W + "p"):
            ppr = para.find(W + "pPr")
            style = None
            if ppr is not None:
                node = ppr.find(W + "pStyle")
                style = node.get(W + "val") if node is not None else None
            chunks = []
            for node in para.iter():
                tag = str(node.tag)
                if tag == W + "t" and not any(str(a.tag) == W + "del"
                                              for a in node.iterancestors()):
                    chunks.append(node.text or "")
                elif tag == W + "tab":
                    chunks.append("\t")
                elif tag == W + "br":
                    chunks.append("\n")
            out.append((name, style, "".join(chunks)))
    return out


def _edit_document(src: Path, dst: Path, mutate) -> None:
    """Copy `src` to `dst`, passing its word/document.xml tree through `mutate`."""
    from lxml import etree
    root = tree_of(src)
    mutate(root)
    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    rewrite_zip(src, dst, lambda n, d: blob if n == DOC_PART else d)


def diff_fixture_facts() -> dict:
    """How report.docx and revised.docx actually differ, measured off the two files.

    V0's subject: the A/B pair is a committed fixture, and every W18 assertion is
    written for the six changes it is supposed to carry. If a later edit to
    report.docx quietly drops one of them, the assertions go looking for a difference
    that is not there.
    """
    a, b = doc_signature(REPORT), doc_signature(REVISED)
    a_texts = [text for part, _, text in a if part == DOC_PART]
    b_texts = [text for part, _, text in b if part == DOC_PART]
    # Keyed by TEXT, not by position: the restyled paragraph is the one whose words
    # did not change, which is exactly what makes a style-only change checkable.
    a_style = {text: style for part, style, text in a if part == DOC_PART}
    b_style = {text: style for part, style, text in b if part == DOC_PART}
    return {
        # Membership, not a set difference. A set difference calls a paragraph whose
        # TEXT was edited both "removed" and "added", so it cannot tell an edit from
        # a deletion — which is the distinction this whole fixture exists to provide.
        "removed_from_a": DIFF_REMOVED in a_texts and DIFF_REMOVED not in b_texts,
        "added_in_b": DIFF_ADDED not in a_texts and DIFF_ADDED in b_texts,
        "restyled": sorted(t for t in set(a_style) & set(b_style)
                           if a_style[t] != b_style[t]),
        "header_a": [text for part, _, text in a if part == DIFF_HEADER_PART],
        "header_b": [text for part, _, text in b if part == DIFF_HEADER_PART],
        "same_parts": sorted(parts_of(REPORT)) == sorted(parts_of(REVISED)),
    }


def make_noise_only(src: Path, dst: Path) -> None:
    """B: everything this capability promises to ignore, and NOT one real change.

    rsid attributes, a spell-check marker, a bookmark and a language tag — the four
    things Word sprinkles through a file on an ordinary save. If any of these counts
    as a difference, the report is one nobody will read.
    """
    from lxml import etree

    def mutate(root) -> None:
        for i, para in enumerate(root.iter(W + "p")):
            para.set(W + "rsidR", f"00{i:06X}")
            para.set(W + "rsidRDefault", f"00{i:06X}")
            proof = etree.Element(W + "proofErr")
            proof.set(W + "type", "spellStart")
            para.insert(0, proof)
            for r in para.findall(W + "r"):
                rpr = r.find(W + "rPr")
                if rpr is None:
                    continue
                lang = etree.SubElement(rpr, W + "lang")
                lang.set(W + "val", "zh-CN")
        body = root.find(W + "body")
        start = etree.Element(W + "bookmarkStart")
        start.set(W + "id", "900")
        start.set(W + "name", "_noise")
        end = etree.Element(W + "bookmarkEnd")
        end.set(W + "id", "900")
        body.insert(0, start)
        body.insert(1, end)
    _edit_document(src, dst, mutate)


def make_row_added(src: Path, dst: Path) -> None:
    """B: a table row appears — a real difference this capability cannot redline."""
    import copy as _copy

    def mutate(root) -> None:
        tbl = next(root.iter(W + "tbl"))
        rows = tbl.findall(W + "tr")
        extra = _copy.deepcopy(rows[-1])
        for node in extra.iter(W + "t"):
            node.text = "新增行"
        rows[-1].addnext(extra)
    _edit_document(src, dst, mutate)


def make_moved(src: Path, dst: Path) -> None:
    """B: a paragraph is where a different one used to be, with nothing rewritten."""
    def mutate(root) -> None:
        body = root.find(W + "body")
        paras = [p for p in body if str(p.tag) == W + "p"]
        mover = paras[DIFF_MOVED_FROM]
        body.remove(mover)
        paras[DIFF_MOVED_TO].addprevious(mover)
    _edit_document(src, dst, mutate)


def with_image(src: Path, dst: Path) -> None:
    """A of the picture pair: report.docx with a chart in it."""
    r = run_script("docx_image.py", "--in", src, "--out", dst,
                   "--insert", CHART, "--width-cm", "6", "--alt", "季度收入趋势图")
    if not dst.is_file():
        raise RuntimeError(f"could not build the picture fixture: {r.stderr[:200]}")


def repaint_media(src: Path, dst: Path) -> None:
    """B: the same picture slot holding different bytes."""
    from PIL import Image
    import io
    buf = io.BytesIO()
    Image.new("RGB", (60, 40), (10, 90, 200)).save(buf, format="PNG")
    replacement = buf.getvalue()

    def mutate(name: str, data: bytes) -> bytes:
        return replacement if name.startswith("word/media/") else data
    rewrite_zip(src, dst, mutate)


def collect_diff(work: Path) -> dict:
    out: dict = {}

    def compare(tag: str, left: Path, right: Path, *extra: str) -> dict:
        redline = work / f"redline-{tag}.docx"
        report = work / f"diff-{tag}.json"
        r = run_script("docx_diff.py", "--a", left, "--b", right,
                       "--redline", redline, "--report", report, *extra)
        payload = json.loads(report.read_text(encoding="utf-8")) \
            if report.is_file() else {}
        return {"exit": r.returncode, "stderr": r.stderr.strip(),
                "report": payload, "wrote": redline.is_file(),
                "redline": str(redline)}

    b = REVISED
    out["fixture"] = diff_fixture_facts()
    out["main"] = compare("main", REPORT, b)

    # The round trip, reproduced from OUTSIDE docx_diff.py: its own check is in-process
    # and could agree with itself. This one resolves the redline with the shipped W7
    # implementation and reads the result with this file's own signature reader.
    ids = out["main"]["report"].get("revision_ids") or []
    out["signatures"] = {"a": doc_signature(REPORT), "b": doc_signature(b)}
    for mode in ("accept", "reject"):
        resolved = work / f"{mode}ed.docx"
        args = ["--in", out["main"]["redline"], "--out", resolved]
        for rid in ids:
            args += [f"--{mode}-id", str(rid)]
        rr = run_script("docx_revise.py", *args)
        out[mode] = {"exit": rr.returncode,
                     "signature": doc_signature(resolved) if resolved.is_file() else []}

    noise = work / "variant-noise.docx"
    make_noise_only(REPORT, noise)
    out["noise"] = compare("noise", REPORT, noise)

    rows = work / "variant-rows.docx"
    make_row_added(REPORT, rows)
    out["rows"] = compare("rows", REPORT, rows)
    strict = run_script("docx_diff.py", "--a", REPORT, "--b", rows,
                        "--redline", work / "strict.docx", "--strict")
    out["rows_strict"] = {"exit": strict.returncode, "stderr": strict.stderr.strip(),
                          "wrote": (work / "strict.docx").exists()}

    moved = work / "variant-moved.docx"
    make_moved(REPORT, moved)
    out["moved"] = compare("moved", REPORT, moved)

    illustrated = work / "with-image.docx"
    with_image(REPORT, illustrated)
    repainted = work / "repainted.docx"
    repaint_media(illustrated, repainted)
    out["image"] = compare("image", illustrated, repainted)

    validated = run_script("docx_validate.py", "--in", out["main"]["redline"],
                           "--report", work / "redline-valid.json")
    out["redline_valid"] = {
        "exit": validated.returncode,
        "report": json.loads((work / "redline-valid.json").read_text(encoding="utf-8"))
        if (work / "redline-valid.json").is_file() else {}}
    listed = run_script("docx_revise.py", "--in", out["main"]["redline"], "--list",
                        "--report", work / "redline-revisions.json")
    out["redline_revisions"] = json.loads(
        (work / "redline-revisions.json").read_text(encoding="utf-8")) \
        if (work / "redline-revisions.json").is_file() else {}
    out["redline_parts_changed"] = sorted(
        n for n, d in parts_of(Path(out["main"]["redline"])).items()
        if parts_of(REPORT).get(n) != d)
    return out


# ── W19: table presets ────────────────────────────────────────────────────────
def wide_table(src: Path, dst: Path) -> None:
    """report.docx with its table's cells replaced by WIDE_ROWS."""
    from lxml import etree
    parts = parts_of(src)
    root = etree.fromstring(parts["word/document.xml"])
    tbl = next(root.iter(W + "tbl"))
    for tr, row in zip(tbl.findall(W + "tr"), WIDE_ROWS):
        for tc, text in zip(tr.findall(W + "tc"), row):
            nodes = list(tc.iter(W + "t"))
            for extra in nodes[1:]:
                extra.text = ""
            if nodes:
                nodes[0].text = text
    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                          standalone=True)
    rewrite_zip(src, dst,
                lambda name, data: blob if name == "word/document.xml" else data)


def naive_fit(src: Path, dst: Path) -> list[int]:
    """Size the columns by `len()` — the implementation everybody writes first.

    Not a description of the defect but the defect itself, executed with the same
    arithmetic and the same constants as the real one, so the only difference between
    the two arms is which function measured the text.
    """
    from lxml import etree
    parts = parts_of(src)
    root = etree.fromstring(parts["word/document.xml"])
    tbl = next(root.iter(W + "tbl"))
    widths = [n * WIDE_DXA_PER_CELL + 2 * WIDE_MARGIN["finance"] for n in WIDE_LENS]
    grid = tbl.find(W + "tblGrid")
    for old in grid.findall(W + "gridCol"):
        grid.remove(old)
    for w in widths:
        etree.SubElement(grid, W + "gridCol").set(W + "w", str(w))
    for tr in tbl.findall(W + "tr"):
        for tc, w in zip(tr.findall(W + "tc"), widths):
            tcw = tc.find(W + "tcPr").find(W + "tcW")
            tcw.set(W + "w", str(w))
    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                          standalone=True)
    rewrite_zip(src, dst,
                lambda name, data: blob if name == "word/document.xml" else data)
    return widths


def table_facts(path: Path, index: int = 0) -> dict:
    """What a table measurably IS, read out of the XML by THIS file's own reader.

    Deliberately not the skill's `fingerprint()`: an assertion that measures an
    implementation with that implementation's ruler agrees with it by construction,
    which this suite has already been caught doing once (V0 reading a script report).
    """
    tbl = list(tree_of(path).iter(W + "tbl"))[index]
    tblpr = tbl.find(W + "tblPr")
    borders = tblpr.find(W + "tblBorders") if tblpr is not None else None

    def weight(edge: str) -> int:
        node = borders.find(W + edge) if borders is not None else None
        if node is None:
            return -1
        return 0 if node.get(W + "val") == "none" else int(node.get(W + "sz") or 0)

    margin = -1
    if tblpr is not None:
        mar = tblpr.find(W + "tblCellMar")
        left = mar.find(W + "left") if mar is not None else None
        if left is not None:
            margin = int(left.get(W + "w"))
    rows = tbl.findall(W + "tr")
    repeats, fills = [], []
    for i, tr in enumerate(rows):
        trpr = tr.find(W + "trPr")
        if trpr is not None and trpr.find(W + "tblHeader") is not None:
            repeats.append(i)
        first = tr.find(W + "tc")
        tcpr = first.find(W + "tcPr") if first is not None else None
        shd = tcpr.find(W + "shd") if tcpr is not None else None
        fills.append(shd.get(W + "fill") if shd is not None else None)
    rule = -1
    if rows:
        tc = rows[0].find(W + "tc")
        tcpr = tc.find(W + "tcPr") if tc is not None else None
        tcb = tcpr.find(W + "tcBorders") if tcpr is not None else None
        bottom = tcb.find(W + "bottom") if tcb is not None else None
        if bottom is not None:
            rule = 0 if bottom.get(W + "val") == "none" else int(bottom.get(W + "sz"))
    grid = tbl.find(W + "tblGrid")
    layout = tblpr.find(W + "tblLayout") if tblpr is not None else None
    return {
        "borders": {e: weight(e) for e in
                    ("top", "left", "bottom", "right", "insideH", "insideV")},
        "cell_margin": margin,
        "repeat_rows": repeats,
        "row_fills": fills,
        "header_rule": rule,
        "layout": layout.get(W + "type") if layout is not None else None,
        "widths": [int(g.get(W + "w")) for g in grid.findall(W + "gridCol")]
        if grid is not None else [],
    }


def rendered_lines(pdf: Path, page: int = 0) -> list[str]:
    """One string per text baseline on a rendered page."""
    import pdfplumber
    with pdfplumber.open(str(pdf)) as doc:
        if page >= len(doc.pages):
            return []
        tops: dict[int, list] = {}
        for word in doc.pages[page].extract_words():
            tops.setdefault(round(word["top"]), []).append((word["x0"], word["text"]))
    return ["".join(t for _, t in sorted(v)) for _, v in sorted(tops.items())]


def rendered_pages(pdf: Path) -> int:
    import pdfplumber
    with pdfplumber.open(str(pdf)) as doc:
        return len(doc.pages)


def collect_tables(work: Path) -> dict:
    out: dict = {"presets": {}}
    base = work / "wide.docx"
    wide_table(REPORT, base)
    # How many paragraphs of the document that actually gets RENDERED contain each
    # probe string. Exactly one is the only safe answer: Q5 asks "did this cell's
    # text survive on one line" by looking for it among the rendered lines, so a
    # second home for the same string answers the question off the wrong line.
    # Counted against the wide document rather than report.docx — the cells are
    # replaced, so report.docx's own "+1.2pt" is not on the page Q5 reads.
    out["probe_strings_unique"] = {
        s: sum(1 for t in paragraph_texts(base) if s in t)
        for row in WIDE_ROWS for s in row
    }

    listed = run_script("docx_table.py", "--list-presets")
    out["listed"] = {"exit": listed.returncode,
                     "presets": json.loads(listed.stdout)["presets"]
                     if listed.returncode == 0 else []}

    for name in TABLE_PRESETS:
        made = work / f"preset-{name}.docx"
        r = run_script("docx_table.py", "--in", base, "--out", made,
                       "--preset", name, "--report", work / f"{name}.json")
        report = json.loads((work / f"{name}.json").read_text(encoding="utf-8"))
        out["presets"][name] = {
            "exit": r.returncode,
            "facts": table_facts(made),
            # The skill's own reading of the same file. Compared against `facts`
            # rather than trusted: a report that agrees with its intent proves
            # nothing, and this skill family has shipped that defect once already.
            "reported": report["fingerprints"][0],
            "parts_changed": report["parts_changed"],
            "path": str(made),
        }

    # The refusals.
    empty = run_script("docx_table.py", "--in", OUTLINE, "--out", work / "no.docx",
                       "--preset", "grid")
    out["no_tables"] = {"exit": empty.returncode, "stderr": empty.stderr.strip(),
                        "wrote": (work / "no.docx").exists(),
                        "traceback": "Traceback" in empty.stderr}
    bad = run_script("docx_table.py", "--in", base, "--out", work / "bad.docx",
                     "--preset", "grid", "--table", "7")
    out["bad_index"] = {"exit": bad.returncode, "stderr": bad.stderr.strip(),
                        "wrote": (work / "bad.docx").exists()}

    # The naive arm: same table, same constants, `len()` instead of display width.
    naive = work / "naive-fit.docx"
    out["naive_widths"] = naive_fit(work / "preset-finance.docx", naive)

    sys.path.insert(0, str(SKILL / "scripts"))
    from office.soffice import find_soffice
    if not find_soffice():
        SKIPS.append("W19 render (Q3, Q5): LibreOffice is not installed on this host, "
                     "so neither the header-repeat measurement nor the column-width "
                     "one could run, and nor could their negative controls. CI runs "
                     "them on Linux, where libreoffice-writer is installed")
        out["render"] = {"skipped": "no LibreOffice"}
        return out

    render: dict = {}
    for tag, src in (("fitted", work / "preset-finance.docx"), ("naive", naive)):
        pdf = work / f"{tag}.pdf"
        run_script("docx_pdf.py", "--in", src, "--out", pdf)
        lines = rendered_lines(pdf) if pdf.is_file() else []
        render[tag] = {"produced": pdf.is_file(),
                       "longest_intact": any(WIDE_LONGEST in ln for ln in lines),
                       "lines": len(lines)}

    # Header repeat, asked of a table long enough to break twice.
    tall = work / "tall-wide.docx"
    tall_table(base, tall, TALL_PRESET_ROWS)
    for name in ("grid", "finance"):
        styled = work / f"tall-{name}.docx"
        run_script("docx_table.py", "--in", tall, "--out", styled, "--preset", name)
        pdf = work / f"tall-{name}.pdf"
        run_script("docx_pdf.py", "--in", styled, "--out", pdf)
        if not pdf.is_file():
            render[name] = {"produced": False, "pages": 0, "header_pages": []}
            continue
        pages = rendered_pages(pdf)
        render[name] = {
            "produced": True, "pages": pages,
            "header_pages": [i + 1 for i in range(pages)
                             if any(WIDE_ROWS[0][0] in ln
                                    for ln in rendered_lines(pdf, i))],
        }
    out["render"] = render
    return out


# ── W4: generating from Markdown ──────────────────────────────────────────────
def numbering_of(path: Path) -> dict:
    """`{numId: [ilvl...]}` for the lists the document defines."""
    if "word/numbering.xml" not in parts_of(path):
        return {}
    root = tree_of(path, "word/numbering.xml")
    abstracts = {a.get(W + "abstractNumId"): [l.get(W + "ilvl")
                                              for l in a.findall(W + "lvl")]
                 for a in root.findall(W + "abstractNum")}
    out = {}
    for num in root.findall(W + "num"):
        ref = num.find(W + "abstractNumId")
        out[num.get(W + "numId")] = abstracts.get(
            ref.get(W + "val") if ref is not None else None, [])
    return out


def list_levels_used(path: Path) -> set:
    out = set()
    for numpr in tree_of(path).iter(W + "numPr"):
        ilvl = numpr.find(W + "ilvl")
        out.add(ilvl.get(W + "val") if ilvl is not None else "0")
    return out


def collect_markdown(work: Path) -> dict:
    generated = work / "generated.docx"
    r = run_script("docx_from_md.py", "--in", SAMPLE_MD, "--out", generated,
                   "--report", work / "generate.json")
    out = {
        "exit": r.returncode,
        "report": json.loads((work / "generate.json").read_text(encoding="utf-8")),
        "paragraphs": len(paragraph_texts(generated)),
        "styles": sorted(styles_of(generated)),
        "paragraph_styles": paragraph_styles(generated),
        "numbering": numbering_of(generated),
        "levels_used": sorted(list_levels_used(generated)),
        "body_children": body_child_names(generated),
        "document_xml": part_text(generated, "word/document.xml"),
        "parts": sorted(parts_of(generated)),
    }
    strict = run_script("docx_from_md.py", "--in", SAMPLE_MD,
                        "--out", work / "strict.docx", "--strict")
    out["strict"] = {"exit": strict.returncode, "stderr": strict.stderr.strip(),
                     "wrote": (work / "strict.docx").exists()}
    # A template's styles, numbering, headers and footers survive; only its body is
    # replaced. report.docx has all four, so it is the one that can prove it.
    tpl = work / "from-template.docx"
    tr = run_script("docx_from_md.py", "--in", SAMPLE_MD, "--out", tpl,
                    "--template", REPORT, "--report", work / "tpl.json")
    out["template"] = {
        "exit": tr.returncode,
        "report": json.loads((work / "tpl.json").read_text(encoding="utf-8")),
        "parts": sorted(parts_of(tpl)),
        "lost": sorted(set(parts_of(REPORT)) - set(parts_of(tpl))),
        "texts": paragraph_texts(tpl),
    }
    return out


# ── W14: schema validation ────────────────────────────────────────────────────
def strip_required_attribute(src: Path, dst: Path) -> None:
    """Reintroduce the exact defect python-docx's own template carries: a
    `<w:zoom>` with no `w:percent`, which Transitional makes REQUIRED."""
    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/settings.xml":
            return data
        return data.replace(b'<w:zoom w:percent="100"/>', b'<w:zoom w:val="bestFit"/>')
    rewrite_zip(src, dst, mutate)


def with_ignorable_namespace(src: Path, dst: Path) -> None:
    """A document carrying an element from a namespace the schema predates, marked
    `mc:Ignorable` — which is what every document Word itself writes looks like."""
    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/document.xml":
            return data
        text = data.decode()
        text = text.replace(
            '<w:document ',
            '<w:document xmlns:mc="http://schemas.openxmlformats.org/'
            'markup-compatibility/2006" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'mc:Ignorable="w14" ', 1)
        return text.replace("<w:body>", '<w:body><w14:conflictIns w14:id="7"/>',
                            1).encode()
    rewrite_zip(src, dst, mutate)


def collect_validate(work: Path) -> dict:
    out: dict = {"schema_files": sorted(p.name for p in SCHEMA_DIR.glob("*.xsd"))}
    r = run_script("docx_validate.py", "--in", REPORT, "--report", work / "ok.json")
    out["good"] = {"exit": r.returncode,
                   "report": json.loads((work / "ok.json").read_text(encoding="utf-8"))}
    broken = work / "schema-broken.docx"
    strip_required_attribute(REPORT, broken)
    b = run_script("docx_validate.py", "--in", broken, "--report", work / "bad.json")
    out["invalid"] = {"exit": b.returncode,
                      "report": json.loads((work / "bad.json").read_text(encoding="utf-8"))}
    mce = work / "ignorable.docx"
    with_ignorable_namespace(REPORT, mce)
    m = run_script("docx_validate.py", "--in", mce, "--report", work / "mce.json")
    out["ignorable"] = {"exit": m.returncode,
                        "report": json.loads((work / "mce.json").read_text(encoding="utf-8"))}
    # No schemas anywhere: the loud degradation. Pointed at an empty directory
    # rather than by moving the shipped one, so a failure here cannot leave the
    # working tree without its schemas.
    empty = work / "no-schemas"
    empty.mkdir(exist_ok=True)
    n = run_script("docx_validate.py", "--in", REPORT, "--schemas", empty)
    out["missing"] = {"exit": n.returncode, "stderr": n.stderr.strip(),
                      "stdout": n.stdout.strip()}
    return out

def blank_document(src: Path, dst: Path) -> None:
    """A document whose body holds one empty paragraph and nothing else.

    The header and footer references go too, and that is not tidiness: the first
    version of this kept them, so LibreOffice rendered a page carrying the
    letterhead and the page number — a page with ink on it. The assertion fired on
    a correct implementation, which is the honest way to find out that the fixture
    was not the thing it was named after.
    """
    import re
    def mutate(name: str, data: bytes) -> bytes:
        if name != "word/document.xml":
            return data
        text = data.decode()
        head = text[:text.index("<w:body>") + len("<w:body>")]
        sect = text[text.index("<w:sectPr>"):]
        sect = re.sub(r"<w:(header|footer)Reference[^>]*/>", "", sect)
        return (head + "<w:p/>" + sect).encode()
    rewrite_zip(src, dst, mutate)


def collect_pdf(work: Path) -> dict:
    """Render, and the blank-render refusal. Skipped and NAMED without LibreOffice."""
    sys.path.insert(0, str(SKILL / "scripts"))
    from office.soffice import find_soffice
    if not find_soffice():
        SKIPS.append("W17 render (Y1-Y3): LibreOffice is not installed on this host, "
                     "so neither the conversion nor its negative controls could run. "
                     "CI runs them on Linux, where libreoffice-writer is installed")
        return {"skipped": "no LibreOffice"}
    out = work / "preview.pdf"
    r = run_script("docx_pdf.py", "--in", REPORT, "--out", out,
                   "--png", work / "pages", "--dpi", "120",
                   "--report", work / "pdf.json")
    report = json.loads((work / "pdf.json").read_text(encoding="utf-8")) \
        if (work / "pdf.json").exists() else {}
    blank = work / "blank.docx"
    blank_document(REPORT, blank)
    b = run_script("docx_pdf.py", "--in", blank, "--out", work / "blank.pdf")
    return {
        "exit": r.returncode, "report": report,
        "produced": out.is_file() and out.stat().st_size > 0,
        "images": len(list((work / "pages").glob("*.png"))) if (work / "pages").is_dir()
                  else 0,
        "blank": {"exit": b.returncode, "stderr": b.stderr.strip(),
                  "wrote": (work / "blank.pdf").exists()},
    }


def revision_elements(path: Path) -> list[dict]:
    """Every w:ins / w:del in document.xml, read straight from the XML."""
    out = []
    for el in tree_of(path).iter():
        name = str(el.tag).rsplit("}", 1)[-1]
        if name not in ("ins", "del"):
            continue
        parent = el.getparent()
        mark = (parent is not None and str(parent.tag).endswith("}rPr")
                and parent.getparent() is not None
                and str(parent.getparent().tag).endswith("}pPr"))
        out.append({
            "kind": name, "id": el.get(W + "id"), "author": el.get(W + "author"),
            "date": el.get(W + "date"), "paragraph_mark": mark,
            "runs": len(el.findall(W + "r")),
            "text": "".join(x.text or "" for x in el.iter(W + "t", W + "delText")),
            "holds_wt": bool(name == "del" and el.findall(f".//{W}t")),
        })
    return out


def body_paragraph_texts(path: Path) -> list[str]:
    root = tree_of(path)
    body = root.find(W + "body")
    return [t for c in body if str(c.tag) == W + "p"
            for t in ["".join(
                (n.text or "") for n in c.iter()
                if str(n.tag) == W + "t"
                and not any(str(a.tag) == W + "del" for a in n.iterancestors()))]]


def collect_revisions(work: Path) -> dict:
    tracked = work / "tracked.docx"
    r = run_script("docx_revise.py", "--in", REPORT, "--out", tracked,
                   "--author", REVISER, "--replace", f"{CROSS_RUN}={CROSS_RUN_NEW}",
                   "--report", work / "revise.json")
    report = json.loads((work / "revise.json").read_text(encoding="utf-8"))

    acc, rej = work / "acc.docx", work / "rej.docx"
    a = run_script("docx_revise.py", "--in", tracked, "--out", acc,
                   "--accept-all", "--strict", "--report", work / "acc.json")
    j = run_script("docx_revise.py", "--in", tracked, "--out", rej,
                   "--reject-all", "--strict", "--report", work / "rej.json")

    added = work / "added.docx"
    run_script("docx_revise.py", "--in", REPORT, "--out", added, "--author", REVISER,
               "--insert-paragraph", INSERTED_PARAGRAPH)
    added_acc, added_rej = work / "added-acc.docx", work / "added-rej.docx"
    run_script("docx_revise.py", "--in", added, "--out", added_acc, "--accept-all")
    run_script("docx_revise.py", "--in", added, "--out", added_rej, "--reject-all")

    # Per-author: accept only the reviser's, leaving the fixture author's alone.
    by_author = work / "by-author.docx"
    ba = run_script("docx_revise.py", "--in", tracked, "--out", by_author,
                    "--accept-author", REVISER, "--report", work / "author.json")
    strict_partial = run_script("docx_revise.py", "--in", tracked,
                                "--out", work / "sp.docx",
                                "--accept-author", REVISER, "--strict")

    # Inserted, then deleted: the nesting the L2 gate used to reject. Both accepting
    # and rejecting must make the text go away, for opposite reasons.
    nested = work / "nested.docx"
    run_script("docx_revise.py", "--in", tracked, "--out", nested,
               "--author", REVISER, "--delete", CROSS_RUN_NEW)
    n_acc, n_rej = work / "nested-acc.docx", work / "nested-rej.docx"
    run_script("docx_revise.py", "--in", nested, "--out", n_acc, "--accept-all")
    run_script("docx_revise.py", "--in", nested, "--out", n_rej, "--reject-all")

    return {
        "exit": r.returncode, "report": report,
        "elements": revision_elements(tracked),
        "accept": {"exit": a.returncode,
                   "report": json.loads((work / "acc.json").read_text(encoding="utf-8")),
                   "texts": body_paragraph_texts(acc),
                   "elements": revision_elements(acc)},
        "reject": {"exit": j.returncode,
                   "report": json.loads((work / "rej.json").read_text(encoding="utf-8")),
                   "texts": body_paragraph_texts(rej),
                   "elements": revision_elements(rej)},
        "original_texts": body_paragraph_texts(REPORT),
        "inserted": {
            "elements": revision_elements(added),
            "accepted": body_paragraph_texts(added_acc),
            "rejected": body_paragraph_texts(added_rej),
        },
        "by_author": {
            "exit": ba.returncode,
            "report": json.loads((work / "author.json").read_text(encoding="utf-8")),
            "elements": revision_elements(by_author),
            "strict_exit": strict_partial.returncode,
            "strict_stderr": strict_partial.stderr.strip(),
            "strict_wrote": (work / "sp.docx").exists(),
        },
        "nested": {
            "elements": revision_elements(nested),
            "accepted": body_paragraph_texts(n_acc),
            "rejected": body_paragraph_texts(n_rej),
        },
    }


def comment_pieces(path: Path) -> dict:
    parts = parts_of(path)
    body = parts.get("word/document.xml", b"")
    return {
        "part": "word/comments.xml" in parts,
        "override": b"comments+xml" in parts.get("[Content_Types].xml", b""),
        "relationship": b"/comments\"" in parts.get("word/_rels/document.xml.rels", b""),
        "range_starts": body.count(b"commentRangeStart"),
        "range_ends": body.count(b"commentRangeEnd"),
        "references": body.count(b"<w:commentReference"),
    }


def bare_document(src: Path, dst: Path) -> None:
    """report.docx with its comments part and all five of its pieces removed."""
    def mutate(name: str, data: bytes) -> bytes | None:
        if name == "word/comments.xml":
            return None
        if name == "[Content_Types].xml":
            return data.replace(
                b'<Override PartName="/word/comments.xml" ContentType="application/'
                b'vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>',
                b"")
        if name == "word/_rels/document.xml.rels":
            return data.replace(
                b'<Relationship Id="rId7" Type="http://schemas.openxmlformats.org/'
                b'officeDocument/2006/relationships/comments" Target="comments.xml"/>',
                b"")
        if name == "word/document.xml":
            for junk in (b'<w:commentRangeStart w:id="1"/>',
                         b'<w:commentRangeEnd w:id="1"/>',
                         b'<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
                         b'<w:commentReference w:id="1"/></w:r>'):
                data = data.replace(junk, b"")
            return data
        return data
    rewrite_zip(src, dst, mutate)


def collect_comments(work: Path) -> dict:
    added = work / "commented.docx"
    a = run_script("docx_comment.py", "--in", REPORT, "--out", added,
                   "--add-on", COMMENT_ANCHOR_SPLIT, "--text", NEW_COMMENT,
                   "--author", "王审计", "--report", work / "comment.json")
    report = json.loads((work / "comment.json").read_text(encoding="utf-8"))

    bare = work / "bare.docx"
    bare_document(REPORT, bare)
    scratch = work / "scratch.docx"
    sc = run_script("docx_comment.py", "--in", bare, "--out", scratch,
                    "--add-on", "应收账款", "--text", "从零新建批注部件",
                    "--report", work / "scratch.json")

    gone = work / "gone.docx"
    dl = run_script("docx_comment.py", "--in", added, "--out", gone,
                    "--delete", "1", "--delete", "2", "--report", work / "gone.json")

    missing = run_script("docx_comment.py", "--in", REPORT, "--out", work / "nope.docx",
                         "--add-on", "文档里没有这句话", "--text", "x")
    listed = run_script("docx_comment.py", "--in", REPORT, "--list")
    return {
        "exit": a.returncode, "report": report,
        "pieces": comment_pieces(added),
        "before_pieces": comment_pieces(REPORT),
        "scratch": {"exit": sc.returncode,
                    "report": json.loads((work / "scratch.json")
                                         .read_text(encoding="utf-8")),
                    "pieces": comment_pieces(scratch)},
        "deleted": {"exit": dl.returncode,
                    "report": json.loads((work / "gone.json")
                                         .read_text(encoding="utf-8")),
                    "pieces": comment_pieces(gone)},
        "missing_anchor": {"exit": missing.returncode,
                           "stderr": missing.stderr.strip(),
                           "wrote": (work / "nope.docx").exists()},
        "list_exit": listed.returncode,
    }


def rpr_of_first_run(path: Path, paragraph_index: int) -> str | None:
    from lxml import etree
    paras = list(tree_of(path).iter(W + "p"))
    if paragraph_index >= len(paras):
        return None
    run = paras[paragraph_index].find(W + "r")
    rpr = run.find(W + "rPr") if run is not None else None
    return etree.tostring(rpr).decode() if rpr is not None else None


def header_text(path: Path) -> str:
    root = tree_of(path, "word/header1.xml")
    return "".join(t.text or "" for t in root.iter(W + "t"))


def appended_run_fonts(path: Path) -> dict:
    """The @ascii / @eastAsia of the run in the last body paragraph."""
    body = tree_of(path).find(W + "body")
    paras = [c for c in body if str(c.tag) == W + "p"]
    if not paras:
        return {}
    rfonts = paras[-1].find(f"{W}r/{W}rPr/{W}rFonts")
    if rfonts is None:
        return {}
    return {"ascii": rfonts.get(W + "ascii"), "eastAsia": rfonts.get(W + "eastAsia")}


def num_ids_of(path: Path) -> list[str]:
    return [n.get(W + "val") for n in tree_of(path).iter(W + "numId")]


def make_xlsx_like(path: Path) -> None:
    """A valid OOXML package that is not a Word document."""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml",
                   '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats'
                   '.org/package/2006/content-types"><Default Extension="xml" '
                   'ContentType="application/xml"/></Types>')
        z.writestr("xl/workbook.xml", '<?xml version="1.0"?><workbook/>')


def traversal_refused(work: Path) -> dict:
    """A part named to escape the unpack directory must be refused, not written."""
    evil = work / "evil.docx"
    def mutate(name: str, data: bytes) -> bytes:
        return data
    rewrite_zip(REPORT, evil, mutate)
    with zipfile.ZipFile(evil, "a", zipfile.ZIP_DEFLATED) as z:
        z.writestr("../escaped.xml", "<x/>")
    r = run_script("docx_package.py", "--in", evil, "--unpack", work / "evil-unpacked")
    return {"exit": r.returncode, "stderr": r.stderr.strip(),
            "escaped_written": (work / "escaped.xml").exists()}


# ── the assertions ────────────────────────────────────────────────────────────
CHECKS: dict[str, dict] = {}


def check(cid: str, title: str):
    def deco(fn):
        CHECKS[cid] = {"title": title, "fn": fn}
        return fn
    return deco


@check("V0", "the fixtures actually give every assertion something to look at")
def v0_not_vacuous(ctx: dict) -> list[str]:
    f = ctx["fixture"]
    out = []
    # If the phrase did NOT span runs, E1's whole point evaporates and the naive
    # control would agree with the real implementation.
    if f["phrase_runs"] < 2:
        out.append(f"V0 {CROSS_RUN!r} is stored in {f['phrase_runs']} run(s) in "
                   f"report.docx — the fixture no longer exercises the case this "
                   f"whole capability exists for")
    if f["naive_hits"] != NAIVE_HITS:
        out.append(f"V0 the per-run search found {f['naive_hits']} hit(s) for "
                   f"{CROSS_RUN!r}, expected {NAIVE_HITS} — the fixture no longer "
                   f"separates the two implementations by exactly the spanning "
                   f"occurrence, which is the whole comparison")
    for name in ("insertions", "deletions", "tables", "comments", "list_paragraphs",
                 "footer_fields"):
        if f[name] < 1:
            out.append(f"V0 report.docx carries no {name}; the assertions that read "
                       f"them have nothing to look at")
    if not f["custom_xml"]:
        out.append("V0 report.docx has no customXml part, so the fidelity claim has "
                   "nothing to lose")
    if f["header_footer_parts"] < 2:
        out.append("V0 report.docx has no header and footer pair")
    if f["parts"] != PARTS_TOTAL:
        out.append(f"V0 report.docx has {f['parts']} parts, the assertions are "
                   f"written for {PARTS_TOTAL}")
    if f["order_defects"] != ORDER_DEFECTS:
        out.append(f"V0 unordered.docx carries {f['order_defects']} order defect(s), "
                   f"expected {ORDER_DEFECTS} — the repair assertions are measuring a "
                   f"different fixture than they were written for")

    # outline.docx earns its place by what it does NOT have. The moment it gains a
    # header, a png Default or a media part, every "create" assertion silently
    # becomes a "replace" assertion and stops testing the package wiring.
    if f["outline_header_footer_parts"] or f["outline_sect_refs"]:
        out.append(f"V0 outline.docx now carries {f['outline_header_footer_parts']} "
                   f"header/footer part(s) and {f['outline_sect_refs']} binding(s) — "
                   f"W9 would then be exercising replacement, not creation")
    if f["outline_png_default"] or f["outline_media_parts"]:
        out.append("V0 outline.docx already declares a picture content type or holds "
                   "media, so inserting one no longer has to declare anything")
    if f["outline_paragraphs"] != OUTLINE_PARAGRAPHS:
        out.append(f"V0 outline.docx has {f['outline_paragraphs']} paragraphs, the "
                   f"assertions are written for {OUTLINE_PARAGRAPHS}")
    if sorted(set(f["outline_heading_levels"])) != [1, 2, 3] or \
            f["outline_heading_styles"] < 3:
        out.append(f"V0 outline.docx's heading levels are "
                   f"{sorted(set(f['outline_heading_levels']))}; a contents page "
                   f"built from one level proves nothing about levels")
    if not f["chart_has_phys"] or f["chart_square"]:
        out.append("V0 chart.png is square or no longer states its own density — a "
                   "square picture hides a swapped extent, and without pHYs the "
                   "96-dpi assumption gives the same answer as reading the file")

    # fontless.docx must carry all three shapes, and the STYLE must be the only
    # place one of the answers lives — otherwise "keep what the document said" and
    # "write the default everywhere" produce the same output.
    if f["fontless_unbound"] != FONTLESS_UNBOUND_RUNS:
        out.append(f"V0 fontless.docx has {f['fontless_unbound']} unbound CJK run(s), "
                   f"expected {FONTLESS_UNBOUND_RUNS}")
    if not f["fontless_style_says"]:
        out.append(f"V0 no style in fontless.docx names {FONTLESS_STYLE_EA!r}, so a "
                   f"repair that writes the fallback everywhere cannot be told apart "
                   f"from one that honours the document")
    if f["fontless_docdefaults_say"]:
        out.append("V0 fontless.docx's w:docDefaults state a font again, so no run in "
                   "it is genuinely unbound and --strict has nothing to refuse")
    if not f["fontless_kept_ascii"]:
        out.append(f"V0 fontless.docx no longer carries a deliberate "
                   f"{FONTLESS_KEPT_ASCII!r}, so 'the repair kept it' is untestable")
    if f["report_unbound"]:
        out.append(f"V0 report.docx now has {f['report_unbound']} unbound CJK run(s), "
                   f"so the clean --check case is no longer clean")

    # W19's probe strings have to occur NOWHERE ELSE in report.docx. Q5 answers "did
    # this cell's text survive on one line" by looking for it among the rendered
    # lines, and a string that is also a substring of the document's title is found
    # on the title's line whether the cell wrapped or not. That is not a hypothetical:
    # the first draft used one, and Q5 passed for both implementations.
    # revised.docx has to carry all six changes W18's assertions are written for. It
    # is generated by substituting into report.docx's markup, so an edit to
    # report.docx that stops one substitution matching would otherwise leave the
    # assertions hunting a difference nobody put there.
    d = ctx.get("diff", {}).get("fixture") or {}
    if d:
        if not d["removed_from_a"]:
            out.append(f"V0 revised.docx no longer drops the paragraph "
                       f"{DIFF_REMOVED!r}, so U1's paragraph-removed finding has no "
                       f"subject")
        if not d["added_in_b"]:
            out.append(f"V0 revised.docx no longer adds {DIFF_ADDED!r}")
        if len(d["restyled"]) != 1:
            out.append(f"V0 revised.docx restyles {d['restyled']}, expected exactly "
                       f"one paragraph; the style-only change is what separates "
                       f"'a paragraph was rewritten' from 'a paragraph was "
                       f"reformatted'")
        if d["header_a"] == d["header_b"]:
            out.append("V0 revised.docx's header is identical to report.docx's, so "
                       "the assertion that a diff reaches the letterhead has no "
                       "subject")
        if not d["same_parts"]:
            out.append("V0 report.docx and revised.docx no longer hold the same "
                       "parts, so a part-level difference would be mixed in with the "
                       "six text ones")

    for text, hits in (ctx.get("tables", {}).get("probe_strings_unique") or {}).items():
        if hits != 1:
            out.append(f"V0 the table probe string {text!r} occurs in {hits} "
                       f"paragraph(s) of the document Q5 renders, not exactly 1 — a "
                       f"render that finds it then proves nothing about the cell")
    return out


# ── W1: reading ───────────────────────────────────────────────────────────────
@check("R1", "a phrase Word split across runs reads back as one string")
def r1_joined(ctx: dict) -> list[str]:
    paras = ctx["read"]["report"]["paragraphs"]
    match = [p for p in paras if p["text"] == SPLIT_PARAGRAPH]
    if not match:
        return [f"R1 no paragraph reads as {SPLIT_PARAGRAPH!r}; got "
                f"{[p['text'][:24] for p in paras[:4]]}"]
    # And it really was stored in several runs, or the join proved nothing.
    if match[0]["runs"] < 2:
        return [f"R1 that paragraph is stored in {match[0]['runs']} run(s), so joining "
                f"them was not the thing that made it readable"]
    return []


@check("R2", "tracked-deleted text is not reported as document text")
def r2_deleted(ctx: dict) -> list[str]:
    report = ctx["read"]["report"]
    out = []
    if any(DELETED in p["text"] for p in report["paragraphs"]):
        out.append(f"R2 {DELETED!r} was deleted with change tracking on and is being "
                   f"reported as document text — the document now says the opposite "
                   f"of what it says")
    deletions = report.get("revisions", {}).get("deletions", [])
    if not any(d["text"] == DELETED for d in deletions):
        out.append(f"R2 {DELETED!r} is not reported under revisions either, so it is "
                   f"simply invisible")
    if deletions and not deletions[0].get("author"):
        out.append("R2 a deletion is reported with no author — who made a change is "
                   "most of what a revision is")
    return out


@check("R3", "tracked-inserted text IS reported as document text")
def r3_inserted(ctx: dict) -> list[str]:
    report = ctx["read"]["report"]
    out = []
    if not any(INSERTED in p["text"] for p in report["paragraphs"]):
        out.append(f"R3 {INSERTED!r} is a tracked INSERTION — it is in the document, "
                   f"and it is missing from the reported text")
    # The measured contrast: python-docx itself gets this wrong, so agreeing with it
    # would mean the skill inherited the bug rather than avoided it.
    if any(INSERTED in t for t in ctx["read"]["python_docx_paragraph_texts"]):
        out.append("R3 python-docx now reports inserted text too, so this assertion "
                   "no longer distinguishes the two readers")
    return out


@check("R4", "tables come back as a grid, not as flattened text")
def r4_tables(ctx: dict) -> list[str]:
    tables = ctx["read"]["report"].get("table_contents") or []
    if not tables:
        return ["R4 no table contents were reported at all"]
    t = tables[0]
    out = []
    if t["rows"] < 3 or t["columns"] < 3:
        out.append(f"R4 the table is reported as {t['rows']}x{t['columns']}; the "
                   f"fixture's is 3x3")
    flat = [c for row in t["cells"] for c in row]
    if TABLE_CELL not in flat:
        out.append(f"R4 {TABLE_CELL!r} is not in any reported cell")
    if any(len(row) != t["columns"] for row in t["cells"]):
        out.append("R4 the reported rows are ragged — a grid whose rows have "
                   "different lengths cannot be indexed by column")
    return out


@check("R5", "list level, style and section structure are reported")
def r5_structure(ctx: dict) -> list[str]:
    report = ctx["read"]["report"]
    out = []
    listed = [p for p in report["paragraphs"] if p.get("list")]
    if not listed:
        out.append("R5 no paragraph is reported as belonging to a list, though the "
                   "fixture has two")
    elif not any(p["text"] == LIST_TEXT for p in listed):
        out.append(f"R5 the list paragraphs do not include {LIST_TEXT!r}")
    if "Heading1" not in report["styles_used"]:
        out.append("R5 the heading style is not reported as used")
    sections = report.get("sections") or []
    if not sections or not sections[0].get("headers"):
        out.append("R5 the section is reported without its header binding, so a "
                   "reader cannot tell the document has one")
    return out


@check("R6", "headers and footers are reported separately, fields named as fields")
def r6_headers(ctx: dict) -> list[str]:
    hf = ctx["read"]["report"].get("headers_and_footers") or []
    out = []
    if len(hf) != 2:
        out.append(f"R6 {len(hf)} header/footer part(s) reported, expected 2")
    header = next((h for h in hf if h.get("kind") == "header"), None)
    footer = next((h for h in hf if h.get("kind") == "footer"), None)
    if not header or HEADER_TEXT not in header.get("text", ""):
        out.append(f"R6 the header's text does not contain {HEADER_TEXT!r}")
    if not footer:
        out.append("R6 no footer was reported")
    elif footer.get("fields", 0) < 1:
        out.append("R6 the footer holds a PAGE field and no field is reported — the "
                   "cached '1' then reads as the footer's literal text, which is how "
                   "a reader concludes every page is page 1")
    return out


# ── W2: replacing ─────────────────────────────────────────────────────────────
@check("E1", "a phrase stored across runs is found, where the per-run search finds none")
def e1_cross_run(ctx: dict) -> list[str]:
    entry = ctx["replace"]["report"]["replacements"][0]
    out = []
    if entry["replaced"] != CROSS_RUN_TOTAL:
        out.append(f"E1 {entry['replaced']} occurrence(s) of {CROSS_RUN!r} replaced, "
                   f"expected {CROSS_RUN_TOTAL} (one inside a single run, one split "
                   f"across two)")
    if entry["cross_run"] != CROSS_RUN_SPANNING:
        out.append(f"E1 {entry['cross_run']} replacement(s) crossed a run boundary, "
                   f"expected {CROSS_RUN_SPANNING}")
    if ctx["fixture"]["naive_hits"] >= entry["replaced"]:
        out.append(f"E1 the per-run implementation found "
                   f"{ctx['fixture']['naive_hits']}, this one found "
                   f"{entry['replaced']} — they are not distinguishable, so this "
                   f"assertion is not measuring anything")
    return out


@check("E2", "the replacement is in the document text and the old phrase is gone")
def e2_text(ctx: dict) -> list[str]:
    texts = ctx["replace"]["texts"]
    out = []
    if not any(CROSS_RUN_NEW in t for t in texts):
        out.append(f"E2 {CROSS_RUN_NEW!r} is not in the document text")
    if any(CROSS_RUN in t for t in texts):
        out.append(f"E2 {CROSS_RUN!r} is still there — the report said it was "
                   f"replaced and the file disagrees")
    if not any(PLACEHOLDER_NEW in t for t in texts):
        out.append(f"E2 the second --replace never landed: no {PLACEHOLDER_NEW!r}")
    if any(PLACEHOLDER in t for t in texts):
        out.append(f"E2 {PLACEHOLDER!r} survived — a template that still shows its "
                   f"own placeholder is the failure this is for")
    return out


@check("E3", "only the parts that changed are rewritten")
def e3_surgical(ctx: dict) -> list[str]:
    r = ctx["replace"]
    out = []
    if r["report"]["parts_changed"] != ["word/document.xml"]:
        out.append(f"E3 the edit rewrote {r['report']['parts_changed']}; a text "
                   f"replacement has business in exactly one part")
    if len(r["parts_after"]) != PARTS_TOTAL:
        out.append(f"E3 the output has {len(r['parts_after'])} parts, the input had "
                   f"{PARTS_TOTAL} — something was dropped or invented")
    if len(r["identical"]) != PARTS_TOTAL - 1:
        lost = sorted(set(r["parts_before"]) - set(r["identical"]))
        out.append(f"E3 {len(r['identical'])}/{PARTS_TOTAL} parts came through "
                   f"byte-identical; these did not: {lost}")
    if r["rpr_before"] != r["rpr_after"]:
        out.append("E3 the run's properties changed — a text replacement that resets "
                   "the font binding is how Chinese turns into tofu two steps later")
    return out


@check("E4", "a phrase broken by a line break is not silently matched, and is named")
def e4_near_miss(ctx: dict) -> list[str]:
    rep = ctx["near_miss"]["report"]
    entry = rep["replacements"][0]
    out = []
    if entry["replaced"] != 0:
        out.append(f"E4 {entry['replaced']} replacement(s) across a line break — the "
                   f"break was treated as if it were not there, which silently "
                   f"reflows the paragraph")
    near = entry.get("near_misses") or []
    if not near:
        out.append("E4 nothing matched and the report does not say why; '0 "
                   "replacements' with no reason is the answer that wastes an "
                   "afternoon")
    elif not any("三季度营业收入" in str(n) for n in near):
        out.append(f"E4 the near miss does not name the phrase it nearly matched: "
                   f"{near}")
    return out


@check("E5", "an edit that lands inside tracked content says so")
def e5_revision_context(ctx: dict) -> list[str]:
    entry = ctx["revision_edit"]["report"]["replacements"][0]
    out = []
    if entry["replaced"] != 1:
        out.append(f"E5 {entry['replaced']} replacement(s) inside the tracked "
                   f"insertion, expected 1")
    if "ins" not in (entry.get("contexts") or []):
        out.append("E5 the replacement happened inside a <w:ins> and the report does "
                   "not mention it — rewriting somebody's tracked change without "
                   "saying so is the one thing a reviewer must not discover later")
    return out


@check("E6", "--in-headers reaches header parts, and only when asked")
def e6_headers(ctx: dict) -> list[str]:
    out = []
    if "公开资料" not in ctx["headers"]["header_text"]:
        out.append("E6 --in-headers did not reach word/header1.xml")
    changed = ctx["headers"]["report"]["parts_changed"]
    if "word/header1.xml" not in changed:
        out.append(f"E6 the report does not list the header as changed: {changed}")
    # Control: without the flag the header's TEXT must be untouched, or --in-headers
    # is decoration. Measured on the header rather than on the part list, because
    # "which parts were rewritten" is E3's question and sharing it would make one
    # defect light two checks.
    if ctx["replace"]["header_text"] != HEADER_TEXT:
        out.append(f"E6 a replacement WITHOUT --in-headers changed the header text to "
                   f"{ctx['replace']['header_text']!r}; the flag decides nothing")
    return out


# ── W3: appending ─────────────────────────────────────────────────────────────
@check("A1", "an appended paragraph lands before <w:sectPr>, which stays last")
def a1_sectpr(ctx: dict) -> list[str]:
    kids = ctx["append"]["body_children"]
    out = []
    if not kids or kids[-1] != "sectPr":
        out.append(f"A1 the body ends with {kids[-1:] or ['nothing']}, not <w:sectPr> "
                   f"— Word repairs such a file by discarding the section, and with "
                   f"it the page size, margins and header bindings")
    if APPEND_TEXT not in ctx["append"]["texts"]:
        out.append(f"A1 the appended paragraph is not in the document text")
    if not ctx["append"]["section_survives"]:
        out.append("A1 the section properties are gone from the body entirely")
    return out


@check("A2", "--list is refused when there is no list to join, and works when there is")
def a2_list(ctx: dict) -> list[str]:
    without = ctx["append"]["list_without_numbering"]
    with_ = ctx["append"]["list_with_numbering"]
    out = []
    if without["exit"] != 2:
        out.append(f"A2 --list on a document with no numbering.xml exited "
                   f"{without['exit']}; a numPr pointing at a definition that does "
                   f"not exist is a paragraph Word renders without its number")
    if without["wrote"]:
        out.append("A2 it refused and wrote the file anyway")
    if "numbering" not in without["stderr"]:
        out.append(f"A2 the refusal does not say what is missing: {without['stderr']!r}")
    if with_["exit"] != 0:
        out.append(f"A2 --list on a document that HAS numbering.xml exited "
                   f"{with_['exit']} — the guard is refusing everything")
    elif with_["num_ids"].count("1") < 3:
        out.append(f"A2 the appended paragraph did not join list 1: numIds are "
                   f"{with_['num_ids']}")
    return out


@check("A3", "a Chinese paragraph this skill writes binds both font faces")
def a3_fonts(ctx: dict) -> list[str]:
    fonts = ctx["append"]["fonts"]
    out = []
    if not fonts.get("eastAsia"):
        out.append("A3 the appended run has no @w:eastAsia; Word then picks the CJK "
                   "face from the theme and the same file renders differently on "
                   "another machine")
    if not fonts.get("ascii"):
        out.append("A3 the appended run has no @w:ascii, so mixed '2026 年' text "
                   "renders in two unrelated typefaces")
    return out


# ── W13: unpack / pack ────────────────────────────────────────────────────────
@check("P1", "the round trip gives back every part, byte for byte")
def p1_roundtrip(ctx: dict) -> list[str]:
    p = ctx["package"]
    out = []
    if p["identical"] != PARTS_TOTAL:
        out.append(f"P1 {p['identical']}/{PARTS_TOTAL} parts came back identical")
    if p["report"].get("parts_lost"):
        out.append(f"P1 parts lost in the round trip: {p['report']['parts_lost']}")
    if not p["report"].get("order_preserved"):
        out.append("P1 the report itself says the part order changed")
    return out


@check("P2", "the part order survives, and the manifest is what makes it survive")
def p2_order(ctx: dict) -> list[str]:
    p = ctx["package"]
    out = []
    if p["order_after"] != p["order_before"]:
        out.append(f"P2 part order changed: {p['order_after'][:3]} vs "
                   f"{p['order_before'][:3]}")
    # Control: without the manifest the order is the filesystem's, not the
    # document's. If the two agree, the manifest is not the thing preserving it.
    without = p["without_manifest_order"]
    if without and without == p["order_before"]:
        out.append("P2 packing WITHOUT the manifest produced the same order, so this "
                   "assertion cannot tell whether the manifest does anything")
    return out


@check("P3", "a part added between unpack and pack is kept")
def p3_added(ctx: dict) -> list[str]:
    added = ctx["package"]["added_part"]
    if added["exit"] != 0:
        return [f"P3 packing a directory with a hand-added part exited {added['exit']}"]
    if not added["present"]:
        return ["P3 the added part is not in the packed file — editing a part is the "
                "entire reason to unpack one"]
    return []


@check("P4", "a part name that escapes the target directory is refused")
def p4_traversal(ctx: dict) -> list[str]:
    t = ctx["package"]["traversal"]
    out = []
    if t["escaped_written"]:
        out.append("P4 a part named ../escaped.xml was written outside the unpack "
                   "directory")
    if t["exit"] == 0:
        out.append("P4 the unpack reported success on a package carrying a traversal "
                   "part name")
    elif "escapes" not in t["stderr"]:
        out.append(f"P4 the refusal does not say why: {t['stderr']!r}")
    return out


# ── W15: element order ────────────────────────────────────────────────────────
@check("O1", "the unordered input really is rejected, by name, all three times")
def o1_input_rejected(ctx: dict) -> list[str]:
    o = ctx["order"]
    out = []
    if o["check_exit"] != 2:
        out.append(f"O1 --check exited {o['check_exit']} on a document with "
                   f"{ORDER_DEFECTS} order defects")
    findings = " ".join(o["check_findings"])
    for element in ("pPr", "rPr", "sectPr"):
        if element not in findings:
            out.append(f"O1 the check does not name <w:{element}> — a repair whose "
                       f"input was never shown to be broken proves nothing")
    return out


@check("O2", "--fix-order repairs every one of them and the output passes")
def o2_repaired(ctx: dict) -> list[str]:
    o = ctx["order"]
    out = []
    if o["fix_exit"] != 0:
        out.append(f"O2 --fix-order exited {o['fix_exit']}")
    fixes = o["fix_report"].get("fixes") or []
    moved = sum(f["elements_reordered"] for f in fixes)
    if moved != ORDER_DEFECTS:
        out.append(f"O2 {moved} element(s) reordered, expected {ORDER_DEFECTS}")
    if o["after_findings"]:
        out.append(f"O2 the repaired document still has: {o['after_findings']}")
    if o["after_exit"] != 0:
        out.append(f"O2 --check on the repaired document exited {o['after_exit']}")
    return out


@check("O3", "the repair moves elements and changes nothing else")
def o3_nothing_else(ctx: dict) -> list[str]:
    o = ctx["order"]
    out = []
    if o["text_after"] != o["text_before"]:
        out.append("O3 the document's text changed while its element order was being "
                   "repaired")
    if o["other_parts_identical"] != o["other_parts_total"]:
        out.append(f"O3 {o['other_parts_identical']}/{o['other_parts_total']} of the "
                   f"other parts are byte-identical; a reordering has business in "
                   f"word/document.xml alone")
    return out


# ── contracts ─────────────────────────────────────────────────────────────────
@check("C1", "stdout stays a summary on a long document")
def c1_stdout(ctx: dict) -> list[str]:
    c = ctx["contracts"]
    out = []
    if c["stdout_exit"] != 0:
        out.append(f"C1 reading a {SCALE_PARAGRAPHS}-paragraph document exited "
                   f"{c['stdout_exit']}")
    if c["stdout_bytes"] > STDOUT_BUDGET:
        out.append(f"C1 stdout was {c['stdout_bytes']} bytes for a "
                   f"{SCALE_PARAGRAPHS}-paragraph document (budget {STDOUT_BUDGET}); "
                   f"an agent pays for every one of them, twice under delegation")
    if c["stdout_bytes"] < 200:
        out.append("C1 stdout is nearly empty — trimming that drops the answer is not "
                   "trimming")
    return out


@check("C3", "a report with FEW but enormous entries is trimmed too")
def c3_byte_budget(ctx: dict) -> list[str]:
    c = ctx["contracts"]
    out = []
    if c["tall_exit"] != 0:
        out.append(f"C3 reading a {SCALE_TABLE_ROWS}-row table exited {c['tall_exit']}")
    if c["tall_stdout_bytes"] > STDOUT_BUDGET:
        out.append(f"C3 one table printed {c['tall_stdout_bytes']} bytes to stdout "
                   f"(budget {STDOUT_BUDGET}). The item-count trimmer cannot see this "
                   f"shape: the list has ONE entry")
    # The over-correction is its own defect: a trimmer that also removes the answer
    # is not a trimmer. The report must still say a table is there and how big.
    if "table_contents_count" not in c["tall_stdout"]:
        out.append("C3 the trimmed report does not say how many tables it dropped, "
                   "so the reader cannot tell a huge table from no table")
    if c["tall_stdout_bytes"] < 300:
        out.append(f"C3 stdout is {c['tall_stdout_bytes']} bytes — the trimming took "
                   f"the answer with it")
    return out


@check("C4", "every entry point survives a Windows code page instead of crashing on "
             "its own Chinese")
def c4_windows_encoding(ctx: dict) -> list[str]:
    e = ctx["encoding"]
    out = []
    # Vacuity first. If a bare print of Chinese under cp1252 exits 0, this host is
    # not reproducing the condition and nothing below means anything — which is the
    # state every non-Windows machine would be in if PYTHONIOENCODING did not bite.
    if e["bare_print_exit"] == 0:
        return ["C4 a bare print of Chinese under cp1252 exited 0, so this host is "
                "not reproducing the Windows code page at all and this assertion "
                "proves nothing"]
    missed = [n for n in e["entry_points"] if n not in e["probed"]]
    if missed:
        out.append(f"C4 {missed} have no probe, so they are not covered — a new entry "
                   f"point inherits this defect for free")
    for key, r in sorted(e["runs"].items()):
        if r["exit"] != 0:
            out.append(f"C4 {key} exited {r['exit']} — on Windows an agent capturing "
                       f"this output gets UnicodeEncodeError instead of a report: "
                       f"{r['stderr'][:120]!r}")
    speaking = sum(1 for k, r in e["runs"].items()
                   if k.startswith("cp1252:") and r["cjk_in_stdout"])
    if speaking < 5:
        out.append(f"C4 only {speaking} probe(s) printed any Chinese, so most of them "
                   f"never crossed the code path that used to crash")
    return out


@check("C2", "bad inputs get one sentence, not a traceback")
def c2_contracts(ctx: dict) -> list[str]:
    c = ctx["contracts"]
    out = []
    if c["in_place"]["exit"] != 2 or "same file" not in c["in_place"]["stderr"]:
        out.append(f"C2 writing --out over --in was not refused clearly: "
                   f"{c['in_place']}")
    if c["missing"]["exit"] != 2 or c["missing"]["traceback"]:
        out.append(f"C2 a missing file produced exit {c['missing']['exit']}"
                   + (" with a traceback" if c["missing"]["traceback"] else ""))
    if c["wrong_kind"]["exit"] != 2 or c["wrong_kind"]["traceback"]:
        out.append(f"C2 an OOXML package that is not a Word document produced exit "
                   f"{c['wrong_kind']['exit']}"
                   + (" with a traceback" if c["wrong_kind"]["traceback"] else ""))
    elif "not a Word document" not in c["wrong_kind"]["stderr"]:
        out.append(f"C2 the refusal does not say what is wrong with the file: "
                   f"{c['wrong_kind']['stderr']!r}")
    return out


@check("F0", "the committed fixtures are byte-reproducible")
def f0_reproducible(ctx: dict) -> list[str]:
    f = ctx["fixtures"]
    if f["exit"] != 0:
        return [f"F0 make_fixtures.py exited {f['exit']}"]
    bad = [n for n, ok in f["identical"].items() if not ok]
    if bad:
        return [f"F0 regenerating produced different bytes for {bad} — every file "
                f"under skills/builtin/ feeds .builtin-version, so a fixture that is "
                f"not reproducible reinstalls the built-in skills on every desktop "
                f"whenever anyone reruns the generator"]
    return []



# ── W5: template fill ─────────────────────────────────────────────────────────
@check("T1", "a placeholder Word split across runs is filled")
def t1_cross_run(ctx: dict) -> list[str]:
    filled = {e["name"]: e for e in ctx["template"]["report"]["filled"]}
    out = []
    e = filled.get(FILL_SPLIT)
    if not e or e["replaced"] != 1:
        out.append(f"T1 {FILL_SPLIT!r} was filled {e['replaced'] if e else 0} time(s), "
                   f"expected 1")
    elif e["cross_run"] != 1:
        out.append(f"T1 the fill did not cross a run boundary, so the template's "
                   f"split placeholder is not what was matched")
    if any("{{" + FILL_SPLIT + "}}" in t for t in ctx["template"]["texts"]):
        out.append(f"T1 the document still says {{{{{FILL_SPLIT}}}}}")
    return out


@check("T2", "a placeholder left unfilled is named, not shipped in silence")
def t2_unfilled(ctx: dict) -> list[str]:
    report = ctx["template"]["report"]
    names = [p["name"] for p in report["unfilled"]]
    out = []
    if FILL_LEFT not in names:
        out.append(f"T2 {FILL_LEFT!r} is still in the document and is not in "
                   f"`unfilled` ({names}) — a contract going out with a hole in it "
                   f"and nothing saying so")
    if not report.get("warning"):
        out.append("T2 placeholders remain and the report carries no warning")
    return out


@check("T3", "a supplied value that matched nothing is named")
def t3_unused(ctx: dict) -> list[str]:
    unused = ctx["template"]["report"]["unused_values"]
    if FILL_TYPO not in unused:
        return [f"T3 {FILL_TYPO!r} matched no placeholder and is not reported as "
                f"unused ({unused}); a typo in the key is otherwise indistinguishable "
                f"from 'that placeholder is not in this template'"]
    return []


@check("T4", "--strict refuses and writes nothing")
def t4_strict(ctx: dict) -> list[str]:
    st = ctx["template"]["strict"]
    out = []
    if st["exit"] != 2:
        out.append(f"T4 --strict with two placeholders left over exited {st['exit']}")
    if st["wrote"]:
        out.append("T4 --strict refused and wrote the file anyway")
    if "placeholder" not in st["stderr"]:
        out.append(f"T4 the refusal does not say what is unfilled: {st['stderr']!r}")
    return out


@check("T5", "placeholders in the header are filled too")
def t5_header(ctx: dict) -> list[str]:
    filled = {e["name"]: e for e in ctx["template"]["report"]["filled"]}
    out = []
    e = filled.get(FILL_HEADER)
    if not e or not any("header" in p for p in e["parts"]):
        out.append(f"T5 {FILL_HEADER!r} lives in word/header1.xml and was not filled "
                   f"there — half of a real template's placeholders are in the "
                   f"letterhead")
    if "{{" in ctx["template"]["header_text"]:
        out.append(f"T5 the header still reads {ctx['template']['header_text']!r}")
    return out


# ── W17: render ───────────────────────────────────────────────────────────────
@check("Y1", "the document renders to a PDF with ink on it")
def y1_render(ctx: dict) -> list[str]:
    p = ctx["pdf"]
    if p.get("skipped"):
        return []
    out = []
    if p["exit"] != 0 or not p["produced"]:
        out.append(f"Y1 the conversion exited {p['exit']} / produced="
                   f"{p['produced']}")
    r = p["report"]
    if r.get("pages", 0) < 1:
        out.append(f"Y1 the PDF has {r.get('pages')} page(s)")
    if r.get("blank_pages"):
        out.append(f"Y1 page(s) {r['blank_pages']} carry no ink")
    if p["images"] < 1:
        out.append("Y1 --png was asked for and no image was written; a preview that "
                   "silently produces nothing is the failure this refuses")
    return out


@check("Y2", "a document that renders blank is refused, not handed back as a preview")
def y2_blank(ctx: dict) -> list[str]:
    p = ctx["pdf"]
    if p.get("skipped"):
        return []
    b = p["blank"]
    out = []
    if b["exit"] != 2:
        out.append(f"Y2 rendering an empty document exited {b['exit']}; a blank "
                   f"preview looks exactly like lost content")
    if b["wrote"]:
        out.append("Y2 it refused and left the PDF on disk anyway")
    if "no ink" not in b["stderr"]:
        out.append(f"Y2 the refusal does not say why: {b['stderr']!r}")
    return out


@check("Y3", "pending revisions and cached fields are named, not rendered in silence")
def y3_warnings(ctx: dict) -> list[str]:
    p = ctx["pdf"]
    if p.get("skipped"):
        return []
    r = p["report"]
    out = []
    if not r.get("warning"):
        out.append("Y3 the document has tracked changes and the report says nothing — "
                   "the PDF shows one resolution of them, which is not the document "
                   "anyone has approved")
    if not r.get("fields_note"):
        out.append("Y3 the footer's PAGE field renders from its CACHED result and the "
                   "report does not say so")
    return out



# ── W6: making tracked changes ────────────────────────────────────────────────
@check("K1", "a cross-run phrase is tracked as ONE deletion plus one insertion")
def k1_tracked_replace(ctx: dict) -> list[str]:
    r = ctx["revise"]
    out = []
    if r["exit"] != 0:
        out.append(f"K1 the tracked replace exited {r['exit']}")
    dels = [e for e in r["elements"] if e["kind"] == "del" and not e["paragraph_mark"]]
    ins = [e for e in r["elements"] if e["kind"] == "ins" and not e["paragraph_mark"]]
    mine_d = [e for e in dels if e["author"] == REVISER]
    mine_i = [e for e in ins if e["author"] == REVISER]
    if len(mine_d) != CROSS_RUN_TOTAL or len(mine_i) != CROSS_RUN_TOTAL:
        out.append(f"K1 {len(mine_d)} deletion(s) and {len(mine_i)} insertion(s) by "
                   f"{REVISER}, expected {CROSS_RUN_TOTAL} of each")
    split = [e for e in mine_d if e["runs"] > 1]
    if not split:
        out.append("K1 no deletion wraps more than one run — the phrase that spans "
                   "two runs was not split out and wrapped as a unit, which is the "
                   "case this capability exists for")
    if any(e["text"] != CROSS_RUN for e in mine_d):
        out.append(f"K1 a deletion holds {[e['text'] for e in mine_d]}, expected "
                   f"{CROSS_RUN!r} — the wrapping took the wrong characters")
    if any(e["holds_wt"] for e in mine_d):
        out.append("K1 a <w:del> holds <w:t>; deleted text must be <w:delText> or "
                   "Word treats the file as damaged")
    ids = [e["id"] for e in r["elements"]]
    if len(ids) != len(set(ids)):
        out.append(f"K1 revision ids are not unique: {ids}")
    if any(not e["author"] or not e["date"] for e in r["elements"]):
        out.append("K1 a revision carries no author or no date — who changed what is "
                   "most of what a revision is")
    return out


# ── W7: resolving them ────────────────────────────────────────────────────────
@check("K2", "accept applies the change, reject restores the original text")
def k2_round_trip(ctx: dict) -> list[str]:
    r = ctx["revise"]
    out = []
    i = TRACKED_PARAGRAPH
    if CROSS_RUN_NEW not in r["accept"]["texts"][i]:
        out.append(f"K2 accepting left {r['accept']['texts'][i]!r}, which does not "
                   f"carry {CROSS_RUN_NEW!r}")
    if r["reject"]["texts"][i] != r["original_texts"][i]:
        out.append(f"K2 rejecting left {r['reject']['texts'][i]!r}, not the original "
                   f"{r['original_texts'][i]!r} — a rejected change must leave the "
                   f"document exactly as it was")
    # The fixture's OWN revision resolves too, in opposite directions. Blanket
    # equality with the original would be the wrong assertion here and it is worth
    # saying why: reject-all resolves the pre-existing revision as well, correctly.
    if r["accept"]["texts"][REVISION_PARAGRAPH] != ACCEPTED_REVISION:
        out.append(f"K2 accepting the fixture's own revision gave "
                   f"{r['accept']['texts'][REVISION_PARAGRAPH]!r}, expected "
                   f"{ACCEPTED_REVISION!r}")
    if r["reject"]["texts"][REVISION_PARAGRAPH] != REJECTED_REVISION:
        out.append(f"K2 rejecting the fixture's own revision gave "
                   f"{r['reject']['texts'][REVISION_PARAGRAPH]!r}, expected "
                   f"{REJECTED_REVISION!r} — the deleted text must come back")
    return out


@check("K3", "an inserted paragraph carries BOTH marks, so rejecting removes it whole")
def k3_paragraph_mark(ctx: dict) -> list[str]:
    ins = ctx["revise"]["inserted"]
    out = []
    marks = [e for e in ins["elements"] if e["paragraph_mark"]]
    if not marks:
        out.append("K3 the appended paragraph's MARK is not marked as inserted; "
                   "rejecting then removes the text and leaves an empty paragraph")
    original = len(ctx["revise"]["original_texts"])
    if len(ins["accepted"]) != original + 1 or ins["accepted"][-1] != INSERTED_PARAGRAPH:
        out.append(f"K3 accepting gave {len(ins['accepted'])} paragraph(s) ending "
                   f"{ins['accepted'][-1:]!r}, expected {original + 1} ending with "
                   f"the inserted text")
    if len(ins["rejected"]) != original:
        out.append(f"K3 rejecting gave {len(ins['rejected'])} paragraph(s), expected "
                   f"{original} — an empty paragraph was left behind")
    return out


@check("K4", "accepting by author leaves the other author's revisions alone")
def k4_by_author(ctx: dict) -> list[str]:
    b = ctx["revise"]["by_author"]
    out = []
    if b["exit"] != 0:
        out.append(f"K4 --accept-author exited {b['exit']}")
    mine = [e for e in b["elements"] if e["author"] == REVISER]
    theirs = [e for e in b["elements"] if e["author"] == FIXTURE_AUTHOR]
    if mine:
        out.append(f"K4 {len(mine)} revision(s) by {REVISER} survived a filter that "
                   f"named that author")
    if not theirs:
        out.append(f"K4 {FIXTURE_AUTHOR}'s revisions are gone too — a per-author "
                   f"filter that resolves everything is not a filter")
    return out


@check("K5", "leftovers are reported, and --strict refuses to write them")
def k5_remaining(ctx: dict) -> list[str]:
    r = ctx["revise"]
    b = r["by_author"]
    out = []
    if r["accept"]["report"].get("remaining"):
        out.append(f"K5 accept-all left {r['accept']['report']['remaining']} behind "
                   f"and still reported success")
    if r["accept"]["elements"]:
        out.append(f"K5 accept-all wrote a document that still holds "
                   f"{len(r['accept']['elements'])} revision element(s)")
    if "remaining" not in r["accept"]["report"]:
        out.append("K5 the report has no `remaining` field, so a form the resolver "
                   "does not handle would be indistinguishable from success")
    if not b["report"].get("remaining"):
        out.append("K5 a per-author pass left revisions behind and did not say so")
    if b["strict_exit"] != 2 or b["strict_wrote"]:
        out.append(f"K5 --strict with leftovers exited {b['strict_exit']} "
                   f"(wrote={b['strict_wrote']})")
    return out


@check("K6", "text inserted and then deleted disappears whichever way it is resolved")
def k6_nested(ctx: dict) -> list[str]:
    n = ctx["revise"]["nested"]
    out = []
    nested = [e for e in n["elements"] if e["kind"] == "del"
              and e["text"] == CROSS_RUN_NEW]
    if not nested:
        out.append(f"K6 deleting the tracked insertion did not produce a deletion "
                   f"holding {CROSS_RUN_NEW!r}")
    i = TRACKED_PARAGRAPH
    if CROSS_RUN_NEW in n["accepted"][i]:
        out.append(f"K6 accepting left {CROSS_RUN_NEW!r} in place; it was inserted "
                   f"and then deleted, so accepting both means it is gone")
    if CROSS_RUN_NEW in n["rejected"][i]:
        out.append(f"K6 rejecting left {CROSS_RUN_NEW!r} in place; rejecting the "
                   f"insertion removes it whatever happened inside")
    return out


# ── W8: comments ──────────────────────────────────────────────────────────────
@check("B1", "adding a comment writes all five pieces, anchored across runs")
def b1_add(ctx: dict) -> list[str]:
    c = ctx["comment"]
    out = []
    if c["exit"] != 0:
        out.append(f"B1 adding a comment exited {c['exit']}")
    added = c["report"].get("added") or {}
    if not added.get("cross_run"):
        out.append(f"B1 the anchor {COMMENT_ANCHOR_SPLIT!r} spans two runs in the "
                   f"fixture and the comment did not report a cross-run anchor")
    if added.get("runs_wrapped", 0) < 2:
        out.append("B1 the range wraps a single run, so the split phrase was not "
                   "isolated — a range marker cannot start mid-run")
    p, before = c["pieces"], c["before_pieces"]
    for piece in ("part", "override", "relationship"):
        if not p[piece]:
            out.append(f"B1 the {piece} is missing from the result")
    for piece in ("range_starts", "range_ends", "references"):
        if p[piece] != before[piece] + 1:
            out.append(f"B1 {piece} went {before[piece]} -> {p[piece]}, expected one "
                       f"more")
    return out


@check("B2", "a document with no comments part gets one, wired up")
def b2_from_scratch(ctx: dict) -> list[str]:
    sc = ctx["comment"]["scratch"]
    out = []
    if sc["exit"] != 0:
        out.append(f"B2 adding to a document with no comments part exited {sc['exit']}")
    created = (sc["report"].get("added") or {}).get("package_pieces_created") or []
    for piece in ("part", "content-type Override", "relationship"):
        if piece not in created:
            out.append(f"B2 the report does not say it created the {piece}")
    for piece in ("part", "override", "relationship"):
        if not sc["pieces"][piece]:
            out.append(f"B2 the {piece} is missing — a comments part nothing points "
                       f"at is a comment Word never shows")
    return out


@check("B3", "deleting the last comment takes the part and both its indexes")
def b3_delete(ctx: dict) -> list[str]:
    d = ctx["comment"]["deleted"]
    out = []
    if d["exit"] != 0:
        out.append(f"B3 deleting exited {d['exit']}")
    rep = d["report"].get("deleted") or {}
    if sorted(rep.get("removed") or []) != ["1", "2"]:
        out.append(f"B3 removed {rep.get('removed')}, expected both comments")
    if not rep.get("comments_part_dropped"):
        out.append("B3 the last comment is gone and word/comments.xml is still there")
    for piece in ("part", "override", "relationship"):
        if d["pieces"][piece]:
            out.append(f"B3 the {piece} survived — removing only the bytes leaves a "
                       f"package pointing at nothing")
    for piece in ("range_starts", "range_ends", "references"):
        if d["pieces"][piece] != 0:
            out.append(f"B3 {d['pieces'][piece]} {piece} left in the body, anchored "
                       f"to a comment that no longer exists")
    return out


@check("B4", "a comment anchored to text that is not there is refused")
def b4_missing_anchor(ctx: dict) -> list[str]:
    m = ctx["comment"]["missing_anchor"]
    out = []
    if m["exit"] != 2:
        out.append(f"B4 anchoring to absent text exited {m['exit']}; a comment with "
                   f"no anchor is a balloon Word draws next to the wrong paragraph")
    if m["wrote"]:
        out.append("B4 it refused and wrote the file anyway")
    if "does not appear" not in m["stderr"]:
        out.append(f"B4 the refusal does not say why: {m['stderr']!r}")
    return out


@check("B5", "the listing says whether each comment is still anchored")
def b5_listing(ctx: dict) -> list[str]:
    c = ctx["comment"]
    out = []
    if c["list_exit"] != 0:
        out.append(f"B5 --list exited {c['list_exit']}")
    listing = c["report"].get("comments") or []
    if len(listing) != 2:
        out.append(f"B5 {len(listing)} comment(s) listed after adding one to a "
                   f"document that had one")
    if not all(e.get("anchored") for e in listing):
        out.append("B5 a comment is reported as unanchored when both have ranges")
    if not all(e.get("author") and e.get("text") for e in listing):
        out.append("B5 a comment is listed without its author or its text")
    return out


# ── W9: headers and footers ───────────────────────────────────────────────────
@check("H1", "a new header is four package pieces, and all four are written")
def h1_pieces(ctx: dict) -> list[str]:
    c = ctx["headerfooter"]["create"]
    out = []
    if c["exit"] != 0:
        out.append(f"H1 creating a header exited {c['exit']}")
    for kind, part in (("header", "word/header1.xml"), ("footer", "word/footer1.xml")):
        if part not in c["parts"]:
            out.append(f"H1 no {part} was written")
            continue
        if part not in c["overrides"]:
            out.append(f"H1 {part} has no content-type Override — Word treats a part "
                       f"nothing declares as damage")
        if part.split("/", 1)[1] not in c["rel_targets"]:
            out.append(f"H1 nothing in word/_rels/document.xml.rels points at {part}; "
                       f"the bytes are in the package and unreachable")
        ref = next((r for r in c["refs"] if r["kind"] == kind), None)
        if ref is None:
            out.append(f"H1 w:sectPr carries no {kind}Reference, so the part is never "
                       f"used by any page")
        elif ref["part"] != part:
            out.append(f"H1 the {kind}Reference resolves to {ref['part']}, not {part}")
    written = c["report"].get("written") or []
    if len(written) != 2 or not all(len(w.get("pieces") or []) >= 4 for w in written):
        out.append("H1 the report does not name the four pieces it wrote — a caller "
                   "cannot tell a complete write from a partial one")
    return out


@check("H2", "a first-page header without <w:titlePg/> is a part Word ignores")
def h2_title_page(ctx: dict) -> list[str]:
    f = ctx["headerfooter"]["first"]
    out = []
    if f["exit"] != 0:
        out.append(f"H2 creating a first-page header exited {f['exit']}")
    if "titlePg" not in f["sect_tags"]:
        out.append("H2 <w:titlePg/> is not in w:sectPr, so the first-page header was "
                   "written correctly in every other respect and page one still "
                   "shows the ordinary one")
    if not any(r["type"] == "first" for r in f["refs"]):
        out.append("H2 no headerReference of type 'first' was added")
    written = (f["report"].get("written") or [{}])[0]
    if "titlePg" not in (written.get("activated_by") or ""):
        out.append("H2 the report does not say what makes the variant take effect, "
                   "which is the one thing about it worth reporting")
    return out


@check("H3", "an even-page header needs a switch in settings.xml, not in the section")
def h3_even_pages(ctx: dict) -> list[str]:
    e = ctx["headerfooter"]["even"]
    out = []
    if e["exit"] != 0:
        out.append(f"H3 creating an even-page header exited {e['exit']}")
    if "evenAndOddHeaders" not in e["settings_tags"]:
        out.append("H3 <w:evenAndOddHeaders/> is not in word/settings.xml — it is a "
                   "DOCUMENT-wide setting, and without it the even header is a part "
                   "no page ever asks for")
    if "evenAndOddHeaders" in e["sect_tags"]:
        out.append("H3 <w:evenAndOddHeaders/> was written into w:sectPr, where it is "
                   "not a member of CT_SectPr and does nothing")
    if not any(r["type"] == "even" for r in e["refs"]):
        out.append("H3 no headerReference of type 'even' was added")
    return out


@check("H4", "removing a header takes its part and its switch with it")
def h4_remove(ctx: dict) -> list[str]:
    r = ctx["headerfooter"]["remove"]
    out = []
    if r["exit"] != 0:
        out.append(f"H4 removing a header exited {r['exit']}")
    if any(x["type"] == "first" for x in r["refs"]):
        out.append("H4 the first-page reference is still in w:sectPr")
    if "word/header2.xml" in r["parts"]:
        out.append("H4 word/header2.xml is still in the package with nothing pointing "
                   "at it — an orphan part is exactly what `drop` exists to prevent")
    if "word/header2.xml" in r["overrides"] or "header2.xml" in r["rel_targets"]:
        out.append("H4 the part is gone but its Override or its Relationship is not — "
                   "removing a part is three things, and this did fewer")
    if "titlePg" in r["sect_tags"]:
        out.append("H4 <w:titlePg/> is still on after the first-page header was "
                   "removed, so page one now has NO header at all — a change nobody "
                   "asked for and nobody sees until it prints")
    removed = (r["report"].get("removed") or [{}])[0]
    if not removed.get("part_dropped") or not removed.get("title_page_switched_off"):
        out.append("H4 the report does not say the part was dropped and the switch "
                   "turned off")
    return out


@check("H5", "the page number is a field, not a digit somebody cached")
def h5_page_field(ctx: dict) -> list[str]:
    c = ctx["headerfooter"]["create"]
    xml = c["footer_xml"]
    out = []
    if "fldChar" not in xml or "PAGE" not in xml:
        out.append("H5 the footer holds no PAGE field, so the page number is whatever "
                   "literal text was written there")
    if "<w:t>1</w:t>" in xml or "<w:t>1<" in xml:
        out.append("H5 the footer caches the digit 1 as the field result — nothing "
                   "here has laid the document out, and a cached 1 is how every page "
                   "ends up saying page 1")
    if 'w:fldCharType="separate"' in xml:
        out.append("H5 the field carries a `separate` marker, which opens a cached "
                   "RESULT; this field was written precisely because there is no "
                   "result to cache")
    footer = next((w for w in (c["report"].get("written") or [])
                   if w.get("kind") == "footer"), {})
    if not footer.get("page_number_field"):
        out.append("H5 the report does not say a page-number field was written")
    return out


# ── W10: contents and outline numbering ───────────────────────────────────────
@check("G1", "every heading is listed, at its own level, in document order")
def g1_headings(ctx: dict) -> list[str]:
    t = ctx["toc"]
    out = []
    if t["exit"] != 0:
        out.append(f"G1 the TOC run exited {t['exit']}")
    headings = t["report"].get("headings") or []
    if [h["level"] for h in headings] != list(OUTLINE_LEVELS):
        out.append(f"G1 heading levels came back as {[h['level'] for h in headings]}, "
                   f"expected {list(OUTLINE_LEVELS)} — a contents page that flattens "
                   f"levels is a list, not a structure")
    if [h["text"] for h in headings] != list(OUTLINE_HEADINGS):
        out.append(f"G1 the headings came back as {[h['text'] for h in headings]}")
    entries = t["entries"]
    if len(entries) != len(OUTLINE_HEADINGS):
        out.append(f"G1 the cached result holds {len(entries)} entries for "
                   f"{len(OUTLINE_HEADINGS)} headings")
    for heading, entry in zip(OUTLINE_HEADINGS, entries):
        if heading not in entry["text"]:
            out.append(f"G1 the entry for {heading!r} does not carry its text")
    styles = [e["style"] for e in entries]
    if styles != [f"TOC{min(lvl, 3)}" for lvl in OUTLINE_LEVELS]:
        out.append(f"G1 the entries carry styles {styles}, so the indentation does "
                   f"not follow the heading levels")
    if t["paragraphs"] != OUTLINE_PARAGRAPHS + TOC_PARAGRAPHS_ADDED:
        out.append(f"G1 the document has {t['paragraphs']} paragraphs, expected "
                   f"{OUTLINE_PARAGRAPHS + TOC_PARAGRAPHS_ADDED}")
    return out


@check("G2", "the cached result carries no page number, because none was computed")
def g2_no_page_numbers(ctx: dict) -> list[str]:
    t = ctx["toc"]
    out = []
    for entry in t["entries"]:
        tail = entry["text"].replace(entry["link_text"], "")
        if any(ch.isdigit() for ch in tail):
            out.append(f"G2 the entry for {entry['link_text']!r} carries {tail!r} "
                       f"where a page number goes — nothing here laid the document "
                       f"out, so that number was invented")
        elif TOC_PLACEHOLDER not in tail:
            out.append(f"G2 the entry for {entry['link_text']!r} has nothing where a "
                       f"page number goes; a blank is indistinguishable from a tool "
                       f"that forgot")
    toc = t["report"].get("toc") or {}
    if not toc.get("needs_update") or "not written" not in (toc.get("page_numbers") or ""):
        out.append("G2 the report does not say the page numbers are missing on "
                   "purpose, which is the whole product decision")
    if t["no_cache"]["paragraphs"] != OUTLINE_PARAGRAPHS + 2 or \
            t["no_cache"]["cached_entries"] != 0:
        out.append(f"G2 --no-cache produced {t['no_cache']['paragraphs']} paragraphs "
                   f"and {t['no_cache']['cached_entries']} cached entr(ies); it is "
                   f"supposed to write the field and NO result, and it must stay "
                   f"distinguishable from the cached mode")
    return out


@check("G3", "the field asks to be recalculated, in both of the two ways")
def g3_dirty(ctx: dict) -> list[str]:
    t = ctx["toc"]
    out = []
    if 'w:dirty="true"' not in t["document_xml"]:
        out.append("G3 the TOC field is not marked w:dirty, so a reader has no reason "
                   "to recompute the page numbers this document deliberately omits")
    if "updateFields" not in t["settings_tags"]:
        out.append("G3 <w:updateFields/> is not in word/settings.xml — the only "
                   "switch that asks a reader to refresh fields on open")
    if 'TOC \\o "1-3"' not in t["document_xml"].replace("\\\\", "\\"):
        out.append("G3 the field instruction is not a TOC over levels 1-3")
    return out


@check("G4", "every entry links to a bookmark that exists")
def g4_bookmarks(ctx: dict) -> list[str]:
    t = ctx["toc"]
    out = []
    names = set(t["bookmarks"])
    anchors = [a for e in t["entries"] for a in e["anchors"]]
    # Deliberately NOT "there are as many anchors as headings": how COMPLETE the
    # contents page is belongs to G1, and duplicating it here made a flaw that
    # shortens the list light up both checks, which hides which one owns the defect.
    # G4 owns one thing — that the links go somewhere.
    if not anchors:
        out.append("G4 no entry is a hyperlink, so a reader clicking the contents "
                   "page gets nothing and the \\h switch was written for no one")
    for anchor in anchors:
        if anchor not in names:
            out.append(f"G4 an entry anchors to {anchor!r}, which no bookmarkStart "
                       f"defines — Word shows 'Error! Bookmark not defined.'")
    for heading in (h["bookmark"] for h in t["report"].get("headings") or []):
        if heading not in names:
            out.append(f"G4 heading bookmark {heading!r} was reported but is not in "
                       f"the document")
    return out


@check("G5", "outline numbering is written in both halves, or it numbers nothing")
def g5_numbering(ctx: dict) -> list[str]:
    t = ctx["toc"]
    out = []
    levels = t["abstract_levels"]
    if [lvl["style"] for lvl in levels] != ["Heading1", "Heading2", "Heading3"]:
        out.append(f"G5 the abstractNum levels name {[l['style'] for l in levels]}, "
                   f"so they are not bound to the heading styles")
    if [lvl["text"] for lvl in levels] != ["%1.", "%1.%2", "%1.%2.%3"]:
        out.append(f"G5 the level texts are {[l['text'] for l in levels]}, not a "
                   f"1 / 1.1 / 1.1.1 outline")
    numbering = t["style_numbering"]
    num_id = (t["report"].get("numbering") or {}).get("num_id")
    for style in ("Heading1", "Heading2", "Heading3"):
        if style not in numbering:
            out.append(f"G5 {style} carries no w:numPr — an abstractNum whose levels "
                       f"name a w:pStyle numbers NOTHING until the style points back "
                       f"at it, and the XML looks complete either way")
        elif numbering[style] != str(num_id):
            out.append(f"G5 {style} points at numId {numbering[style]}, but the "
                       f"list that was authored is numId {num_id}")
    return out


@check("G6", "the contents heading does not take chapter number 1 for itself")
def g6_contents_not_numbered(ctx: dict) -> list[str]:
    t = ctx["toc"]
    numbering = t["style_numbering"]
    based = t["style_based_on"]
    out = []
    if TOC_HEADING_STYLE not in based:
        return [f"G6 {TOC_HEADING_STYLE} was not created, so the contents title is "
                f"formatted as body text"]
    # Walk the basedOn chain: numbering attached to a style is inherited by every
    # style built on it, and TOCHeading is built on Heading1 (as Word's own is).
    chain, current = [], based.get(TOC_HEADING_STYLE)
    while current and current not in chain:
        chain.append(current)
        current = based.get(current)
    inherits = [s for s in chain if s in numbering and numbering[s] != "0"]
    cancels = numbering.get(TOC_HEADING_STYLE) == "0"
    if inherits and not cancels:
        out.append(f"G6 {TOC_HEADING_STYLE} is based on {inherits[0]}, which is "
                   f"numbered, and does not cancel it with <w:numId w:val=\"0\"/> — "
                   f"measured in a rendered PDF, the contents page then takes number "
                   f"1 and the real first chapter becomes 2. Nothing in the package "
                   f"is invalid and no other check can see it")
    return out


# ── W11: pictures ─────────────────────────────────────────────────────────────
@check("M1", "a picture is four package pieces, and a package that has never held "
             "one has no content type for it")
def m1_pieces(ctx: dict) -> list[str]:
    i = ctx["image"]
    out = []
    if i["exit"] != 0:
        out.append(f"M1 inserting a picture exited {i['exit']}")
    if not any(n.startswith("word/media/") for n in i["parts"]):
        out.append("M1 no media part was written")
    if i["defaults"].get("png") != "image/png":
        out.append("M1 [Content_Types].xml declares no Default for .png — outline.docx "
                   "had never held one, and a package with an undeclared part is one "
                   "Word offers to repair")
    drawings = i["drawings"]
    if len(drawings) != 1:
        out.append(f"M1 {len(drawings)} drawing(s) in the document, expected 1")
    elif not drawings[0]["part"]:
        out.append(f"M1 the drawing's r:embed={drawings[0]['rid']} resolves to no "
                   f"relationship, so it points at nothing")
    if i["paragraphs"] != OUTLINE_PARAGRAPHS + 1:
        out.append(f"M1 the document has {i['paragraphs']} paragraphs, expected "
                   f"{OUTLINE_PARAGRAPHS + 1}")
    if i["body_children"] and i["body_children"][-1] != "sectPr":
        out.append("M1 <w:sectPr> is no longer last in the body")
    if not (i["report"].get("inserted") or {}).get("content_type_default_added"):
        out.append("M1 the report does not say the content type had to be declared")
    return out


@check("M2", "the size comes from the picture's own density, not from 96 dpi")
def m2_intrinsic(ctx: dict) -> list[str]:
    drawings = ctx["image"]["intrinsic"]["drawings"]
    out = []
    if not drawings or not drawings[0]["extent"]:
        return ["M2 the intrinsic insert produced no sized drawing"]
    extent = drawings[0]["extent"]
    if extent != list(CHART_INTRINSIC_EMU):
        out.append(f"M2 the extent is {extent} EMU, expected "
                   f"{list(CHART_INTRINSIC_EMU)} — chart.png declares 150 dpi in its "
                   f"pHYs chunk, and {list(CHART_96DPI_EMU)} is what assuming the "
                   f"web's 96 dpi gives instead")
    if extent[0] < 100000:
        out.append(f"M2 the extent is {extent[0]} EMU, which is a pixel count written "
                   f"into a field measured in 914400ths of an inch — the picture is "
                   f"a fraction of a millimetre wide and renders as nothing")
    described = (ctx["image"]["intrinsic"]["report"].get("inserted") or {})
    if described.get("pixels") != list(CHART_PX):
        out.append(f"M2 the report says the picture is {described.get('pixels')} "
                   f"pixels, measured {list(CHART_PX)}")
    if "read from the file" not in (described.get("density") or ""):
        out.append("M2 the report does not say whether the density was measured or "
                   "assumed, which is the difference between the two numbers above")
    return out


@check("M3", "--width-cm scales it and keeps the aspect ratio")
def m3_scaled(ctx: dict) -> list[str]:
    drawings = ctx["image"]["drawings"]
    out = []
    if not drawings or not drawings[0]["extent"]:
        return ["M3 the scaled insert produced no sized drawing"]
    extent = drawings[0]["extent"]
    if extent != list(IMAGE_WIDTH_EMU):
        out.append(f"M3 the extent is {extent} EMU; {IMAGE_WIDTH_CM}cm at the "
                   f"fixture's 2:1 ratio is {list(IMAGE_WIDTH_EMU)}")
    if extent[1] and abs(extent[0] / extent[1] - CHART_PX[0] / CHART_PX[1]) > 0.01:
        out.append(f"M3 the ratio came out {extent[0] / extent[1]:.3f}, the picture's "
                   f"is {CHART_PX[0] / CHART_PX[1]:.3f} — scaling one side and "
                   f"leaving the other stretches it, and nothing errors")
    return out


@check("M4", "the two numbers that state one size agree")
def m4_extent_agrees(ctx: dict) -> list[str]:
    out = []
    for label, drawings in (("scaled", ctx["image"]["drawings"]),
                            ("intrinsic", ctx["image"]["intrinsic"]["drawings"])):
        for d in drawings:
            if d["extent"] and d["inner"] and d["extent"] != d["inner"]:
                out.append(f"M4 ({label}) <wp:extent> says {d['extent']} and <a:ext> "
                           f"says {d['inner']} — Word lays out a box of one size and "
                           f"stretches a picture of the other into it, which reads as "
                           f"a blurry export rather than as a defect")
    return out


@check("M5", "replacing a picture rewires it without orphaning the old part")
def m5_replace(ctx: dict) -> list[str]:
    same = ctx["image"]["replace_same"]
    other = ctx["image"]["replace_other"]
    out = []
    if same["exit"] != 0 or other["exit"] != 0:
        out.append(f"M5 replace exited {same['exit']} / {other['exit']}")
    media = [n for n in same["parts"] if n.startswith("word/media/")]
    if len(media) != 1:
        out.append(f"M5 same-format replace left {len(media)} media part(s): {media} — "
                   f"the relationship and the content type already said the right "
                   f"thing, so only the bytes had to change")
    if same["media_bytes"].get("word/media/image1.png") == CHART.stat().st_size:
        out.append("M5 same-format replace did not change the bytes at all")
    if same["drawings"] and same["drawings"][0]["part"] != "word/media/image1.png":
        out.append("M5 same-format replace repointed the drawing when it did not "
                   "need to")

    media = [n for n in other["parts"] if n.startswith("word/media/")]
    if any(n.endswith(".png") for n in media):
        out.append(f"M5 the replaced .png is still in the package: {media} — nothing "
                   f"points at it, which is precisely the orphan `drop` exists for")
    if "image1.png" in other["rel_targets"]:
        out.append("M5 the old relationship survived the replacement, so the package "
                   "has a link to bytes that are gone")
    if other["defaults"].get("gif") != "image/gif":
        out.append("M5 no Default was declared for the new format's extension")
    if other["drawings"] and not other["drawings"][0]["part"]:
        out.append("M5 the drawing's r:embed no longer resolves after the replacement")
    return out


# ── W16: font bindings ────────────────────────────────────────────────────────
@check("N1", "the audit finds all three shapes of a missing binding")
def n1_audit(ctx: dict) -> list[str]:
    f = ctx["fonts"]
    out = []
    if f["exit"] != 0:
        out.append(f"N1 --fix exited {f['exit']}")
    problems = f["report"].get("problems") or []
    if len(problems) != FONTLESS_UNBOUND_RUNS:
        out.append(f"N1 {len(problems)} unbound run(s) found, the fixture carries "
                   f"{FONTLESS_UNBOUND_RUNS} and they are three different shapes: no "
                   f"w:rPr at all, a style that already answers, and a w:rFonts with "
                   f"only @w:ascii")
    shapes = {tuple(p["missing"]) for p in problems}
    if ("eastAsia",) not in shapes:
        out.append("N1 the run carrying @w:ascii and no @w:eastAsia was not found — "
                   "an audit that only looks for a missing w:rFonts misses the half "
                   "that is hardest to see")
    if f["still_unbound"]:
        out.append(f"N1 after --fix, {len(f['still_unbound'])} CJK run(s) still state "
                   f"neither face: {f['still_unbound'][:2]}")
    if f["parts_changed"] != ["word/document.xml"]:
        out.append(f"N1 --fix rewrote {f['parts_changed']}; only the part holding the "
                   f"unbound runs should have changed")
    return out


@check("N2", "a face the document already stated is kept, not replaced by the default")
def n2_inherited(ctx: dict) -> list[str]:
    f = ctx["fonts"]
    out = []
    inherited = [p for p in f["report"].get("problems") or []
                 if any(v.get("from", "").startswith("style:")
                        for v in (p.get("written") or {}).values())]
    if not inherited:
        out.append("N2 no run resolved its face from a style, so 'keep what the "
                   "document said' was never exercised")
    for entry in inherited:
        value = entry["written"].get("eastAsia", {}).get("value")
        if value != FONTLESS_STYLE_EA:
            out.append(f"N2 the run under Heading2 was bound to {value!r}; its style "
                       f"already says {FONTLESS_STYLE_EA!r}, and writing the default "
                       f"over it restyles a heading the author chose a face for — a "
                       f"change D6 cannot tell apart from a repair")
    if FONTLESS_STYLE_EA not in f["document_xml"]:
        out.append(f"N2 {FONTLESS_STYLE_EA!r} appears nowhere in the repaired document")
    sources = f["report"].get("sources") or {}
    if not any(k.startswith("style:") for k in sources):
        out.append("N2 the report does not say which faces came from the document and "
                   "which from the fallback")
    return out


@check("N3", "a latin face the run chose deliberately survives the repair")
def n3_keeps_ascii(ctx: dict) -> list[str]:
    f = ctx["fonts"]
    out = []
    if FONTLESS_KEPT_ASCII not in f["document_xml"]:
        out.append(f"N3 {FONTLESS_KEPT_ASCII!r} is gone from the repaired document — "
                   f"that run was missing only its @w:eastAsia, and overwriting the "
                   f"@w:ascii it did state changes text nobody asked to change")
    half = [p for p in f["report"].get("problems") or [] if p["missing"] == ["eastAsia"]]
    if half and "ascii" in (half[0].get("written") or {}):
        out.append("N3 the repair wrote an @w:ascii onto a run that already had one")
    return out


@check("N4", "--check reports and writes nothing, and it reads headers and footers too")
def n4_check_only(ctx: dict) -> list[str]:
    c = ctx["fonts"]["check_clean"]
    out = []
    if c["exit"] != 0:
        out.append(f"N4 --check on an already-correct document exited {c['exit']}")
    if c["report"].get("unbound_runs") != 0:
        out.append(f"N4 --check reports {c['report'].get('unbound_runs')} unbound "
                   f"run(s) in report.docx, whose every run states both faces — an "
                   f"audit that cries wolf is one people turn off")
    parts = c["report"].get("parts_examined") or []
    if not any(p.startswith("word/header") for p in parts) or \
            not any(p.startswith("word/footer") for p in parts):
        out.append(f"N4 --check examined {parts}; a letterhead is exactly the text "
                   f"most likely to have been pasted in with its own fonts")
    if c["report"].get("out") or c["wrote_anything"]:
        out.append("N4 --check wrote a document; it is an inspection")
    if not c["report"].get("verdict"):
        out.append("N4 --check gives no verdict, so a clean document and a crashed "
                   "run read the same")
    return out


@check("N5", "--strict refuses when the face came from this tool and not the document")
def n5_strict(ctx: dict) -> list[str]:
    f = ctx["fonts"]
    s = f["strict"]
    out = []
    if s["exit"] == 0:
        out.append("N5 --strict exited 0 on a document where two runs had no face "
                   "stated anywhere — not on the run, not on a style, not in "
                   "w:docDefaults. That value is this tool's choice, not the "
                   "document's, and --strict exists to say so")
    if s["wrote"]:
        out.append("N5 --strict wrote the file anyway, which makes the refusal a "
                   "message rather than a refusal")
    if "docDefaults" not in s["stderr"] and "fallback" not in s["stderr"]:
        out.append(f"N5 the refusal does not say why: {s['stderr'][:120]!r}")
    if not f["report"].get("fallback_used"):
        out.append("N5 the non-strict report does not name the runs whose face came "
                   "from the fallback, so the same information is unavailable to a "
                   "caller who did not pass --strict")
    return out


# ── W12: styles ───────────────────────────────────────────────────────────────
@check("S1", "a new style is written in CT_Style order and touches nothing else")
def s1_create(ctx: dict) -> list[str]:
    c = ctx["styles"]["create"]
    out = []
    if c["exit"] != 0:
        out.append(f"S1 creating a style exited {c['exit']}")
    style = c["styles"].get(NEW_STYLE)
    if style is None:
        return out + [f"S1 {NEW_STYLE!r} is not in word/styles.xml"]
    # CT_Style is an xsd:sequence: name, basedOn, next, ..., pPr, rPr. Out of order
    # is invalid even though every element in it is spelled right (gotchas §21.2 ㉓).
    order = ["name", "basedOn", "next", "uiPriority", "pPr", "rPr"]
    ranks = [order.index(k) for k in style["children"] if k in order]
    if ranks != sorted(ranks):
        out.append(f"S1 the style's children are {style['children']}, which is not "
                   f"CT_Style order — Word repairs the file and drops what it cannot "
                   f"place")
    if not style["custom"]:
        out.append("S1 the new style is not marked w:customStyle, so Word files it "
                   "with its own built-ins")
    if style["based_on"] != "Normal":
        out.append(f"S1 w:basedOn is {style['based_on']!r}, expected 'Normal'")
    if "rFonts" not in style["rpr"]:
        out.append("S1 the style sets a size but binds no font faces — a CJK style "
                   "with no @w:eastAsia renders in whatever the reader's theme picks")
    if c["parts_changed"] != ["word/styles.xml"]:
        out.append(f"S1 the edit rewrote {c['parts_changed']}; a style change belongs "
                   f"in word/styles.xml and nowhere else")
    return out


@check("S2", "modifying a style that exists is refused until the caller says so")
def s2_overwrite(ctx: dict) -> list[str]:
    s = ctx["styles"]
    out = []
    if s["no_overwrite"]["exit"] == 0:
        out.append("S2 --set on an existing style succeeded without --overwrite. A "
                   "style is shared: the change repaints every paragraph using it, "
                   "and the request that asked for it usually named one")
    if s["no_overwrite"]["wrote"]:
        out.append("S2 it wrote the file anyway, which makes the refusal a message "
                   "rather than a refusal")
    if str(HEADING_USERS) not in s["no_overwrite"]["stderr"]:
        out.append(f"S2 the refusal does not say how many would be repainted: "
                   f"{s['no_overwrite']['stderr'][:120]!r}")
    if s["overwrite"]["exit"] != 0:
        out.append(f"S2 --overwrite exited {s['overwrite']['exit']}")
    elif s["overwrite"]["report"]["set"].get("repainted") != HEADING_USERS:
        out.append(f"S2 the report says it repainted "
                   f"{s['overwrite']['report']['set'].get('repainted')}, "
                   f"expected {HEADING_USERS}")
    return out


@check("S3", "the report names every property it set, and the file agrees")
def s3_report_truthful(ctx: dict) -> list[str]:
    ow = ctx["styles"]["overwrite"]
    out = []
    said = set(ow["report"]["set"].get("properties_set") or [])
    got = set(ow["styles"].get("Heading1", {}).get("rpr") or [])
    # The defect this pins: `insert_ordered` REPARENTS each child, so counting them
    # after the merge counted an empty shell — the report said `properties_set: []`
    # on a call that had just set six. The document was right and the REPORT was
    # wrong, which is the half nobody checks.
    for want in ("color", "u", "i", "sz"):
        if want not in said:
            out.append(f"S3 the report does not mention w:{want}, which was asked for")
        if want not in got:
            out.append(f"S3 w:{want} is not in the style, though it was asked for")
    if ow["styles"].get("Heading1", {}).get("sz") != "40":
        out.append(f"S3 --size 20 wrote w:sz={ow['styles'].get('Heading1', {}).get('sz')!r}; "
                   f"w:sz is in HALF-points, so 20pt is 40")
    return out


@check("S4", "deleting a style in use is refused, and --reassign moves what used it")
def s4_delete(ctx: dict) -> list[str]:
    s = ctx["styles"]
    out = []
    if s["delete_in_use"]["exit"] == 0 or s["delete_in_use"]["wrote"]:
        out.append("S4 a style still in use was deleted. Word raises nothing — the "
                   "paragraphs fall back to Normal and the document quietly loses "
                   "its headings")
    if s["reassign"]["exit"] != 0:
        out.append(f"S4 --reassign exited {s['reassign']['exit']}")
        return out
    if "Heading1" in s["reassign"]["styles"]:
        out.append("S4 the style is still in word/styles.xml after --delete")
    if "Heading1" in s["reassign"]["paragraph_styles"]:
        out.append("S4 a paragraph still names the deleted style, so it now points "
                   "at nothing")
    if s["reassign"]["report"]["deleted"].get("reassigned") != HEADING_USERS:
        out.append(f"S4 the report says it moved "
                   f"{s['reassign']['report']['deleted'].get('reassigned')} "
                   f"paragraph(s), expected {HEADING_USERS}")
    # Everything based on the removed style has to be repointed, not left dangling:
    # Word treats a missing w:basedOn as no inheritance at all.
    for sid, style in s["reassign"]["styles"].items():
        if style["based_on"] == "Heading1":
            out.append(f"S4 {sid!r} is still based on the style that was deleted")
    return out


@check("S5", "a w:basedOn cycle is refused rather than written")
def s5_cycle(ctx: dict) -> list[str]:
    c = ctx["styles"]["cycle"]
    out = []
    if c["exit"] == 0 or c["wrote"]:
        out.append("S5 a style was made to inherit from itself. Word stops resolving "
                   "formatting at the loop and renders it as Normal, with no error "
                   "anywhere — so nothing downstream would report this either")
    if "itself" not in c["stderr"] and "->" not in c["stderr"]:
        out.append(f"S5 the refusal does not name the loop: {c['stderr'][:120]!r}")
    return out



# ── W4: generating from Markdown ──────────────────────────────────────────────
@check("D1", "every block kind in the Markdown reaches the document")
def d1_blocks(ctx: dict) -> list[str]:
    m = ctx["markdown"]
    out = []
    if m["exit"] != 0:
        out.append(f"D1 generating exited {m['exit']}")
    if m["paragraphs"] != MD_PARAGRAPHS:
        out.append(f"D1 the document has {m['paragraphs']} paragraphs, expected "
                   f"{MD_PARAGRAPHS}")
    blocks = m["report"]["written"]["blocks"]
    for kind in ("heading", "paragraph", "list_item", "quote", "table", "code",
                 "rule", "image"):
        if not blocks.get(kind):
            out.append(f"D1 no {kind} block was written, though sample.md has one — "
                       f"a generator that drops a block kind produces a document "
                       f"that is missing content with nothing saying so")
    if m["body_children"] and m["body_children"][-1] != "sectPr":
        out.append("D1 <w:sectPr> is no longer last in the body after appending "
                   "every block")
    return out


@check("D2", "what it cannot map is named with its line, not dropped")
def d2_unsupported(ctx: dict) -> list[str]:
    m = ctx["markdown"]
    out = []
    reported = m["report"]["unsupported"]
    if len(reported) != MD_UNSUPPORTED:
        out.append(f"D2 {len(reported)} unmappable construct(s) reported, sample.md "
                   f"carries {MD_UNSUPPORTED} on purpose (a footnote and two raw "
                   f"HTML tags). A generator that silently discards them still "
                   f"produces a legal document that passes every other check")
    kinds = {u["construct"] for u in reported}
    for want in ("footnote reference", "raw HTML tag"):
        if want not in kinds:
            out.append(f"D2 a {want} in the source is not reported")
    if any(not u.get("line") for u in reported):
        out.append("D2 a reported construct carries no line number, so nobody can "
                   "find it in the source")
    return out


@check("D3", "--strict refuses and writes nothing when something could not be mapped")
def d3_strict(ctx: dict) -> list[str]:
    s = ctx["markdown"]["strict"]
    out = []
    if s["exit"] == 0:
        out.append("D3 --strict exited 0 on a source carrying constructs the "
                   "generator cannot map")
    if s["wrote"]:
        out.append("D3 --strict wrote the document anyway, which makes the refusal "
                   "a message rather than a refusal")
    if "strict" not in s["stderr"]:
        out.append(f"D3 the refusal does not say why: {s['stderr'][:120]!r}")
    return out


@check("D4", "the styles the document names are the styles it defines")
def d4_styles(ctx: dict) -> list[str]:
    m = ctx["markdown"]
    out = []
    defined = set(m["styles"])
    for style in MD_STYLES:
        if style not in defined:
            out.append(f"D4 {style!r} is used but not defined — a w:pStyle naming a "
                       f"style the document does not have is valid XML that silently "
                       f"formats as Normal, and nothing anywhere explains why")
    for used in m["paragraph_styles"]:
        if used and used not in defined:
            out.append(f"D4 a paragraph names style {used!r}, which is not defined")
    return out


@check("D5", "nested list levels reach the numbering they claim")
def d5_lists(ctx: dict) -> list[str]:
    m = ctx["markdown"]
    out = []
    numbering = m["numbering"]
    if len(numbering) < 2:
        out.append(f"D5 {len(numbering)} list(s) defined; a bullet list and an "
                   f"ordered list are two different abstractNums")
    for num_id, levels in numbering.items():
        if len(levels) <= MD_LIST_DEPTH:
            out.append(f"D5 numId {num_id} defines {len(levels)} level(s); the "
                       f"fixture nests {MD_LIST_DEPTH + 1} deep and a level that is "
                       f"not defined renders flat")
    if str(MD_LIST_DEPTH) not in m["levels_used"]:
        out.append(f"D5 no paragraph uses ilvl={MD_LIST_DEPTH}, so the nested bullet "
                   f"in sample.md came out at the top level")
    return out


@check("D6", "--template keeps the house style and replaces only the body")
def d6_template(ctx: dict) -> list[str]:
    tpl = ctx["markdown"]["template"]
    out = []
    if tpl["exit"] != 0:
        out.append(f"D6 --template exited {tpl['exit']}")
        return out
    if tpl["lost"]:
        out.append(f"D6 the template lost {tpl['lost']} — its styles, numbering, "
                   f"headers and footers are the whole reason to pass one")
    if not any("header" in p for p in tpl["parts"]):
        out.append("D6 the template's header is gone")
    if tpl["report"]["template_blocks_replaced"] < 1:
        out.append("D6 the report says no template block was replaced, so the "
                   "generated content was appended to the template's own text")
    for old in ("二零二六年第三季度经营分析报告",):
        if any(old in text for text in tpl["texts"]):
            out.append(f"D6 the template's own body text {old!r} is still in the "
                       f"document — only its styles were meant to survive")
    return out



# ── W14: schema validation ────────────────────────────────────────────────────
@check("Z1", "the schemas ship, and a legal document validates against them")
def z1_ships(ctx: dict) -> list[str]:
    v = ctx["validate"]
    out = []
    if len(v["schema_files"]) != SCHEMA_FILES:
        out.append(f"Z1 {len(v['schema_files'])} schema file(s) ship with the skill, "
                   f"expected {SCHEMA_FILES} — the closure of wml.xsd. A validator "
                   f"whose grammar the user has to fetch first is one nobody runs")
    if v["good"]["exit"] != 0:
        out.append(f"Z1 a document known to be legal exited {v['good']['exit']}")
    report = v["good"]["report"]
    if not report.get("valid"):
        out.append(f"Z1 report.docx is reported invalid: "
                   f"{report.get('violations', [])[:1]}")
    if report["parts_checked"] != VALIDATED_PARTS:
        out.append(f"Z1 {report['parts_checked']} part(s) checked, expected "
                   f"{VALIDATED_PARTS}")
    if report["schemas"]["resolved_from"] != "bundled with the skill":
        out.append(f"Z1 the grammar came from {report['schemas']['resolved_from']!r} "
                   f"rather than the copy that ships — on a user's machine there is "
                   f"nothing else to fall back to")
    return out


@check("Z2", "a genuinely invalid document is rejected, with the part and the line")
def z2_rejects(ctx: dict) -> list[str]:
    inv = ctx["validate"]["invalid"]
    out = []
    if inv["exit"] == 0:
        out.append("Z2 a document whose w:zoom is missing the REQUIRED w:percent was "
                   "accepted — that is the exact defect python-docx's own template "
                   "carries, so this is not a hypothetical")
    report = inv["report"]
    if report.get("valid"):
        out.append("Z2 the report says valid")
    if not report.get("violations"):
        out.append("Z2 no violation was reported")
        return out
    first = report["violations"][0]
    if first.get("part") != "word/settings.xml" or not first.get("line"):
        out.append(f"Z2 the violation does not locate itself: {first}")
    return out


@check("Z3", "no schemas is a loud failure, never a quiet pass")
def z3_loud(ctx: dict) -> list[str]:
    m = ctx["validate"]["missing"]
    out = []
    if m["exit"] == 0:
        out.append("Z3 with no schemas anywhere it exited 0. A run that checked "
                   "NOTHING must not be reportable as a pass — that is the single "
                   "outcome this capability was not allowed to have")
    if '"valid": true' in m["stdout"].replace(" ", "").lower():
        out.append("Z3 it printed a valid verdict from a run that validated nothing")
    if "tried" not in m["stderr"]:
        out.append(f"Z3 the failure does not name where it looked: "
                   f"{m['stderr'][:120]!r}")
    return out


@check("Z4", "mc:Ignorable is honoured, or every real Word document cries wolf")
def z4_ignorable(ctx: dict) -> list[str]:
    ig = ctx["validate"]["ignorable"]
    out = []
    if ig["exit"] != 0 or not ig["report"].get("valid"):
        out.append(f"Z4 a document carrying an element from a namespace marked "
                   f"mc:Ignorable was reported invalid "
                   f"({ig['report'].get('violation_count')} violation(s)). Every "
                   f"document Word itself writes looks like this (w14, w15, wp14), "
                   f"so a validator that does not strip them reports a wall of "
                   f"non-defects and gets switched off")
    return out


@check("Z5", "parts with no grammar are named, so 'valid' cannot be misread")
def z5_not_checked(ctx: dict) -> list[str]:
    report = ctx["validate"]["good"]["report"]
    out = []
    if "not_checked" not in report:
        out.append("Z5 the report does not say which parts were left unchecked, so "
                   "'valid' reads as 'everything was checked' when it means 'valid "
                   "where a grammar existed'")
    checked = set(report.get("checked") or [])
    if "word/document.xml" not in checked:
        out.append("Z5 word/document.xml is not among the checked parts")
    return out


# ── the negative controls ─────────────────────────────────────────────────────
# Each one is a defect an assertion above claims to catch, applied to the collected
# context. They are the implementations somebody reaches for first, not invented
# damage.

# ── W18: document diff ────────────────────────────────────────────────────────
@check("U1", "every whitelisted change is found, and nothing else is called a change")
def u1_findings(ctx: dict) -> list[str]:
    d = ctx["diff"]["main"]
    out = []
    if d["exit"] != 0:
        return [f"U1 the comparison exited {d['exit']}: {d['stderr'][:120]}"]
    found = d["report"]["differences"]
    kinds = sorted(f["kind"] for f in found)
    want = sorted(["text", "text", "text", "style", "paragraph-removed",
                   "paragraph-added"])
    if kinds != want:
        out.append(f"U1 the six changes came back as {kinds}, expected {want}")
    if d["report"]["counted"] != DIFF_EXPECTED:
        out.append(f"U1 counted {d['report']['counted']} differences, expected "
                   f"{DIFF_EXPECTED}")
    # And each one has to be found WHERE it is, not just in the right quantity.
    places = {f["kind"]: [] for f in found}
    for f in found:
        places[f["kind"]].append(f.get("where", ""))
    if not any(DIFF_HEADER_PART in w for w in places.get("text", [])):
        out.append(f"U1 nothing was found in {DIFF_HEADER_PART}; a diff that only "
                   f"walks word/document.xml misses every letterhead edit")
    if not any("table" in w for w in places.get("text", [])):
        out.append("U1 the table cell edit was not found; table cell text is on the "
                   "whitelist and its paragraphs are nested two levels down")
    style = [f for f in found if f["kind"] == "style"]
    if style and style[0].get("after") != DIFF_NEW_STYLE:
        out.append(f"U1 the style change reads {style[0].get('after')!r}, expected "
                   f"{DIFF_NEW_STYLE!r}")
    return out


@check("U2", "accept every marked change and it reads like B; reject them and it reads like A")
def u2_roundtrip(ctx: dict) -> list[str]:
    d = ctx["diff"]
    out = []
    if d["accept"]["signature"] != d["signatures"]["b"]:
        out.append("U2 accepting the redline's changes does not give B — the redline "
                   "describes an edit that is not the edit that was made"
                   + _sig_gap(d["accept"]["signature"], d["signatures"]["b"]))
    if d["reject"]["signature"] != d["signatures"]["a"]:
        out.append("U2 rejecting the redline's changes does not give A back"
                   + _sig_gap(d["reject"]["signature"], d["signatures"]["a"]))
    # The script's own in-process verdict has to agree with the one measured here by
    # resolving the file with docx_revise.py. If it claims a round trip this file
    # cannot reproduce, the claim is the defect.
    claimed = d["main"]["report"].get("roundtrip", {})
    if not claimed.get("exact"):
        out.append(f"U2 the script does not claim an exact round trip on a pair where "
                   f"every change is on its whitelist: {claimed}")
    return out


def _sig_gap(got: list, want: list) -> str:
    for i, (x, y) in enumerate(zip(got, want)):
        if x != y:
            return f" — first difference at paragraph {i}: {x!r} vs {y!r}"
    return f" — lengths {len(got)} vs {len(want)}"


@check("U3", "what Word rewrites on every save is counted, and counted as NOT a change")
def u3_noise(ctx: dict) -> list[str]:
    d = ctx["diff"]
    out = []
    ignored = d["main"]["report"].get("ignored_not_counted_as_differences", {})
    for key in DIFF_IGNORED_KEYS:
        if key not in ignored:
            out.append(f"U3 the report never mentions {key!r}; a category that is "
                       f"named only when it fires cannot be told from one nobody "
                       f"looked at")
    noise = d["noise"]
    if noise["exit"] != 0:
        return out + [f"U3 the noise-only comparison exited {noise['exit']}"]
    if noise["report"]["counted"] != 0:
        out.append(f"U3 a document differing ONLY in rsids, proofErr, bookmarks and "
                   f"lang produced {noise['report']['counted']} difference(s): "
                   f"{[f['kind'] for f in noise['report']['differences']][:4]} — this "
                   f"is the diff nobody reads, and it is what the capability exists "
                   f"not to be")
    seen = noise["report"].get("ignored_not_counted_as_differences", {})
    for key in ("rsid", "proofErr", "bookmark", "lang"):
        pair = seen.get(key) or {}
        if not isinstance(pair, dict) or pair.get("a") == pair.get("b"):
            out.append(f"U3 {key} reads {pair} for a pair that differs in exactly "
                       f"that — the category is being reported without being counted")
    return out


@check("U4", "a change it cannot mark is named, and never left looking reviewed")
def u4_honest_about_gaps(ctx: dict) -> list[str]:
    rows = ctx["diff"]["rows"]
    out = []
    if rows["exit"] != 0:
        return [f"U4 the row-added comparison exited {rows['exit']}"]
    report = rows["report"]
    if not any(f["kind"] == "table-shape" for f in report["differences"]):
        out.append("U4 a table row appeared and no difference was reported for it")
    if not report["not_redlined"]:
        out.append("U4 the row change is not expressible as a tracked change by this "
                   "capability, yet nothing said so — four of five differences marked "
                   "is worse than none, because the fifth now looks reviewed")
    if report.get("roundtrip", {}).get("exact"):
        out.append("U4 an exact round trip was claimed for a comparison that left a "
                   "difference unmarked")
    strict = ctx["diff"]["rows_strict"]
    if strict["exit"] == 0 or strict["wrote"]:
        out.append("U4 --strict wrote a redline that does not round-trip")
    return out


@check("U5", "a paragraph that moved is reported as moved, not as a delete and an insert")
def u5_move(ctx: dict) -> list[str]:
    moved = ctx["diff"]["moved"]
    if moved["exit"] != 0:
        return [f"U5 the move comparison exited {moved['exit']}"]
    kinds = [f["kind"] for f in moved["report"]["differences"]]
    out = []
    if "paragraph-moved" not in kinds:
        out.append(f"U5 moving a paragraph was reported as {sorted(set(kinds))} — a "
                   f"reader then checks two findings that are really one")
    if "paragraph-removed" in kinds or "paragraph-added" in kinds:
        out.append(f"U5 the move was ALSO reported as a removal or an addition: "
                   f"{sorted(set(kinds))}")
    return out


@check("U6", "the redline is a legal Word document carrying real tracked changes")
def u6_redline(ctx: dict) -> list[str]:
    d = ctx["diff"]
    out = []
    valid = d["redline_valid"]["report"]
    if not valid.get("valid"):
        out.append(f"U6 the redline does not validate against the schema: "
                   f"{valid.get('violations', [])[:2]}")
    kinds = {r["kind"] for r in d["redline_revisions"].get("revisions", [])}
    for needed in ("ins", "del", "pPrChange"):
        if needed not in kinds:
            out.append(f"U6 the redline carries no <w:{needed}>; the {needed} half of "
                       f"the whitelist was reported but not written into the document")
    # Marking a paragraph as deleted has to mark its BREAK too, or accepting the
    # change empties the paragraph and leaves it there.
    marks = [r for r in d["redline_revisions"].get("revisions", [])
             if r.get("scope") == "paragraph-mark"]
    if not marks:
        out.append("U6 no revision is on a paragraph MARK, so the added and removed "
                   "paragraphs were marked by their text alone — accepting that "
                   "leaves an empty paragraph behind")
    touched = set(d["redline_parts_changed"])
    stray = {n for n in touched
             if not (n == DOC_PART or n.startswith(("word/header", "word/footer")))}
    if stray:
        out.append(f"U6 the redline rewrote {sorted(stray)}; a redline is tracked "
                   f"changes in the text parts and nothing else")
    return out


@check("U7", "a picture that changed is reported, and reported as not redlined")
def u7_images(ctx: dict) -> list[str]:
    image = ctx["diff"]["image"]
    if image["exit"] != 0:
        return [f"U7 the image comparison exited {image['exit']}: "
                f"{image['stderr'][:120]}"]
    report = image["report"]
    kinds = [f["kind"] for f in report["differences"]]
    out = []
    if "image-replaced" not in kinds:
        out.append(f"U7 the same picture slot holding different bytes was reported as "
                   f"{sorted(set(kinds)) or 'nothing'} — a diff that reads only text "
                   f"cannot see a swapped chart, which is the change most worth "
                   f"catching in a report")
    if not any("media" in note for note in report["not_redlined"]):
        out.append("U7 the picture change was counted but nothing says it was left "
                   "for a person; copying a media part across is a package edit, not "
                   "a tracked change")
    if report.get("roundtrip", {}).get("exact"):
        out.append("U7 an exact round trip was claimed while a picture change went "
                   "unmarked")
    return out


# ── W19: table presets ────────────────────────────────────────────────────────
@check("Q1", "each preset writes the border weights it advertises, in eighths of a point")
def q1_borders(ctx: dict) -> list[str]:
    t = ctx["tables"]
    out = []
    listed = {p["name"]: p for p in t["listed"]["presets"]}
    for name in TABLE_PRESETS:
        got = t["presets"][name]["facts"]["borders"]
        if got != PRESET_BORDERS[name]:
            out.append(f"Q1 {name} produced borders {got}, expected "
                       f"{PRESET_BORDERS[name]} (eighths of a point; a preset that "
                       f"writes points draws hairlines and nothing complains)")
        # --list-presets is how a caller chooses. If it describes something the file
        # does not carry, the choice was made on a false description.
        claimed = listed.get(name, {}).get("border_eighths")
        if claimed != PRESET_BORDERS[name]:
            out.append(f"Q1 --list-presets says {name} draws {claimed}, but the "
                       f"expected weights are {PRESET_BORDERS[name]}")
        if t["presets"][name]["facts"]["cell_margin"] != WIDE_MARGIN[name]:
            out.append(f"Q1 {name} wrote cell margin "
                       f"{t['presets'][name]['facts']['cell_margin']}, expected "
                       f"{WIDE_MARGIN[name]}")
    # An omitted border is not the same as `w:val="none"`: the sample's table names a
    # w:tblStyle, whose borders show through anything this skill declines to state.
    fin = t["presets"]["finance"]["facts"]["borders"]
    if any(v == -1 for v in fin.values()):
        out.append("Q1 finance left an edge unstated rather than writing an explicit "
                   "none, so the table style's own border shows through")
    return out


@check("Q2", "no two presets produce the same measured fingerprint")
def q2_distinct(ctx: dict) -> list[str]:
    t = ctx["tables"]
    out = []
    seen: dict[str, str] = {}
    for name in TABLE_PRESETS:
        key = json.dumps(t["presets"][name]["facts"], sort_keys=True,
                         ensure_ascii=False)
        if key in seen:
            out.append(f"Q2 {name} and {seen[key]} produce identical measured "
                       f"tables — one of them is a name with nothing behind it, "
                       f"which is the whole reason this capability shipped three "
                       f"presets and not thirteen")
        seen[key] = name
    # And the skill's own reading of the file has to agree with this file's reading
    # of it. A truthful report is a separate claim from a correct document.
    for name in TABLE_PRESETS:
        mine = t["presets"][name]["facts"]
        theirs = t["presets"][name]["reported"]
        for field, got in (("borders", theirs.get("borders")),
                           ("cell_margin", theirs.get("cell_margin")),
                           ("header_rule", theirs.get("header_rule"))):
            if got != mine[field]:
                out.append(f"Q2 {name}: the report says {field}={got} while the "
                           f"document says {mine[field]} — the file and the report "
                           f"disagree, and the report is what a caller acts on")
    return out


@check("Q3", "the header row really does come back on page 2 of a table that spans pages")
def q3_header_repeat(ctx: dict) -> list[str]:
    t = ctx["tables"]
    if t["render"].get("skipped"):
        return []
    out = []
    fin = t["render"].get("finance", {})
    grid = t["render"].get("grid", {})
    if not fin.get("produced") or not grid.get("produced"):
        return ["Q3 the tall tables did not render, so nothing was measured"]
    if fin["pages"] < 2 or grid["pages"] < 2:
        return [f"Q3 the tall table fits on {fin['pages']} page(s); with nothing to "
                f"break across, 'the header came back' has no subject"]
    if fin["header_pages"] != list(range(1, fin["pages"] + 1)):
        out.append(f"Q3 finance sets w:tblHeader, but the header row rendered only on "
                   f"page(s) {fin['header_pages']} of {fin['pages']} — the flag is in "
                   f"the file and the layout engine did not act on it")
    # The control arm, and it has to stay negative: if the header repeats WITHOUT the
    # flag, then page 2 carrying it says nothing about the preset.
    if grid["header_pages"] != [1]:
        out.append(f"Q3 grid does NOT set w:tblHeader, yet its header rendered on "
                   f"page(s) {grid['header_pages']} — repeating is happening for some "
                   f"other reason and this measurement cannot be attributed to the "
                   f"preset")
    return out


@check("Q4", "column widths come from east-asian display width, not from len()")
def q4_widths(ctx: dict) -> list[str]:
    t = ctx["tables"]
    out = []
    for name in ("finance", "banded"):
        got = t["presets"][name]["facts"]["widths"]
        want = [n * WIDE_DXA_PER_CELL + 2 * WIDE_MARGIN[name] for n in WIDE_DISPLAY]
        if got != want:
            out.append(f"Q4 {name} sized the columns {got}, expected {want} "
                       f"(display width {WIDE_DISPLAY} x {WIDE_DXA_PER_CELL} dxa plus "
                       f"padding)")
        if t["presets"][name]["facts"]["layout"] != "fixed":
            out.append(f"Q4 {name} fitted the columns but left the layout auto, so "
                       f"the widths it computed are a suggestion the renderer may "
                       f"ignore")
    # grid declines to refit, and that has to stay visible — otherwise `fit_columns`
    # is not a property that distinguishes anything.
    if t["presets"]["grid"]["facts"]["layout"] == "fixed":
        out.append("Q4 grid is declared as leaving the author's widths alone, but it "
                   "wrote a fixed layout")
    # The vacuity guard: the two rulers have to DISAGREE on this fixture, or the
    # assertion above is satisfied by both implementations at once.
    if WIDE_DISPLAY == WIDE_LENS:
        out.append("Q4 display width and len() give the same answer for every column "
                   "of this fixture, so nothing here separates the two ways of "
                   "measuring text")
    return out


@check("Q5", "the fitted table renders without wrapping, where len()-sized columns wrap")
def q5_no_wrap(ctx: dict) -> list[str]:
    t = ctx["tables"]
    if t["render"].get("skipped"):
        return []
    out = []
    fitted = t["render"].get("fitted", {})
    naive = t["render"].get("naive", {})
    if not fitted.get("produced") or not naive.get("produced"):
        return ["Q5 the comparison did not render, so nothing was measured"]
    if not fitted["longest_intact"]:
        out.append(f"Q5 {WIDE_LONGEST!r} did not survive on one line in the fitted "
                   f"table — the column it sits in was allocated "
                   f"{t['presets']['finance']['facts']['widths'][0]} dxa and that was "
                   f"not enough")
    # The control arm. If len()-sized columns ALSO fit, this host's fonts have made
    # the fixture stop separating the two implementations, and Q5 proves nothing —
    # which is a finding, not a pass.
    if naive["longest_intact"]:
        out.append(f"Q5 the len()-sized columns ({t['naive_widths']}) also rendered "
                   f"{WIDE_LONGEST!r} intact, so on this host the fixture no longer "
                   f"tells a correct allocation from a naive one")
    return out


@check("Q6", "a preset is formatting: one part changes, and an empty request is refused")
def q6_scope(ctx: dict) -> list[str]:
    t = ctx["tables"]
    out = []
    for name in TABLE_PRESETS:
        changed = t["presets"][name]["parts_changed"]
        if changed != ["word/document.xml"]:
            out.append(f"Q6 {name} changed {changed}; a table preset writes direct "
                       f"formatting and has no business in any other part")
        if t["presets"][name]["exit"] != 0:
            out.append(f"Q6 {name} exited {t['presets'][name]['exit']}")
    empty = t["no_tables"]
    if empty["exit"] == 0 or empty["wrote"]:
        out.append("Q6 a document with no tables was accepted; a call that changed "
                   "nothing and reported success cannot be told from one that worked")
    if empty["traceback"]:
        out.append("Q6 the no-tables refusal arrived as a traceback rather than a "
                   "sentence")
    if t["bad_index"]["exit"] == 0 or t["bad_index"]["wrote"]:
        out.append("Q6 --table 7 was accepted on a document with one table")
    if "1" not in t["bad_index"]["stderr"]:
        out.append(f"Q6 the out-of-range refusal does not say how many tables there "
                   f"are: {t['bad_index']['stderr'][:80]!r}")
    return out


@check("Q7", "banding shades the rows it claims, and never counts the header as one")
def q7_banding(ctx: dict) -> list[str]:
    t = ctx["tables"]
    out = []
    banded = t["presets"]["banded"]["facts"]
    fills = banded["row_fills"]
    if not fills or fills[0] != "D9E2F3":
        out.append(f"Q7 the banded header row carries fill {fills[0] if fills else None}, "
                   f"expected D9E2F3")
    # Bands count from the first DATA row. If the header is band 1, every stripe
    # below it lands on the wrong row and the table is harder to read, not easier.
    body = fills[1:]
    want = [None if i % 2 == 0 else "F2F2F2" for i in range(len(body))]
    if body != want:
        out.append(f"Q7 the data rows are filled {body}, expected {want} — banding "
                   f"has to start from the first data row, not from the header")
    for name in ("grid", "finance"):
        if any(t["presets"][name]["facts"]["row_fills"]):
            out.append(f"Q7 {name} shaded a row; only banded declares a fill, and if "
                       f"they all shade then shading distinguishes nothing")
    if banded["header_rule"] != -1:
        out.append("Q7 banded wrote a header rule; that is finance's property, and "
                   "two presets sharing every property are one preset")
    if t["presets"]["finance"]["facts"]["header_rule"] != 6:
        out.append(f"Q7 finance's header rule is "
                   f"{t['presets']['finance']['facts']['header_rule']}, expected 6 — "
                   f"a three-line table is a rule under the header, not insideH")
    return out


def flaw_replace_run_by_run(ctx, work):
    """THE defect: iterate paragraph.runs and call str.replace on each.

    The replaced count is the number that implementation really produces on this
    file (measured, 0), not a number chosen to make the assertion fire.
    """
    r = copy.deepcopy(ctx["replace"])
    entry = r["report"]["replacements"][0]
    entry["replaced"] = ctx["fixture"]["naive_hits"]
    entry["cross_run"] = 0
    r["texts"] = [t.replace(CROSS_RUN_NEW, CROSS_RUN) for t in r["texts"]]
    ctx["replace"] = r
    return ctx


def flaw_replace_reports_more_than_it_did(ctx, work):
    r = copy.deepcopy(ctx["replace"])
    r["report"]["replacements"][0]["replaced"] = 5
    ctx["replace"] = r
    return ctx


def flaw_replace_leaves_the_old_text(ctx, work):
    r = copy.deepcopy(ctx["replace"])
    r["texts"] = [t.replace(CROSS_RUN_NEW, CROSS_RUN) for t in r["texts"]]
    ctx["replace"] = r
    return ctx


def flaw_replace_misses_the_placeholder(ctx, work):
    r = copy.deepcopy(ctx["replace"])
    r["texts"] = [t.replace(PLACEHOLDER_NEW, PLACEHOLDER) for t in r["texts"]]
    ctx["replace"] = r
    return ctx


def flaw_edit_rebuilds_the_package(ctx, work):
    """python-docx load→save keeps the parts but rewrites their bytes."""
    r = copy.deepcopy(ctx["replace"])
    r["identical"] = [n for n in r["identical"] if n.startswith("docProps")]
    r["report"]["parts_changed"] = sorted(set(r["parts_before"]) - set(r["identical"]))
    ctx["replace"] = r
    return ctx


def flaw_edit_drops_the_custom_xml(ctx, work):
    r = copy.deepcopy(ctx["replace"])
    r["parts_after"] = [n for n in r["parts_after"] if not n.startswith("customXml/")]
    r["identical"] = [n for n in r["identical"] if not n.startswith("customXml/")]
    ctx["replace"] = r
    return ctx


def flaw_edit_resets_the_run_font(ctx, work):
    r = copy.deepcopy(ctx["replace"])
    r["rpr_after"] = "<w:rPr/>"
    ctx["replace"] = r
    return ctx


def flaw_replace_ignores_line_breaks(ctx, work):
    n = copy.deepcopy(ctx["near_miss"])
    n["report"]["replacements"][0]["replaced"] = 1
    ctx["near_miss"] = n
    return ctx


def flaw_near_miss_reported_as_plain_zero(ctx, work):
    n = copy.deepcopy(ctx["near_miss"])
    n["report"]["replacements"][0]["near_misses"] = []
    ctx["near_miss"] = n
    return ctx


def flaw_revision_edit_says_nothing(ctx, work):
    r = copy.deepcopy(ctx["revision_edit"])
    r["report"]["replacements"][0]["contexts"] = []
    ctx["revision_edit"] = r
    return ctx


def flaw_headers_always_edited(ctx, work):
    """The flag is decoration: header parts are edited whether or not it was passed."""
    r = copy.deepcopy(ctx["replace"])
    r["header_text"] = HEADER_TEXT.replace("内部资料", "公开资料")
    ctx["replace"] = r
    return ctx


def flaw_in_headers_does_nothing(ctx, work):
    h = copy.deepcopy(ctx["headers"])
    h["header_text"] = HEADER_TEXT
    h["report"]["parts_changed"] = ["word/document.xml"]
    ctx["headers"] = h
    return ctx


def flaw_append_after_sectpr(ctx, work):
    """`body.append(p)` — the natural implementation, and an invalid document."""
    a = copy.deepcopy(ctx["append"])
    kids = [k for k in a["body_children"] if k != "sectPr"]
    a["body_children"] = kids[:-1] + ["sectPr", "p"]
    ctx["append"] = a
    return ctx


def flaw_append_drops_the_section(ctx, work):
    a = copy.deepcopy(ctx["append"])
    a["body_children"] = [k for k in a["body_children"] if k != "sectPr"]
    a["section_survives"] = False
    ctx["append"] = a
    return ctx


def flaw_list_accepted_without_numbering(ctx, work):
    a = copy.deepcopy(ctx["append"])
    a["list_without_numbering"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["append"] = a
    return ctx


def flaw_list_refused_always(ctx, work):
    a = copy.deepcopy(ctx["append"])
    a["list_with_numbering"] = {"exit": 2, "num_ids": ["1"]}
    ctx["append"] = a
    return ctx


def flaw_appended_run_has_no_eastasia(ctx, work):
    a = copy.deepcopy(ctx["append"])
    a["fonts"] = {"ascii": "Calibri", "eastAsia": None}
    ctx["append"] = a
    return ctx


def flaw_pack_loses_a_part(ctx, work):
    p = copy.deepcopy(ctx["package"])
    p["identical"] -= 1
    p["report"]["parts_lost"] = ["customXml/item1.xml"]
    ctx["package"] = p
    return ctx


def flaw_pack_reorders(ctx, work):
    p = copy.deepcopy(ctx["package"])
    p["order_after"] = sorted(p["order_before"])
    p["report"]["order_preserved"] = False
    ctx["package"] = p
    return ctx


def flaw_manifest_makes_no_difference(ctx, work):
    """The control's control: if name order equals the recorded order, P2 is blind."""
    p = copy.deepcopy(ctx["package"])
    p["without_manifest_order"] = list(p["order_before"])
    ctx["package"] = p
    return ctx


def flaw_pack_drops_added_parts(ctx, work):
    p = copy.deepcopy(ctx["package"])
    p["added_part"] = {"exit": 0, "present": False}
    ctx["package"] = p
    return ctx


def flaw_unpack_trusts_part_names(ctx, work):
    p = copy.deepcopy(ctx["package"])
    p["traversal"] = {"exit": 0, "stderr": "", "escaped_written": True}
    ctx["package"] = p
    return ctx


def flaw_check_ignores_element_order(ctx, work):
    o = copy.deepcopy(ctx["order"])
    o["check_exit"] = 0
    o["check_findings"] = []
    ctx["order"] = o
    return ctx


def flaw_check_only_finds_sectpr(ctx, work):
    o = copy.deepcopy(ctx["order"])
    o["check_findings"] = [f for f in o["check_findings"] if "sectPr" in f]
    ctx["order"] = o
    return ctx


def flaw_fix_repairs_only_the_body(ctx, work):
    o = copy.deepcopy(ctx["order"])
    o["fix_report"]["fixes"] = [{"part": "word/document.xml",
                                 "elements_reordered": 1, "detail": []}]
    o["after_findings"] = ["word/document.xml: document/body[1]/p[1] has pPr out of "
                           "the ECMA-376 order"]
    o["after_exit"] = 2
    ctx["order"] = o
    return ctx


def flaw_fix_rewrites_the_text(ctx, work):
    o = copy.deepcopy(ctx["order"])
    o["text_after"] = [t.replace("第三季度", "") for t in o["text_after"]]
    ctx["order"] = o
    return ctx


def flaw_fix_rewrites_every_part(ctx, work):
    o = copy.deepcopy(ctx["order"])
    o["other_parts_identical"] = 0
    ctx["order"] = o
    return ctx


def flaw_stdout_dumps_every_paragraph(ctx, work):
    c = copy.deepcopy(ctx["contracts"])
    c["stdout_bytes"] = 400_000
    ctx["contracts"] = c
    return ctx


def flaw_trimming_drops_the_report(ctx, work):
    c = copy.deepcopy(ctx["contracts"])
    c["stdout_bytes"] = 12
    ctx["contracts"] = c
    return ctx


def flaw_tall_table_dumped_whole(ctx, work):
    """The defect as measured before the byte budget existed: 130,602 bytes."""
    c = copy.deepcopy(ctx["contracts"])
    c["tall_stdout_bytes"] = 130_602
    ctx["contracts"] = c
    return ctx


def flaw_byte_budget_drops_the_answer(ctx, work):
    """The over-correction: trimmed so hard the report no longer names the table."""
    c = copy.deepcopy(ctx["contracts"])
    c["tall_stdout"] = "{}"
    c["tall_stdout_bytes"] = 2
    ctx["contracts"] = c
    return ctx


def flaw_in_place_write_allowed(ctx, work):
    c = copy.deepcopy(ctx["contracts"])
    c["in_place"] = {"exit": 0, "stderr": ""}
    ctx["contracts"] = c
    return ctx


def flaw_missing_file_raises(ctx, work):
    c = copy.deepcopy(ctx["contracts"])
    c["missing"] = {"exit": 1, "stderr": "Traceback (most recent call last): ...",
                    "traceback": True}
    ctx["contracts"] = c
    return ctx


def flaw_xlsx_accepted_as_docx(ctx, work):
    c = copy.deepcopy(ctx["contracts"])
    c["wrong_kind"] = {"exit": 1, "stderr": "KeyError: 'word/document.xml'",
                       "traceback": True}
    ctx["contracts"] = c
    return ctx


def flaw_fixture_not_reproducible(ctx, work):
    f = copy.deepcopy(ctx["fixtures"])
    f["identical"]["report.docx"] = False
    ctx["fixtures"] = f
    return ctx


def flaw_reader_hands_back_runs(ctx, work):
    r = copy.deepcopy(ctx["read"])
    for p in r["report"]["paragraphs"]:
        if p["text"] == SPLIT_PARAGRAPH:
            p["text"] = "2026 年第"
    ctx["read"] = r
    return ctx


def flaw_reader_folds_in_deleted_text(ctx, work):
    r = copy.deepcopy(ctx["read"])
    for p in r["report"]["paragraphs"]:
        if "同比增长。" in p["text"]:
            p["text"] = p["text"].replace("同比增长。", DELETED + "同比增长。")
    ctx["read"] = r
    return ctx


def flaw_reader_drops_revisions_entirely(ctx, work):
    r = copy.deepcopy(ctx["read"])
    r["report"].pop("revisions", None)
    ctx["read"] = r
    return ctx


def flaw_reader_skips_inserted_text(ctx, work):
    """What python-docx itself does: walk direct <w:r> children only."""
    r = copy.deepcopy(ctx["read"])
    for p in r["report"]["paragraphs"]:
        p["text"] = p["text"].replace(INSERTED, "")
    ctx["read"] = r
    return ctx


def flaw_tables_flattened(ctx, work):
    r = copy.deepcopy(ctx["read"])
    r["report"]["table_contents"] = []
    ctx["read"] = r
    return ctx


def flaw_table_rows_ragged(ctx, work):
    r = copy.deepcopy(ctx["read"])
    r["report"]["table_contents"][0]["cells"][1] = ["营业收入"]
    ctx["read"] = r
    return ctx


def flaw_list_level_not_reported(ctx, work):
    r = copy.deepcopy(ctx["read"])
    for p in r["report"]["paragraphs"]:
        p.pop("list", None)
    ctx["read"] = r
    return ctx


def flaw_section_reported_without_bindings(ctx, work):
    r = copy.deepcopy(ctx["read"])
    for s in r["report"]["sections"]:
        s["headers"] = []
    ctx["read"] = r
    return ctx


def flaw_footer_field_read_as_text(ctx, work):
    r = copy.deepcopy(ctx["read"])
    for h in r["report"]["headers_and_footers"]:
        h["fields"] = 0
    ctx["read"] = r
    return ctx


def flaw_headers_not_reported(ctx, work):
    r = copy.deepcopy(ctx["read"])
    r["report"]["headers_and_footers"] = []
    ctx["read"] = r
    return ctx


def flaw_fixture_stops_splitting_runs(ctx, work):
    """CONTROL: the fixture stops exercising the case, so V0 must say so."""
    f = copy.deepcopy(ctx["fixture"])
    f["phrase_runs"] = 1
    f["naive_hits"] = CROSS_RUN_TOTAL
    ctx["fixture"] = f
    return ctx


def flaw_fixture_loses_its_custom_xml(ctx, work):
    f = copy.deepcopy(ctx["fixture"])
    f["custom_xml"] = False
    ctx["fixture"] = f
    return ctx


def flaw_fixture_stops_being_unordered(ctx, work):
    """CONTROL: unordered.docx stops carrying the defects the repair is judged on."""
    f = copy.deepcopy(ctx["fixture"])
    f["order_defects"] = 0
    ctx["fixture"] = f
    return ctx



def flaw_fill_only_walks_runs(ctx, work):
    """The scanner walks runs, so a split placeholder is invisible to it."""
    r = copy.deepcopy(ctx["template"])
    for e in r["report"]["filled"]:
        if e["name"] == FILL_SPLIT:
            e["replaced"], e["cross_run"] = 0, 0
    ctx["template"] = r
    return ctx


def flaw_fill_leaves_the_placeholder_visible(ctx, work):
    r = copy.deepcopy(ctx["template"])
    r["texts"] = [t.replace("示例科技有限公司", "{{" + FILL_SPLIT + "}}")
                  for t in r["texts"]]
    ctx["template"] = r
    return ctx


def flaw_unfilled_not_reported(ctx, work):
    r = copy.deepcopy(ctx["template"])
    r["report"]["unfilled"] = []
    r["report"].pop("warning", None)
    ctx["template"] = r
    return ctx


def flaw_unused_value_not_reported(ctx, work):
    r = copy.deepcopy(ctx["template"])
    r["report"]["unused_values"] = []
    ctx["template"] = r
    return ctx


def flaw_strict_writes_anyway(ctx, work):
    r = copy.deepcopy(ctx["template"])
    r["strict"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["template"] = r
    return ctx


def flaw_fill_skips_headers(ctx, work):
    r = copy.deepcopy(ctx["template"])
    for e in r["report"]["filled"]:
        if e["name"] == FILL_HEADER:
            e["parts"] = {}
            e["replaced"] = 0
    r["header_text"] = HEADER_TEXT
    ctx["template"] = r
    return ctx


def flaw_render_produces_nothing(ctx, work):
    p = copy.deepcopy(ctx["pdf"])
    if p.get("skipped"):
        return ctx
    p["produced"] = False
    p["exit"] = 0
    ctx["pdf"] = p
    return ctx


def flaw_render_skips_the_images(ctx, work):
    p = copy.deepcopy(ctx["pdf"])
    if p.get("skipped"):
        return ctx
    p["images"] = 0
    ctx["pdf"] = p
    return ctx


def flaw_blank_render_handed_back(ctx, work):
    p = copy.deepcopy(ctx["pdf"])
    if p.get("skipped"):
        return ctx
    p["blank"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["pdf"] = p
    return ctx


def flaw_revisions_rendered_in_silence(ctx, work):
    p = copy.deepcopy(ctx["pdf"])
    if p.get("skipped"):
        return ctx
    p["report"].pop("warning", None)
    ctx["pdf"] = p
    return ctx


def flaw_cached_field_rendered_in_silence(ctx, work):
    p = copy.deepcopy(ctx["pdf"])
    if p.get("skipped"):
        return ctx
    p["report"].pop("fields_note", None)
    ctx["pdf"] = p
    return ctx



def flaw_track_replaces_without_tracking(ctx, work):
    """The edit is made, but silently — no revision markup at all."""
    r = copy.deepcopy(ctx["revise"])
    r["elements"] = [e for e in r["elements"] if e["author"] != REVISER]
    ctx["revise"] = r
    return ctx


def flaw_track_wraps_one_run_only(ctx, work):
    """Only the occurrence that happens to sit in a single run is tracked."""
    r = copy.deepcopy(ctx["revise"])
    for e in r["elements"]:
        if e["author"] == REVISER:
            e["runs"] = 1
    ctx["revise"] = r
    return ctx


def flaw_del_keeps_wt(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    for e in r["elements"]:
        if e["kind"] == "del":
            e["holds_wt"] = True
    ctx["revise"] = r
    return ctx


def flaw_revision_ids_reused(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    for e in r["elements"]:
        e["id"] = "101"
    ctx["revise"] = r
    return ctx


def flaw_revision_without_author(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["elements"][0] = {**r["elements"][0], "author": None}
    ctx["revise"] = r
    return ctx


def flaw_accept_does_nothing(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["accept"]["texts"] = list(r["original_texts"])
    ctx["revise"] = r
    return ctx


def flaw_reject_keeps_the_new_text(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["reject"]["texts"] = list(r["accept"]["texts"])
    ctx["revise"] = r
    return ctx


def flaw_reject_drops_the_deleted_text(ctx, work):
    """Rejecting a deletion unwraps it but forgets to turn delText back into text."""
    r = copy.deepcopy(ctx["revise"])
    t = list(r["reject"]["texts"])
    t[REVISION_PARAGRAPH] = "本季度同比增长。"
    r["reject"]["texts"] = t
    ctx["revise"] = r
    return ctx


def flaw_paragraph_mark_not_tracked(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["inserted"]["elements"] = [e for e in r["inserted"]["elements"]
                                 if not e["paragraph_mark"]]
    ctx["revise"] = r
    return ctx


def flaw_reject_leaves_an_empty_paragraph(ctx, work):
    """The defect this repo actually had before paragraph marks were ordered last."""
    r = copy.deepcopy(ctx["revise"])
    r["inserted"]["rejected"] = r["inserted"]["rejected"] + [""]
    ctx["revise"] = r
    return ctx


def flaw_author_filter_ignored(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["by_author"]["elements"] = []
    ctx["revise"] = r
    return ctx


def flaw_author_filter_does_nothing(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["by_author"]["elements"] = list(r["elements"])
    ctx["revise"] = r
    return ctx


def flaw_remaining_not_reported(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["accept"]["report"].pop("remaining", None)
    ctx["revise"] = r
    return ctx


def flaw_accept_leaves_markup_and_says_nothing(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["accept"]["elements"] = [{"kind": "ins", "id": "1", "author": "x", "date": "y",
                                "paragraph_mark": False, "runs": 1, "text": "",
                                "holds_wt": False}]
    ctx["revise"] = r
    return ctx


def flaw_strict_writes_with_leftovers(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    r["by_author"]["strict_exit"] = 0
    r["by_author"]["strict_wrote"] = True
    ctx["revise"] = r
    return ctx


def flaw_nested_survives_accept(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    t = list(r["nested"]["accepted"])
    t[TRACKED_PARAGRAPH] = t[TRACKED_PARAGRAPH] + CROSS_RUN_NEW
    r["nested"]["accepted"] = t
    ctx["revise"] = r
    return ctx


def flaw_nested_survives_reject(ctx, work):
    r = copy.deepcopy(ctx["revise"])
    t = list(r["nested"]["rejected"])
    t[TRACKED_PARAGRAPH] = t[TRACKED_PARAGRAPH] + CROSS_RUN_NEW
    r["nested"]["rejected"] = t
    ctx["revise"] = r
    return ctx


def flaw_comment_anchor_not_isolated(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["report"]["added"]["cross_run"] = False
    c["report"]["added"]["runs_wrapped"] = 1
    ctx["comment"] = c
    return ctx


def flaw_comment_without_range_markers(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["pieces"] = {**c["pieces"], "range_starts": c["before_pieces"]["range_starts"]}
    ctx["comment"] = c
    return ctx


def flaw_comment_part_not_wired(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["scratch"]["pieces"] = {**c["scratch"]["pieces"], "relationship": False}
    ctx["comment"] = c
    return ctx


def flaw_comment_creation_not_reported(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["scratch"]["report"]["added"]["package_pieces_created"] = []
    ctx["comment"] = c
    return ctx


def flaw_delete_leaves_the_part(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["deleted"]["pieces"] = {**c["deleted"]["pieces"], "part": True}
    c["deleted"]["report"]["deleted"]["comments_part_dropped"] = False
    ctx["comment"] = c
    return ctx


def flaw_delete_leaves_orphan_markers(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["deleted"]["pieces"] = {**c["deleted"]["pieces"], "range_starts": 2,
                              "references": 2}
    ctx["comment"] = c
    return ctx


def flaw_missing_anchor_accepted(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["missing_anchor"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["comment"] = c
    return ctx


def flaw_listing_hides_the_anchor_state(ctx, work):
    c = copy.deepcopy(ctx["comment"])
    c["report"]["comments"] = [{**e, "anchored": False}
                               for e in c["report"]["comments"]]
    ctx["comment"] = c
    return ctx


# ── W9 ────────────────────────────────────────────────────────────────────────
def flaw_header_part_not_wired(ctx, work):
    """The bytes and the content type, and nothing pointing at either."""
    c = copy.deepcopy(ctx["headerfooter"])
    c["create"]["rel_targets"] = [t for t in c["create"]["rel_targets"]
                                  if t != "header1.xml"]
    for ref in c["create"]["refs"]:
        if ref["kind"] == "header":
            ref["part"] = None
    ctx["headerfooter"] = c
    return ctx


def flaw_header_reference_never_added(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    c["create"]["refs"] = [r for r in c["create"]["refs"] if r["kind"] != "header"]
    ctx["headerfooter"] = c
    return ctx


def flaw_footer_never_created(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    c["create"]["parts"] = [p for p in c["create"]["parts"] if "footer" not in p]
    c["create"]["refs"] = [r for r in c["create"]["refs"] if r["kind"] != "footer"]
    ctx["headerfooter"] = c
    return ctx


def flaw_first_page_without_titlepg(ctx, work):
    """All four pieces written, and page one still shows the ordinary header."""
    c = copy.deepcopy(ctx["headerfooter"])
    c["first"]["sect_tags"] = [t for t in c["first"]["sect_tags"] if t != "titlePg"]
    for w in c["first"]["report"].get("written") or []:
        w["activated_by"] = None
    ctx["headerfooter"] = c
    return ctx


def flaw_first_page_activation_not_reported(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    for w in c["first"]["report"].get("written") or []:
        w["activated_by"] = None
    ctx["headerfooter"] = c
    return ctx


def flaw_even_switch_missing(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    c["even"]["settings_tags"] = [t for t in c["even"]["settings_tags"]
                                  if t != "evenAndOddHeaders"]
    ctx["headerfooter"] = c
    return ctx


def flaw_even_switch_put_in_the_section(ctx, work):
    """The natural guess: it looks like a section property, so it goes in sectPr."""
    c = copy.deepcopy(ctx["headerfooter"])
    c["even"]["settings_tags"] = [t for t in c["even"]["settings_tags"]
                                  if t != "evenAndOddHeaders"]
    c["even"]["sect_tags"] = c["even"]["sect_tags"] + ["evenAndOddHeaders"]
    ctx["headerfooter"] = c
    return ctx


def flaw_remove_leaves_the_part(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    c["remove"]["parts"] = sorted(c["remove"]["parts"] + ["word/header2.xml"])
    c["remove"]["overrides"] = sorted(c["remove"]["overrides"] + ["word/header2.xml"])
    for r in c["remove"]["report"].get("removed") or []:
        r["part_dropped"] = False
    ctx["headerfooter"] = c
    return ctx


def flaw_remove_leaves_titlepg_on(ctx, work):
    """Page one now has NO header, and nothing says so."""
    c = copy.deepcopy(ctx["headerfooter"])
    c["remove"]["sect_tags"] = c["remove"]["sect_tags"] + ["titlePg"]
    for r in c["remove"]["report"].get("removed") or []:
        r["title_page_switched_off"] = False
    ctx["headerfooter"] = c
    return ctx


def flaw_page_number_cached_as_a_digit(ctx, work):
    """The field is there and so is a stale '1' — which is what every page then says."""
    c = copy.deepcopy(ctx["headerfooter"])
    c["create"]["footer_xml"] = c["create"]["footer_xml"].replace(
        '<w:fldChar w:fldCharType="end"/>',
        '<w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>1</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/>')
    ctx["headerfooter"] = c
    return ctx


def flaw_page_number_written_as_plain_text(ctx, work):
    c = copy.deepcopy(ctx["headerfooter"])
    xml = c["create"]["footer_xml"]
    xml = re.sub(r"<w:fldChar[^>]*/>|<w:instrText[^>]*>.*?</w:instrText>", "", xml)
    c["create"]["footer_xml"] = xml
    for w in c["create"]["report"].get("written") or []:
        w["page_number_field"] = False
    ctx["headerfooter"] = c
    return ctx


# ── W10 ───────────────────────────────────────────────────────────────────────
def flaw_toc_flattens_the_levels(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    for h in t["report"]["headings"]:
        h["level"] = 1
    for e in t["entries"]:
        e["style"] = "TOC1"
    ctx["toc"] = t
    return ctx


def flaw_toc_lists_only_the_top_level(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    t["report"]["headings"] = [h for h in t["report"]["headings"] if h["level"] == 1]
    t["entries"] = [e for e in t["entries"] if e["style"] == "TOC1"]
    t["paragraphs"] = OUTLINE_PARAGRAPHS + 3
    ctx["toc"] = t
    return ctx


def flaw_toc_caches_invented_page_numbers(ctx, work):
    """The tempting implementation: fill the entries in and look finished."""
    t = copy.deepcopy(ctx["toc"])
    for i, e in enumerate(t["entries"]):
        e["text"] = e["link_text"] + "\t" + str(i + 1)
    ctx["toc"] = t
    return ctx


def flaw_toc_caches_no_placeholder(ctx, work):
    """The over-correction: nothing where the number goes, which reads as a bug."""
    t = copy.deepcopy(ctx["toc"])
    for e in t["entries"]:
        e["text"] = e["link_text"]
    ctx["toc"] = t
    return ctx


def flaw_no_cache_caches_anyway(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    t["no_cache"] = {"exit": 0, "paragraphs": OUTLINE_PARAGRAPHS + TOC_PARAGRAPHS_ADDED,
                     "cached_entries": len(OUTLINE_HEADINGS)}
    ctx["toc"] = t
    return ctx


def flaw_toc_field_not_dirty(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    t["document_xml"] = t["document_xml"].replace(' w:dirty="true"', "")
    ctx["toc"] = t
    return ctx


def flaw_update_fields_not_set(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    t["settings_tags"] = [x for x in t["settings_tags"] if x != "updateFields"]
    ctx["toc"] = t
    return ctx


def flaw_toc_entries_are_plain_text(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    for e in t["entries"]:
        e["anchors"] = []
    ctx["toc"] = t
    return ctx


def flaw_toc_links_to_a_bookmark_that_is_not_there(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    t["bookmarks"] = []
    ctx["toc"] = t
    return ctx


def flaw_numbering_not_attached_to_the_styles(ctx, work):
    """The half everyone writes: an abstractNum that numbers nothing."""
    t = copy.deepcopy(ctx["toc"])
    t["style_numbering"] = {k: v for k, v in t["style_numbering"].items()
                            if not k.startswith("Heading")}
    ctx["toc"] = t
    return ctx


def flaw_numbering_levels_name_no_style(ctx, work):
    t = copy.deepcopy(ctx["toc"])
    for lvl in t["abstract_levels"]:
        lvl["style"] = None
    ctx["toc"] = t
    return ctx


def flaw_contents_heading_takes_chapter_one(ctx, work):
    """Measured in a rendered PDF: '1. 目录 / 2. 经营概况 / 2.1 收入分析'."""
    t = copy.deepcopy(ctx["toc"])
    t["style_numbering"].pop(TOC_HEADING_STYLE, None)
    ctx["toc"] = t
    return ctx


# ── W11 ───────────────────────────────────────────────────────────────────────
def flaw_image_without_a_content_type(ctx, work):
    i = copy.deepcopy(ctx["image"])
    i["defaults"].pop("png", None)
    i["report"]["inserted"]["content_type_default_added"] = False
    ctx["image"] = i
    return ctx


def flaw_image_drawing_points_at_nothing(ctx, work):
    i = copy.deepcopy(ctx["image"])
    for d in i["drawings"]:
        d["part"] = None
    ctx["image"] = i
    return ctx


def flaw_size_assumes_96_dpi(ctx, work):
    """The file says 150 dpi; this is the number you get for not asking."""
    i = copy.deepcopy(ctx["image"])
    for d in i["intrinsic"]["drawings"]:
        d["extent"] = list(CHART_96DPI_EMU)
        d["inner"] = list(CHART_96DPI_EMU)
    i["intrinsic"]["report"]["inserted"]["density"] = "assumed 96 dpi"
    ctx["image"] = i
    return ctx


def flaw_extent_filled_with_the_pixel_count(ctx, work):
    i = copy.deepcopy(ctx["image"])
    for d in i["intrinsic"]["drawings"]:
        d["extent"] = list(CHART_PX)
        d["inner"] = list(CHART_PX)
    ctx["image"] = i
    return ctx


def flaw_width_cm_ignored(ctx, work):
    i = copy.deepcopy(ctx["image"])
    for d in i["drawings"]:
        d["extent"] = list(CHART_INTRINSIC_EMU)
        d["inner"] = list(CHART_INTRINSIC_EMU)
    ctx["image"] = i
    return ctx


def flaw_only_the_width_is_scaled(ctx, work):
    """Height left at its intrinsic value: the picture is stretched, nothing errors."""
    i = copy.deepcopy(ctx["image"])
    for d in i["drawings"]:
        d["extent"] = [IMAGE_WIDTH_EMU[0], CHART_INTRINSIC_EMU[1]]
        d["inner"] = list(d["extent"])
    ctx["image"] = i
    return ctx


def flaw_extent_and_inner_disagree(ctx, work):
    i = copy.deepcopy(ctx["image"])
    for d in i["drawings"]:
        d["inner"] = [d["extent"][0] // 2, d["extent"][1] // 2]
    ctx["image"] = i
    return ctx


def flaw_same_format_replace_adds_a_part(ctx, work):
    i = copy.deepcopy(ctx["image"])
    i["replace_same"]["parts"] = sorted(i["replace_same"]["parts"]
                                        + ["word/media/image2.png"])
    ctx["image"] = i
    return ctx


def flaw_replace_orphans_the_old_part(ctx, work):
    i = copy.deepcopy(ctx["image"])
    i["replace_other"]["parts"] = sorted(i["replace_other"]["parts"]
                                         + ["word/media/image1.png"])
    i["replace_other"]["rel_targets"] = sorted(i["replace_other"]["rel_targets"]
                                               + ["image1.png"])
    ctx["image"] = i
    return ctx


def flaw_replace_repoints_nothing(ctx, work):
    i = copy.deepcopy(ctx["image"])
    for d in i["replace_other"]["drawings"]:
        d["part"] = None
    ctx["image"] = i
    return ctx


# ── W16 ───────────────────────────────────────────────────────────────────────
def flaw_audit_only_looks_for_a_missing_rfonts(ctx, work):
    """Misses the run that carries @w:ascii and no @w:eastAsia — the hardest half."""
    f = copy.deepcopy(ctx["fonts"])
    f["report"]["problems"] = [p for p in f["report"]["problems"]
                               if p["missing"] != ["eastAsia"]]
    ctx["fonts"] = f
    return ctx


def flaw_repair_leaves_runs_unbound(ctx, work):
    f = copy.deepcopy(ctx["fonts"])
    f["still_unbound"] = [{"paragraph": 1, "text": "本季度整体经营情况",
                           "missing": ["eastAsia"]}]
    ctx["fonts"] = f
    return ctx


def flaw_repair_writes_the_fallback_everywhere(ctx, work):
    """Passes D6, renders, and quietly restyles a heading the author chose a face for."""
    f = copy.deepcopy(ctx["fonts"])
    f["document_xml"] = f["document_xml"].replace(FONTLESS_STYLE_EA,
                                                  FONTLESS_FALLBACK_EA)
    for p in f["report"]["problems"]:
        for slot in (p.get("written") or {}).values():
            if slot["value"] == FONTLESS_STYLE_EA:
                slot["value"] = FONTLESS_FALLBACK_EA
    ctx["fonts"] = f
    return ctx


def flaw_repair_overwrites_the_latin_face(ctx, work):
    f = copy.deepcopy(ctx["fonts"])
    f["document_xml"] = f["document_xml"].replace(FONTLESS_KEPT_ASCII, "Calibri")
    for p in f["report"]["problems"]:
        if p["missing"] == ["eastAsia"]:
            p.setdefault("written", {})["ascii"] = {"value": "Calibri",
                                                    "from": "fallback"}
    ctx["fonts"] = f
    return ctx


def flaw_check_cries_wolf(ctx, work):
    f = copy.deepcopy(ctx["fonts"])
    f["check_clean"]["report"]["unbound_runs"] = 4
    ctx["fonts"] = f
    return ctx


def flaw_check_skips_headers_and_footers(ctx, work):
    f = copy.deepcopy(ctx["fonts"])
    f["check_clean"]["report"]["parts_examined"] = ["word/document.xml"]
    ctx["fonts"] = f
    return ctx


def flaw_font_strict_writes_anyway(ctx, work):
    # NOT named flaw_strict_writes_anyway: W5's template control already owns that
    # name, and defining it twice silently rebinds it — the T4 row would then run
    # this mutation instead of its own and fire the wrong check.
    f = copy.deepcopy(ctx["fonts"])
    f["strict"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["fonts"] = f
    return ctx


def flaw_fallback_runs_not_named(ctx, work):
    f = copy.deepcopy(ctx["fonts"])
    f["report"]["fallback_used"] = []
    ctx["fonts"] = f
    return ctx


# ── the new fixtures stop exercising what they were built for ─────────────────
def flaw_outline_gains_a_header(ctx, work):
    f = copy.deepcopy(ctx["fixture"])
    f["outline_header_footer_parts"] = 1
    f["outline_sect_refs"] = 1
    ctx["fixture"] = f
    return ctx


def flaw_outline_already_declares_png(ctx, work):
    f = copy.deepcopy(ctx["fixture"])
    f["outline_png_default"] = True
    ctx["fixture"] = f
    return ctx


def flaw_fontless_becomes_bound(ctx, work):
    f = copy.deepcopy(ctx["fixture"])
    f["fontless_unbound"] = 0
    ctx["fixture"] = f
    return ctx


def flaw_fontless_style_stops_answering(ctx, work):
    """Then "keep what the document said" and "write the default" agree."""
    f = copy.deepcopy(ctx["fixture"])
    f["fontless_style_says"] = False
    ctx["fixture"] = f
    return ctx


def flaw_chart_loses_its_density(ctx, work):
    f = copy.deepcopy(ctx["fixture"])
    f["chart_has_phys"] = False
    ctx["fixture"] = f
    return ctx




def flaw_schemas_not_shipped(ctx, work):
    v = copy.deepcopy(ctx["validate"])
    v["schema_files"] = []
    ctx["validate"] = v
    return ctx


def flaw_grammar_came_from_elsewhere(ctx, work):
    """Green on the author's machine because $ECMA376_XSD_DIR happened to be set."""
    v = copy.deepcopy(ctx["validate"])
    v["good"]["report"]["schemas"]["resolved_from"] = "$ECMA376_XSD_DIR"
    ctx["validate"] = v
    return ctx


def flaw_invalid_document_accepted(ctx, work):
    v = copy.deepcopy(ctx["validate"])
    v["invalid"] = {"exit": 0, "report": {"valid": True, "violations": [],
                                          "violation_count": 0}}
    ctx["validate"] = v
    return ctx


def flaw_violation_without_a_location(ctx, work):
    v = copy.deepcopy(ctx["validate"])
    for viol in v["invalid"]["report"]["violations"]:
        viol["part"], viol["line"] = None, None
    ctx["validate"] = v
    return ctx


def flaw_missing_schemas_pass_quietly(ctx, work):
    """The one outcome that was ruled out before any code was written."""
    v = copy.deepcopy(ctx["validate"])
    v["missing"] = {"exit": 0, "stderr": "",
                    "stdout": '{"valid": true, "violation_count": 0}'}
    ctx["validate"] = v
    return ctx


def flaw_ignorable_namespace_reported_as_invalid(ctx, work):
    v = copy.deepcopy(ctx["validate"])
    v["ignorable"] = {"exit": 1, "report": {"valid": False, "violation_count": 12}}
    ctx["validate"] = v
    return ctx


def flaw_not_checked_never_reported(ctx, work):
    v = copy.deepcopy(ctx["validate"])
    v["good"]["report"].pop("not_checked", None)
    ctx["validate"] = v
    return ctx


def flaw_a_block_kind_is_dropped(ctx, work):
    """The defect the whole contract is about: content that never arrives."""
    m = copy.deepcopy(ctx["markdown"])
    m["report"]["written"]["blocks"].pop("table", None)
    ctx["markdown"] = m
    return ctx


def flaw_sectpr_no_longer_last(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    m["body_children"] = m["body_children"][:-1] + ["sectPr", "p"]
    ctx["markdown"] = m
    return ctx


def flaw_unsupported_silently_dropped(ctx, work):
    """A legal document that is missing content, and nothing says so."""
    m = copy.deepcopy(ctx["markdown"])
    m["report"]["unsupported"] = []
    ctx["markdown"] = m
    return ctx


def flaw_unsupported_reported_without_a_line(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    for u in m["report"]["unsupported"]:
        u["line"] = 0
    ctx["markdown"] = m
    return ctx


def flaw_md_strict_writes_anyway(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    m["strict"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["markdown"] = m
    return ctx


def flaw_style_named_but_never_created(ctx, work):
    """w:pStyle naming a style that does not exist: valid XML, silently Normal."""
    m = copy.deepcopy(ctx["markdown"])
    m["styles"] = [s for s in m["styles"] if s != "SourceCode"]
    ctx["markdown"] = m
    return ctx


def flaw_nested_list_flattened(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    m["levels_used"] = ["0"]
    ctx["markdown"] = m
    return ctx


def flaw_only_one_list_defined(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    m["numbering"] = {"1": ["0", "1", "2"]}
    ctx["markdown"] = m
    return ctx


def flaw_template_parts_lost(ctx, work):
    m = copy.deepcopy(ctx["markdown"])
    m["template"]["lost"] = ["word/numbering.xml", "word/header1.xml"]
    m["template"]["parts"] = [p for p in m["template"]["parts"]
                              if "header" not in p]
    ctx["markdown"] = m
    return ctx


def flaw_template_body_kept_as_well(ctx, work):
    """Appending to the template instead of replacing its body: the generated
    document then carries somebody else's text."""
    m = copy.deepcopy(ctx["markdown"])
    m["template"]["report"]["template_blocks_replaced"] = 0
    m["template"]["texts"] = ["二零二六年第三季度经营分析报告"] + m["template"]["texts"]
    ctx["markdown"] = m
    return ctx


def flaw_style_written_out_of_ct_style_order(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["create"]["styles"][NEW_STYLE]["children"] = ["rPr", "name", "basedOn"]
    ctx["styles"] = s
    return ctx


def flaw_style_edit_rewrites_other_parts(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["create"]["parts_changed"] = ["word/document.xml", "word/styles.xml"]
    ctx["styles"] = s
    return ctx


def flaw_overwrite_allowed_silently(ctx, work):
    """The defect: a style change that looks local and repaints the document."""
    s = copy.deepcopy(ctx["styles"])
    s["no_overwrite"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["styles"] = s
    return ctx


def flaw_overwrite_does_not_say_how_many(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["overwrite"]["report"]["set"]["repainted"] = 0
    s["no_overwrite"]["stderr"] = "error: style already exists"
    ctx["styles"] = s
    return ctx


def flaw_report_says_it_set_nothing(ctx, work):
    """Measured: `properties_set: []` on a call that set six of them."""
    s = copy.deepcopy(ctx["styles"])
    s["overwrite"]["report"]["set"]["properties_set"] = []
    ctx["styles"] = s
    return ctx


def flaw_size_written_in_points_not_half_points(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["overwrite"]["styles"]["Heading1"]["sz"] = "20"
    ctx["styles"] = s
    return ctx


def flaw_delete_in_use_allowed(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["delete_in_use"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["styles"] = s
    return ctx


def flaw_reassign_leaves_paragraphs_pointing_at_nothing(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["reassign"]["paragraph_styles"] = ["Heading1"] + \
        s["reassign"]["paragraph_styles"][1:]
    ctx["styles"] = s
    return ctx


def flaw_children_left_based_on_a_deleted_style(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    first = next(iter(s["reassign"]["styles"]))
    s["reassign"]["styles"][first]["based_on"] = "Heading1"
    ctx["styles"] = s
    return ctx


def flaw_basedon_cycle_written(ctx, work):
    s = copy.deepcopy(ctx["styles"])
    s["cycle"] = {"exit": 0, "stderr": "", "wrote": True}
    ctx["styles"] = s
    return ctx


def flaw_entry_point_dies_on_a_windows_code_page(ctx, work):
    """The measured defect: on Windows every entry point exited 2 with
    UnicodeEncodeError the moment its output was captured."""
    e = copy.deepcopy(ctx["encoding"])
    for key in list(e["runs"]):
        if key.startswith("cp1252:"):
            e["runs"][key] = {"exit": 2, "cjk_in_stdout": False,
                              "stderr": "UnicodeEncodeError: 'charmap' codec can't "
                                        "encode character '\\u2705'"}
    ctx["encoding"] = e
    return ctx


def flaw_a_new_entry_point_is_never_probed(ctx, work):
    e = copy.deepcopy(ctx["encoding"])
    e["entry_points"] = sorted(e["entry_points"] + ["docx_brandnew.py"])
    ctx["encoding"] = e
    return ctx


def flaw_probes_print_nothing_chinese(ctx, work):
    """Passing without having crossed the code path is the other way to be wrong."""
    e = copy.deepcopy(ctx["encoding"])
    for r in e["runs"].values():
        r["cjk_in_stdout"] = False
    ctx["encoding"] = e
    return ctx


def flaw_host_does_not_reproduce_the_code_page(ctx, work):
    """CONTROL: PYTHONIOENCODING stops biting, so C4 is measuring nothing."""
    e = copy.deepcopy(ctx["encoding"])
    e["bare_print_exit"] = 0
    ctx["encoding"] = e
    return ctx


# ── W18 ───────────────────────────────────────────────────────────────────────
def _diff(ctx):
    ctx["diff"] = copy.deepcopy(ctx["diff"])
    return ctx["diff"]


def flaw_revised_stops_removing_a_paragraph(ctx, work):
    """CONTROL: the A/B fixture stops carrying one of the six changes."""
    d = _diff(ctx)
    d["fixture"]["removed_from_a"] = False
    return ctx


def flaw_revised_stops_restyling(ctx, work):
    """CONTROL: the style-only change goes, and U1's style finding has no subject."""
    d = _diff(ctx)
    d["fixture"]["restyled"] = []
    return ctx


def flaw_revised_header_matches(ctx, work):
    """CONTROL: the letterhead stops differing, so 'the diff reached it' is vacuous."""
    d = _diff(ctx)
    d["fixture"]["header_b"] = list(d["fixture"]["header_a"])
    return ctx


def flaw_diff_only_walks_the_body(ctx, work):
    """The implementation everybody writes first: iterate word/document.xml and stop."""
    d = _diff(ctx)
    d["main"]["report"]["differences"] = [
        f for f in d["main"]["report"]["differences"]
        if DIFF_HEADER_PART not in f.get("where", "")]
    d["main"]["report"]["counted"] = len(d["main"]["report"]["differences"])
    return ctx


def flaw_diff_flattens_tables(ctx, work):
    d = _diff(ctx)
    d["main"]["report"]["differences"] = [
        f for f in d["main"]["report"]["differences"]
        if "table" not in f.get("where", "")]
    d["main"]["report"]["counted"] = len(d["main"]["report"]["differences"])
    return ctx


def flaw_style_change_not_noticed(ctx, work):
    d = _diff(ctx)
    d["main"]["report"]["differences"] = [
        f for f in d["main"]["report"]["differences"] if f["kind"] != "style"]
    d["main"]["report"]["counted"] = len(d["main"]["report"]["differences"])
    return ctx


def flaw_accept_does_not_give_b(ctx, work):
    d = _diff(ctx)
    d["accept"]["signature"] = d["signatures"]["a"]
    return ctx


def flaw_reject_does_not_give_a(ctx, work):
    d = _diff(ctx)
    d["reject"]["signature"] = d["signatures"]["b"]
    return ctx


def flaw_script_will_not_claim_the_roundtrip(ctx, work):
    """The round trip holds and the script does not say so — silence reads as failure."""
    d = _diff(ctx)
    d["main"]["report"]["roundtrip"]["exact"] = False
    return ctx


def flaw_noise_counted_as_a_difference(ctx, work):
    """The diff nobody reads: an rsid churn reported as an edit."""
    d = _diff(ctx)
    d["noise"]["report"]["differences"] = [
        {"kind": "text", "where": "word/document.xml ¶1", "before": "x", "after": "x"}]
    d["noise"]["report"]["counted"] = 1
    return ctx


def flaw_ignored_category_never_named(ctx, work):
    d = _diff(ctx)
    d["main"]["report"]["ignored_not_counted_as_differences"].pop("proofErr", None)
    return ctx


def flaw_category_named_but_never_counted(ctx, work):
    """The shape of "absolutely never silent" that is, in fact, silent."""
    d = _diff(ctx)
    for key in ("rsid", "proofErr", "bookmark", "lang"):
        d["noise"]["report"]["ignored_not_counted_as_differences"][key] = {
            "a": 0, "b": 0}
    return ctx


def flaw_unmarkable_change_not_named(ctx, work):
    d = _diff(ctx)
    d["rows"]["report"]["not_redlined"] = []
    return ctx


def flaw_exact_claimed_over_a_gap(ctx, work):
    d = _diff(ctx)
    d["rows"]["report"].setdefault("roundtrip", {})["exact"] = True
    return ctx


def flaw_strict_writes_a_partial_redline(ctx, work):
    d = _diff(ctx)
    d["rows_strict"] = {"exit": 0, "stderr": "", "wrote": True}
    return ctx


def flaw_row_change_not_reported_at_all(ctx, work):
    d = _diff(ctx)
    d["rows"]["report"]["differences"] = [
        f for f in d["rows"]["report"]["differences"] if f["kind"] != "table-shape"]
    return ctx


def flaw_move_reported_as_delete_and_insert(ctx, work):
    d = _diff(ctx)
    for f in d["moved"]["report"]["differences"]:
        if f["kind"] == "paragraph-moved":
            f["kind"] = "paragraph-removed" if f.get("before") else "paragraph-added"
    return ctx


def flaw_redline_does_not_validate(ctx, work):
    d = _diff(ctx)
    d["redline_valid"]["report"] = {"valid": False,
                                    "violations": [{"part": DOC_PART, "line": 2}]}
    return ctx


def flaw_redline_has_no_paragraph_mark(ctx, work):
    """Content marked, break not: accepting empties the paragraph and leaves it."""
    d = _diff(ctx)
    d["redline_revisions"]["revisions"] = [
        r for r in d["redline_revisions"].get("revisions", [])
        if r.get("scope") != "paragraph-mark"]
    return ctx


def flaw_redline_skips_the_style_change(ctx, work):
    d = _diff(ctx)
    d["redline_revisions"]["revisions"] = [
        r for r in d["redline_revisions"].get("revisions", [])
        if r.get("kind") != "pPrChange"]
    return ctx


def flaw_redline_rewrites_the_styles_part(ctx, work):
    d = _diff(ctx)
    d["redline_parts_changed"] = sorted(d["redline_parts_changed"] + ["word/styles.xml"])
    return ctx


def flaw_picture_change_invisible(ctx, work):
    """A swapped chart, in a diff that reads only text."""
    d = _diff(ctx)
    d["image"]["report"]["differences"] = [
        f for f in d["image"]["report"]["differences"]
        if not f["kind"].startswith("image-")]
    d["image"]["report"]["not_redlined"] = [
        n for n in d["image"]["report"]["not_redlined"] if "media" not in n]
    return ctx


def flaw_picture_change_looks_reviewed(ctx, work):
    d = _diff(ctx)
    d["image"]["report"]["not_redlined"] = []
    d["image"]["report"].setdefault("roundtrip", {})["exact"] = True
    return ctx


# ── W19 ───────────────────────────────────────────────────────────────────────
def _tables(ctx):
    ctx["tables"] = copy.deepcopy(ctx["tables"])
    return ctx["tables"]


def _sync(t, preset: str, field: str, value):
    """Apply a defect to BOTH readings of the file — the document's and the report's.

    A defect in the implementation shows up in the document AND in the report the
    implementation writes about it: the two agree, and are both wrong. Mutating only
    this file's reading manufactures a disagreement the modelled defect would never
    produce, and Q2 — which owns "the report and the file disagree" — then goes red
    for a reason that has nothing to do with the flaw under test. Both of the first
    two W19 controls were written that way and both were caught by exactly that.
    """
    t["presets"][preset]["facts"][field] = value
    t["presets"][preset]["reported"][field] = copy.deepcopy(value)


def flaw_border_size_written_in_points(ctx, work):
    """sz is EIGHTHS of a point; writing points gives a hairline and no complaint."""
    t = _tables(ctx)
    for name in TABLE_PRESETS:
        b = dict(t["presets"][name]["facts"]["borders"])
        for edge, v in list(b.items()):
            if v > 0:
                b[edge] = max(1, v // 8)
        _sync(t, name, "borders", b)
    return ctx


def flaw_border_left_unstated(ctx, work):
    """Omitting an edge instead of writing `none` lets the table style show through."""
    t = _tables(ctx)
    b = dict(t["presets"]["finance"]["facts"]["borders"])
    for edge in ("left", "right", "insideV"):
        b[edge] = -1
    _sync(t, "finance", "borders", b)
    return ctx


def flaw_list_presets_describes_something_else(ctx, work):
    t = _tables(ctx)
    for p in t["listed"]["presets"]:
        if p["name"] == "finance":
            p["border_eighths"] = dict(PRESET_BORDERS["grid"])
    return ctx


def flaw_two_presets_are_the_same(ctx, work):
    """A name with nothing behind it — the defect thirteen presets are made of."""
    t = _tables(ctx)
    t["presets"]["banded"]["facts"] = copy.deepcopy(t["presets"]["finance"]["facts"])
    t["presets"]["banded"]["reported"] = copy.deepcopy(
        t["presets"]["finance"]["reported"])
    return ctx


def flaw_report_disagrees_with_the_document(ctx, work):
    t = _tables(ctx)
    t["presets"]["finance"]["reported"]["cell_margin"] = 108
    return ctx


def flaw_header_repeat_not_rendered(ctx, work):
    """The flag is in the file; the header still does not come back."""
    t = _tables(ctx)
    t["render"]["finance"]["header_pages"] = [1]
    return ctx


def flaw_every_preset_repeats_the_header(ctx, work):
    """CONTROL: if it repeats without the flag, page 2 says nothing about the preset."""
    t = _tables(ctx)
    grid = t["render"]["grid"]
    grid["header_pages"] = list(range(1, grid["pages"] + 1))
    return ctx


def flaw_columns_sized_by_len(ctx, work):
    t = _tables(ctx)
    for name in ("finance", "banded"):
        t["presets"][name]["facts"]["widths"] = [
            n * WIDE_DXA_PER_CELL + 2 * WIDE_MARGIN[name] for n in WIDE_LENS]
    return ctx


def flaw_fitted_widths_left_auto(ctx, work):
    """Widths computed and then written as a suggestion the renderer may ignore."""
    t = _tables(ctx)
    t["presets"]["finance"]["facts"]["layout"] = None
    return ctx


def flaw_fitted_table_wraps_anyway(ctx, work):
    t = _tables(ctx)
    t["render"]["fitted"]["longest_intact"] = False
    return ctx


def flaw_naive_widths_fit_too(ctx, work):
    """CONTROL: the fixture stops separating a correct allocation from a naive one."""
    t = _tables(ctx)
    t["render"]["naive"]["longest_intact"] = True
    return ctx


def flaw_preset_rewrites_the_styles_part(ctx, work):
    t = _tables(ctx)
    t["presets"]["banded"]["parts_changed"] = ["word/document.xml", "word/styles.xml"]
    return ctx


def flaw_no_tables_reported_as_done(ctx, work):
    t = _tables(ctx)
    t["no_tables"] = {"exit": 0, "stderr": "", "wrote": True, "traceback": False}
    return ctx


def flaw_out_of_range_table_accepted(ctx, work):
    t = _tables(ctx)
    t["bad_index"] = {"exit": 0, "stderr": "", "wrote": True}
    return ctx


def flaw_banding_starts_at_the_header(ctx, work):
    """Header counted as band 1, so every stripe below it lands one row out."""
    t = _tables(ctx)
    fills = t["presets"]["banded"]["facts"]["row_fills"]
    t["presets"]["banded"]["facts"]["row_fills"] = [
        "D9E2F3"] + ["F2F2F2" if i % 2 == 0 else None for i in range(len(fills) - 1)]
    return ctx


def flaw_every_preset_shades(ctx, work):
    t = _tables(ctx)
    facts = t["presets"]["finance"]["facts"]
    facts["row_fills"] = ["D9E2F3"] + [None] * (len(facts["row_fills"]) - 1)
    return ctx


def flaw_header_rule_drawn_as_insideh(ctx, work):
    """A rule between every pair of rows is a grid, not a three-line table."""
    t = _tables(ctx)
    _sync(t, "finance", "header_rule", -1)
    b = dict(t["presets"]["finance"]["facts"]["borders"])
    b["insideH"] = 6
    _sync(t, "finance", "borders", b)
    return ctx


def flaw_probe_string_also_in_the_document(ctx, work):
    """CONTROL: the string Q5 hunts for is findable somewhere that is not the table."""
    t = _tables(ctx)
    t["probe_strings_unique"][WIDE_LONGEST] = 2
    return ctx


FLAWS = [
    ("replace-run-by-run", flaw_replace_run_by_run, {"E1", "E2"},
     "E2 also fires, and it must: a replacement that never happened leaves the old "
     "phrase in the text. E1 owns the reason (the phrase spans runs), E2 owns the "
     "consequence (the document still says it)"),
    ("replace-reports-more-than-it-did", flaw_replace_reports_more_than_it_did,
     {"E1"}, ""),
    ("replace-leaves-the-old-text", flaw_replace_leaves_the_old_text, {"E2"}, ""),
    ("replace-misses-the-placeholder", flaw_replace_misses_the_placeholder, {"E2"}, ""),
    ("edit-rebuilds-every-part", flaw_edit_rebuilds_the_package, {"E3"}, ""),
    ("edit-drops-the-custom-xml", flaw_edit_drops_the_custom_xml, {"E3"}, ""),
    ("edit-resets-the-run-font", flaw_edit_resets_the_run_font, {"E3"}, ""),
    ("replace-steps-over-a-line-break", flaw_replace_ignores_line_breaks, {"E4"}, ""),
    ("near-miss-reported-as-plain-zero", flaw_near_miss_reported_as_plain_zero,
     {"E4"}, ""),
    ("edit-inside-a-tracked-change-says-nothing", flaw_revision_edit_says_nothing,
     {"E5"}, ""),
    ("headers-edited-whether-asked-or-not", flaw_headers_always_edited, {"E6"}, ""),
    ("in-headers-does-nothing", flaw_in_headers_does_nothing, {"E6"}, ""),

    ("append-after-sectpr", flaw_append_after_sectpr, {"A1"}, ""),
    ("append-drops-the-section", flaw_append_drops_the_section, {"A1"}, ""),
    ("list-accepted-with-no-list-to-join", flaw_list_accepted_without_numbering,
     {"A2"}, ""),
    ("list-guard-refuses-everything", flaw_list_refused_always, {"A2"}, ""),
    ("written-run-has-no-eastasia", flaw_appended_run_has_no_eastasia, {"A3"}, ""),

    ("pack-loses-a-part", flaw_pack_loses_a_part, {"P1"}, ""),
    ("pack-reorders-the-parts", flaw_pack_reorders, {"P1", "P2"},
     "P1 also fires: the report's own order_preserved flag is what P1 reads, and a "
     "pack that reorders sets it. P2 owns the order, P1 owns the report agreeing "
     "with the file"),
    ("manifest-makes-no-difference", flaw_manifest_makes_no_difference, {"P2"}, ""),
    ("pack-drops-hand-added-parts", flaw_pack_drops_added_parts, {"P3"}, ""),
    ("unpack-trusts-part-names", flaw_unpack_trusts_part_names, {"P4"}, ""),

    ("check-ignores-element-order", flaw_check_ignores_element_order, {"O1"}, ""),
    ("check-only-looks-at-the-body", flaw_check_only_finds_sectpr, {"O1"}, ""),
    ("fix-repairs-only-one-of-them", flaw_fix_repairs_only_the_body, {"O2"}, ""),
    ("fix-rewrites-the-text", flaw_fix_rewrites_the_text, {"O3"}, ""),
    ("fix-rewrites-every-part", flaw_fix_rewrites_every_part, {"O3"}, ""),

    ("read-hands-back-runs-not-paragraphs", flaw_reader_hands_back_runs, {"R1"}, ""),
    ("read-folds-in-deleted-text", flaw_reader_folds_in_deleted_text, {"R2"}, ""),
    ("read-drops-revisions-entirely", flaw_reader_drops_revisions_entirely, {"R2"}, ""),
    ("read-skips-inserted-text", flaw_reader_skips_inserted_text, {"R3"}, ""),
    ("read-flattens-tables", flaw_tables_flattened, {"R4"}, ""),
    ("read-returns-ragged-rows", flaw_table_rows_ragged, {"R4"}, ""),
    ("read-forgets-list-levels", flaw_list_level_not_reported, {"R5"}, ""),
    ("read-forgets-the-header-binding", flaw_section_reported_without_bindings,
     {"R5"}, ""),
    ("read-reports-a-page-field-as-text", flaw_footer_field_read_as_text, {"R6"}, ""),
    ("read-ignores-headers-and-footers", flaw_headers_not_reported, {"R6"}, ""),

    ("stdout-dumps-every-paragraph", flaw_stdout_dumps_every_paragraph, {"C1"}, ""),
    ("trimming-drops-the-report", flaw_trimming_drops_the_report, {"C1"}, ""),
    ("edit-made-without-tracking-it", flaw_track_replaces_without_tracking, {"K1"}, ""),
    ("tracking-wraps-a-single-run-only", flaw_track_wraps_one_run_only, {"K1"}, ""),
    ("deleted-text-left-as-w-t", flaw_del_keeps_wt, {"K1"}, ""),
    ("revision-ids-reused", flaw_revision_ids_reused, {"K1"}, ""),
    ("revision-without-an-author", flaw_revision_without_author, {"K1"}, ""),
    ("accept-changes-nothing", flaw_accept_does_nothing, {"K2"}, ""),
    ("reject-keeps-the-new-text", flaw_reject_keeps_the_new_text, {"K2"}, ""),
    ("reject-drops-the-deleted-text", flaw_reject_drops_the_deleted_text, {"K2"}, ""),
    ("inserted-paragraph-mark-not-tracked", flaw_paragraph_mark_not_tracked,
     {"K3"}, ""),
    ("reject-leaves-an-empty-paragraph", flaw_reject_leaves_an_empty_paragraph,
     {"K3"}, ""),
    ("author-filter-resolves-everything", flaw_author_filter_ignored, {"K4"}, ""),
    ("author-filter-resolves-nothing", flaw_author_filter_does_nothing, {"K4"}, ""),
    ("remaining-not-reported", flaw_remaining_not_reported, {"K5"}, ""),
    ("accept-leaves-markup-and-claims-success",
     flaw_accept_leaves_markup_and_says_nothing, {"K5"}, ""),
    ("strict-writes-with-leftovers", flaw_strict_writes_with_leftovers, {"K5"}, ""),
    ("insert-then-delete-survives-accept", flaw_nested_survives_accept, {"K6"}, ""),
    ("insert-then-delete-survives-reject", flaw_nested_survives_reject, {"K6"}, ""),

    ("comment-anchor-not-isolated", flaw_comment_anchor_not_isolated, {"B1"}, ""),
    ("comment-without-range-markers", flaw_comment_without_range_markers, {"B1"}, ""),
    ("comments-part-not-wired-up", flaw_comment_part_not_wired, {"B2"}, ""),
    ("comment-part-creation-not-reported", flaw_comment_creation_not_reported,
     {"B2"}, ""),
    ("delete-leaves-the-comments-part", flaw_delete_leaves_the_part, {"B3"}, ""),
    ("delete-leaves-orphan-range-markers", flaw_delete_leaves_orphan_markers,
     {"B3"}, ""),
    ("comment-anchored-to-text-that-is-not-there", flaw_missing_anchor_accepted,
     {"B4"}, ""),
    ("listing-hides-whether-a-comment-is-anchored",
     flaw_listing_hides_the_anchor_state, {"B5"}, ""),

    ("one-huge-table-dumped-to-stdout", flaw_tall_table_dumped_whole, {"C3"}, ""),
    ("byte-budget-drops-the-answer-too", flaw_byte_budget_drops_the_answer, {"C3"}, ""),
    ("writer-overwrites-its-own-input", flaw_in_place_write_allowed, {"C2"}, ""),
    ("missing-file-answers-with-a-traceback", flaw_missing_file_raises, {"C2"}, ""),
    ("an-xlsx-is-accepted-as-a-docx", flaw_xlsx_accepted_as_docx, {"C2"}, ""),
    ("fixture-is-not-reproducible", flaw_fixture_not_reproducible, {"F0"}, ""),

    ("fill-scans-runs-not-paragraphs", flaw_fill_only_walks_runs, {"T1"}, ""),
    ("fill-reports-a-hit-it-did-not-make", flaw_fill_leaves_the_placeholder_visible,
     {"T1"}, ""),
    ("unfilled-placeholders-not-reported", flaw_unfilled_not_reported, {"T2"}, ""),
    ("unused-value-not-reported", flaw_unused_value_not_reported, {"T3"}, ""),
    ("strict-writes-the-file-anyway", flaw_strict_writes_anyway, {"T4"}, ""),
    ("fill-never-reaches-the-header", flaw_fill_skips_headers, {"T5"}, ""),

    ("render-produces-no-file", flaw_render_produces_nothing, {"Y1"}, ""),
    ("render-skips-the-page-images", flaw_render_skips_the_images, {"Y1"}, ""),
    ("blank-render-handed-back-as-a-preview", flaw_blank_render_handed_back,
     {"Y2"}, ""),
    ("tracked-changes-rendered-in-silence", flaw_revisions_rendered_in_silence,
     {"Y3"}, ""),
    ("cached-field-rendered-in-silence", flaw_cached_field_rendered_in_silence,
     {"Y3"}, ""),

    ("CONTROL: fixture stops splitting the phrase across runs",
     flaw_fixture_stops_splitting_runs, {"V0", "E1"},
     "E1 also fires, and that is the point: once the per-run search finds as much as "
     "this implementation does, E1's comparison is meaningless. V0 says the fixture "
     "stopped exercising it; E1 says the assertion stopped distinguishing"),
    ("CONTROL: fixture loses its custom XML part", flaw_fixture_loses_its_custom_xml,
     {"V0"}, ""),
    ("CONTROL: unordered.docx stops being unordered", flaw_fixture_stops_being_unordered,
     {"V0"}, ""),

    ("header-part-written-but-not-wired", flaw_header_part_not_wired, {"H1"}, ""),
    ("header-reference-never-added", flaw_header_reference_never_added, {"H1"}, ""),
    ("footer-never-created", flaw_footer_never_created, {"H1"}, ""),
    ("first-page-header-without-titlepg", flaw_first_page_without_titlepg, {"H2"}, ""),
    ("first-page-activation-not-reported", flaw_first_page_activation_not_reported,
     {"H2"}, ""),
    ("even-header-without-the-settings-switch", flaw_even_switch_missing, {"H3"}, ""),
    ("even-switch-written-into-the-section", flaw_even_switch_put_in_the_section,
     {"H3"}, ""),
    ("remove-leaves-the-orphan-part", flaw_remove_leaves_the_part, {"H4"}, ""),
    ("remove-leaves-titlepg-on", flaw_remove_leaves_titlepg_on, {"H4"}, ""),
    ("page-number-cached-as-a-digit", flaw_page_number_cached_as_a_digit, {"H5"}, ""),
    ("page-number-written-as-plain-text", flaw_page_number_written_as_plain_text,
     {"H5"}, ""),

    ("toc-flattens-the-levels", flaw_toc_flattens_the_levels, {"G1"}, ""),
    ("toc-lists-only-the-top-level", flaw_toc_lists_only_the_top_level, {"G1"}, ""),
    ("toc-caches-invented-page-numbers", flaw_toc_caches_invented_page_numbers,
     {"G2"}, ""),
    ("toc-caches-no-placeholder-either", flaw_toc_caches_no_placeholder, {"G2"}, ""),
    ("no-cache-caches-anyway", flaw_no_cache_caches_anyway, {"G2"}, ""),
    ("toc-field-not-marked-dirty", flaw_toc_field_not_dirty, {"G3"}, ""),
    ("update-fields-not-set", flaw_update_fields_not_set, {"G3"}, ""),
    ("toc-entries-are-plain-text", flaw_toc_entries_are_plain_text, {"G4"}, ""),
    ("toc-links-to-a-bookmark-that-is-not-there",
     flaw_toc_links_to_a_bookmark_that_is_not_there, {"G4"}, ""),
    # G6 stays SILENT here, and that is right rather than a gap: with no heading
    # style numbered there is nothing for the contents heading to inherit, so G6's
    # subject has gone. G5 owns the missing half; G6 owns what happens when the half
    # IS there. The first draft of this row predicted a cascade and got it backwards.
    ("numbering-not-attached-to-the-styles", flaw_numbering_not_attached_to_the_styles,
     {"G5"}, ""),
    ("numbering-levels-name-no-style", flaw_numbering_levels_name_no_style, {"G5"}, ""),
    ("contents-heading-takes-chapter-one", flaw_contents_heading_takes_chapter_one,
     {"G6"}, ""),

    ("image-written-with-no-content-type", flaw_image_without_a_content_type,
     {"M1"}, ""),
    ("image-drawing-points-at-nothing", flaw_image_drawing_points_at_nothing,
     {"M1"}, ""),
    ("size-assumes-96-dpi", flaw_size_assumes_96_dpi, {"M2"}, ""),
    ("extent-filled-with-the-pixel-count", flaw_extent_filled_with_the_pixel_count,
     {"M2"}, ""),
    ("width-cm-ignored", flaw_width_cm_ignored, {"M3"}, ""),
    ("only-the-width-is-scaled", flaw_only_the_width_is_scaled, {"M3"}, ""),
    ("extent-and-inner-disagree", flaw_extent_and_inner_disagree, {"M4"}, ""),
    ("same-format-replace-adds-a-part", flaw_same_format_replace_adds_a_part,
     {"M5"}, ""),
    ("replace-orphans-the-old-part", flaw_replace_orphans_the_old_part, {"M5"}, ""),
    ("replace-repoints-nothing", flaw_replace_repoints_nothing, {"M5"}, ""),

    ("audit-only-looks-for-a-missing-rfonts",
     flaw_audit_only_looks_for_a_missing_rfonts, {"N1"}, ""),
    ("repair-leaves-runs-unbound", flaw_repair_leaves_runs_unbound, {"N1"}, ""),
    ("repair-writes-the-fallback-everywhere",
     flaw_repair_writes_the_fallback_everywhere, {"N2"}, ""),
    ("repair-overwrites-the-latin-face", flaw_repair_overwrites_the_latin_face,
     {"N3"}, ""),
    ("check-cries-wolf", flaw_check_cries_wolf, {"N4"}, ""),
    ("check-skips-headers-and-footers", flaw_check_skips_headers_and_footers,
     {"N4"}, ""),
    ("font-strict-writes-anyway", flaw_font_strict_writes_anyway, {"N5"}, ""),
    ("fallback-runs-not-named", flaw_fallback_runs_not_named, {"N5"}, ""),

    ("CONTROL: outline.docx gains a header", flaw_outline_gains_a_header, {"V0"}, ""),
    ("CONTROL: outline.docx already declares a png content type",
     flaw_outline_already_declares_png, {"V0"}, ""),
    ("CONTROL: fontless.docx becomes correctly bound", flaw_fontless_becomes_bound,
     {"V0"}, ""),
    ("CONTROL: fontless.docx's style stops naming a face",
     flaw_fontless_style_stops_answering, {"V0"}, ""),
    ("CONTROL: chart.png loses its declared density", flaw_chart_loses_its_density,
     {"V0"}, ""),

    ("schemas-not-shipped", flaw_schemas_not_shipped, {"Z1"}, ""),
    ("grammar-came-from-somewhere-else", flaw_grammar_came_from_elsewhere,
     {"Z1"}, ""),
    ("invalid-document-accepted", flaw_invalid_document_accepted, {"Z2"}, ""),
    ("violation-without-a-location", flaw_violation_without_a_location, {"Z2"}, ""),
    ("missing-schemas-pass-quietly", flaw_missing_schemas_pass_quietly, {"Z3"}, ""),
    ("ignorable-namespace-reported-as-invalid",
     flaw_ignorable_namespace_reported_as_invalid, {"Z4"}, ""),
    ("not-checked-never-reported", flaw_not_checked_never_reported, {"Z5"}, ""),

    ("a-block-kind-is-dropped", flaw_a_block_kind_is_dropped, {"D1"}, ""),
    ("sectpr-no-longer-last-after-generating", flaw_sectpr_no_longer_last,
     {"D1"}, ""),
    ("unsupported-silently-dropped", flaw_unsupported_silently_dropped, {"D2"}, ""),
    ("unsupported-reported-without-a-line",
     flaw_unsupported_reported_without_a_line, {"D2"}, ""),
    ("md-strict-writes-anyway", flaw_md_strict_writes_anyway, {"D3"}, ""),
    ("style-named-but-never-created", flaw_style_named_but_never_created,
     {"D4"}, ""),
    ("nested-list-flattened", flaw_nested_list_flattened, {"D5"}, ""),
    ("only-one-list-defined", flaw_only_one_list_defined, {"D5"}, ""),
    ("template-parts-lost", flaw_template_parts_lost, {"D6"}, ""),
    ("template-body-kept-as-well", flaw_template_body_kept_as_well, {"D6"}, ""),

    ("style-written-out-of-ct-style-order", flaw_style_written_out_of_ct_style_order,
     {"S1"}, ""),
    ("style-edit-rewrites-other-parts", flaw_style_edit_rewrites_other_parts,
     {"S1"}, ""),
    ("overwrite-allowed-silently", flaw_overwrite_allowed_silently, {"S2"}, ""),
    ("overwrite-does-not-say-how-many", flaw_overwrite_does_not_say_how_many,
     {"S2"}, ""),
    ("report-says-it-set-nothing", flaw_report_says_it_set_nothing, {"S3"}, ""),
    ("size-written-in-points-not-half-points",
     flaw_size_written_in_points_not_half_points, {"S3"}, ""),
    ("delete-a-style-still-in-use", flaw_delete_in_use_allowed, {"S4"}, ""),
    ("reassign-leaves-paragraphs-pointing-at-nothing",
     flaw_reassign_leaves_paragraphs_pointing_at_nothing, {"S4"}, ""),
    ("children-left-based-on-a-deleted-style",
     flaw_children_left_based_on_a_deleted_style, {"S4"}, ""),
    ("basedon-cycle-written", flaw_basedon_cycle_written, {"S5"}, ""),

    ("CONTROL: revised.docx stops removing a paragraph",
     flaw_revised_stops_removing_a_paragraph, {"V0"}, ""),
    ("CONTROL: revised.docx stops restyling a paragraph",
     flaw_revised_stops_restyling, {"V0"}, ""),
    ("CONTROL: revised.docx's header stops differing", flaw_revised_header_matches,
     {"V0"}, ""),
    ("diff-only-walks-word-document-xml", flaw_diff_only_walks_the_body, {"U1"}, ""),
    ("diff-never-descends-into-a-table", flaw_diff_flattens_tables, {"U1"}, ""),
    ("style-only-change-not-noticed", flaw_style_change_not_noticed, {"U1"}, ""),
    ("accepting-the-redline-does-not-give-b", flaw_accept_does_not_give_b, {"U2"}, ""),
    ("rejecting-the-redline-does-not-give-a", flaw_reject_does_not_give_a, {"U2"}, ""),
    ("script-will-not-claim-a-round-trip-it-made",
     flaw_script_will_not_claim_the_roundtrip, {"U2"}, ""),
    ("rsid-churn-reported-as-an-edit", flaw_noise_counted_as_a_difference, {"U3"}, ""),
    ("ignored-category-never-named", flaw_ignored_category_never_named, {"U3"}, ""),
    ("category-named-but-never-counted", flaw_category_named_but_never_counted,
     {"U3"}, ""),
    ("unmarkable-change-not-named", flaw_unmarkable_change_not_named, {"U4"}, ""),
    ("exact-round-trip-claimed-over-a-gap", flaw_exact_claimed_over_a_gap, {"U4"}, ""),
    ("strict-writes-a-partial-redline", flaw_strict_writes_a_partial_redline,
     {"U4"}, ""),
    ("row-change-not-reported-at-all", flaw_row_change_not_reported_at_all,
     {"U4"}, ""),
    ("move-reported-as-a-delete-and-an-insert",
     flaw_move_reported_as_delete_and_insert, {"U5"}, ""),
    ("redline-does-not-validate", flaw_redline_does_not_validate, {"U6"}, ""),
    ("redline-marks-the-text-but-not-the-paragraph-mark",
     flaw_redline_has_no_paragraph_mark, {"U6"}, ""),
    ("style-change-reported-but-never-written", flaw_redline_skips_the_style_change,
     {"U6"}, ""),
    ("redline-rewrites-the-styles-part", flaw_redline_rewrites_the_styles_part,
     {"U6"}, ""),
    ("swapped-picture-invisible-to-the-diff", flaw_picture_change_invisible,
     {"U7"}, ""),
    ("picture-change-left-looking-reviewed", flaw_picture_change_looks_reviewed,
     {"U7"}, ""),

    ("border-size-written-in-points-not-eighths", flaw_border_size_written_in_points,
     {"Q1"}, ""),
    ("border-edge-left-unstated", flaw_border_left_unstated, {"Q1"}, ""),
    ("list-presets-describes-a-different-table",
     flaw_list_presets_describes_something_else, {"Q1"}, ""),
    # All four are declared rather than left to the cascade note, because a non-empty
    # note masks EVERY unexpected check, not the ones it happens to mention. Declared
    # as {Q2, Q7} first, this row was quietly also lighting Q1 and Q4 and the run
    # still said PASS.
    ("two-presets-are-the-same-table", flaw_two_presets_are_the_same,
     {"Q1", "Q2", "Q4", "Q7"},
     "making banded identical to finance IS banded taking on finance's borders (Q1), "
     "finance's column padding (Q4), and losing the shading that is the only thing "
     "banded is for (Q7). Q2 owns the claim under test — that no two presets measure "
     "the same — and the other three are the same edit seen from each preset's own "
     "contract"),
    ("report-disagrees-with-the-document", flaw_report_disagrees_with_the_document,
     {"Q2"}, ""),
    ("header-repeat-flag-set-but-not-honoured", flaw_header_repeat_not_rendered,
     {"Q3"}, ""),
    ("CONTROL: every preset repeats the header", flaw_every_preset_repeats_the_header,
     {"Q3"}, ""),
    ("columns-sized-by-len", flaw_columns_sized_by_len, {"Q4"}, ""),
    ("fitted-widths-left-as-a-suggestion", flaw_fitted_widths_left_auto, {"Q4"}, ""),
    ("fitted-table-wraps-anyway", flaw_fitted_table_wraps_anyway, {"Q5"}, ""),
    ("CONTROL: len()-sized columns fit on this host too", flaw_naive_widths_fit_too,
     {"Q5"}, ""),
    ("preset-rewrites-the-styles-part", flaw_preset_rewrites_the_styles_part,
     {"Q6"}, ""),
    ("document-with-no-tables-reported-as-done", flaw_no_tables_reported_as_done,
     {"Q6"}, ""),
    ("out-of-range-table-index-accepted", flaw_out_of_range_table_accepted,
     {"Q6"}, ""),
    ("banding-counts-the-header-as-band-one", flaw_banding_starts_at_the_header,
     {"Q7"}, ""),
    ("every-preset-shades-a-row", flaw_every_preset_shades, {"Q7"}, ""),
    ("header-rule-drawn-as-insideh", flaw_header_rule_drawn_as_insideh, {"Q1", "Q7"},
     "Q1 also fires, and that is the shape of the defect rather than a cascade: "
     "drawing the rule with insideH means every pair of rows gets a line, which IS a "
     "change to the advertised border weights. Q1 owns the borders, Q7 owns the "
     "missing header rule"),
    ("CONTROL: a table probe string also occurs in the document",
     flaw_probe_string_also_in_the_document, {"V0"}, ""),

    ("entry-point-dies-on-a-windows-code-page",
     flaw_entry_point_dies_on_a_windows_code_page, {"C4"}, ""),
    ("a-new-entry-point-is-never-probed", flaw_a_new_entry_point_is_never_probed,
     {"C4"}, ""),
    ("probes-never-print-any-chinese", flaw_probes_print_nothing_chinese, {"C4"}, ""),
    ("CONTROL: the host stops reproducing the Windows code page",
     flaw_host_does_not_reproduce_the_code_page, {"C4"}, ""),
]


def fired(ctx: dict) -> dict[str, list[str]]:
    out = {}
    for cid, c in CHECKS.items():
        findings = c["fn"](ctx)
        if findings:
            out[cid] = findings
    return out


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--json", action="store_true")
    args = ap.parse_args()

    for f in (REPORT, REVISED, UNORDERED, OUTLINE, FONTLESS, CHART):
        if not f.is_file():
            print(f"[error] fixture missing: {f} (run fixtures/make_fixtures.py)",
                  file=sys.stderr)
            return 1

    results = []
    SKIPS.clear()
    with tempfile.TemporaryDirectory(prefix="docx-skill-test-") as td:
        work = Path(td)
        base = collect(work)

        clean = fired(base)
        results.append({"case": "real output of the real scripts", "expect": "silence",
                        "ok": not clean,
                        "detail": [f for v in clean.values() for f in v]})

        no_soffice = bool(base["pdf"].get("skipped"))
        matrix = []
        for name, mutate, expected, cascade in FLAWS:
            # A negative control whose check needs a tier this host lacks cannot
            # fire, and calling that a failure makes the suite red for a reason that
            # has nothing to do with the code. It is skipped and NAMED — never
            # folded into the pass count, and never quietly dropped either.
            if no_soffice and expected <= SOFFICE_CHECKS:
                SKIPS.append(f"negative control {name!r}: needs LibreOffice")
                continue
            ctx = mutate(copy.deepcopy(base), work)
            got = fired(ctx)
            unexpected = set(got) - expected
            missing = expected - set(got)
            matrix.append({"flaw": name, "expected": sorted(expected),
                           "fired": sorted(got), "cascade_note": cascade})
            detail = []
            if missing:
                detail.append(f"expected {sorted(missing)} to fire and it did not")
            if unexpected and not cascade:
                detail.append(f"unexpected checks fired: {sorted(unexpected)} — either "
                              f"a real cascade that needs documenting, or a check "
                              f"measuring the wrong thing")
            results.append({"case": f"flaw: {name}",
                            "expect": f"fires {sorted(expected)}",
                            "ok": not detail, "detail": detail,
                            "fired": sorted(got)})

    failed = [r for r in results if not r["ok"]]
    if args.json:
        print(json.dumps({"results": results, "matrix": matrix,
                          "skipped": SKIPS, "failed": len(failed)},
                         ensure_ascii=False, indent=2))
        return 1 if failed else 0

    for r in results:
        print(f"{'PASS' if r['ok'] else 'FAIL'}  [{r['expect']}] {r['case']}")
        for d in r["detail"][:6]:
            print(f"      · {d}")
    print("\n[flaw -> fired] every row must light the check that owns the defect, "
          "and nothing else:")
    for m in matrix:
        extra = sorted(set(m["fired"]) - set(m["expected"]))
        note = f"   (also {', '.join(extra)}: {m['cascade_note']})" if extra else ""
        print(f"  {m['flaw']:<52} -> {', '.join(m['fired']) or '(nothing)'}{note}")
    if SKIPS:
        print("\n[skipped] claims this host could not exercise:")
        for note in SKIPS:
            print(f"  - {note}")
    print(f"\n[docx-skill] {len(results) - len(failed)} passed, {len(failed)} failed, "
          f"{len(CHECKS)} assertions"
          + (f", {len(SKIPS)} skipped" if SKIPS else ""))
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
