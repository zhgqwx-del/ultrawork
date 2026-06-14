#!/usr/bin/env python3
"""Edit a .docx: find/replace text, append a paragraph. ultrawork doc-edit skill."""
from __future__ import annotations
import argparse, sys

def _require():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError:
        print("Missing dependency: python-docx (pip install python-docx)", file=sys.stderr)
        raise SystemExit(1)

def _replace_in_paragraph(paragraph, old, new):
    n = 0
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            n += run.text.count(new) if new else 1
    # 跨 run 回退：整段含 old 但单 run 不含时，在首 run 重写全段
    if old in paragraph.text and not any(old in r.text for r in paragraph.runs):
        if paragraph.runs:
            paragraph.runs[0].text = paragraph.text.replace(old, new)
            for r in paragraph.runs[1:]:
                r.text = ""
            n += 1
    return n

def main(argv):
    ap = argparse.ArgumentParser(description="Edit a .docx in place (or --out)")
    ap.add_argument("file")
    ap.add_argument("--replace", nargs=2, metavar=("OLD", "NEW"), action="append", default=[])
    ap.add_argument("--append-paragraph", metavar="TEXT", action="append", default=[])
    ap.add_argument("--out", help="write to this path instead of in place")
    args = ap.parse_args(argv)
    docx = _require()
    try:
        doc = docx.Document(args.file)
    except Exception as exc:  # noqa: BLE001
        print(f"Error opening {args.file}: {exc}", file=sys.stderr)
        return 1
    total = 0
    for old, new in args.replace:
        for p in doc.paragraphs:
            total += _replace_in_paragraph(p, old, new)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        total += _replace_in_paragraph(p, old, new)
    for text in args.append_paragraph:
        doc.add_paragraph(text)
    out = args.out or args.file
    doc.save(out)
    print(f"Saved {out} (replacements: {total}, appended: {len(args.append_paragraph)})")
    return 0

if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
