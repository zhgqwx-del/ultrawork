#!/usr/bin/env python3
"""Artifact-legality gate for the office skills (discussions/059 §5 L2).

L0 proves the code is ours, L1 proves the capability matrix is backed by something
that runs. Neither says the *artifacts* are valid. This does: it opens a generated
.docx/.xlsx/.pdf and asserts the things that actually break real documents —
element order, revision markup, CJK font binding, formula errors, field overflow,
tofu rendering.

Two modes:

    python3 scripts/office-skills-selftest.py                  # selftest (default)
    python3 scripts/office-skills-selftest.py --check out.docx --expect e.json

The selftest is the point. Every assertion below is exercised twice: once against a
synthesized *good* artifact (must stay silent) and once against an artifact carrying
exactly the defect it hunts (must fire). L0 and L1 both shipped a first version that
looked green while the guard itself was blind — a gate nobody has watched fail is
not evidence of anything.

--check is what S2-S4 call on real output. Expectations come from a JSON file:

    {"paragraphs": 6, "tables": 1, "contains": ["季度经营分析"],
     "sheets": {"利润表": {"B2": 1000, "D2": "=C2-B2"}},
     "recalc": {"D2": 200}, "finance_colors": true, "allow_blank_pages": false}

Tiers. Pure-python checks always run. Anything needing LibreOffice is skipped when
soffice is absent — never red, but always named in the report (no silent caps).

Exit 0 = all assertions behaved, 1 = something failed.
"""
from __future__ import annotations

import argparse
import json
import os
import platform
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import zipfile
from contextlib import closing
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


# ── thresholds ────────────────────────────────────────────────────────────────
# Calibrated 2026-08-01 against rendered samples (gotchas §10⑭). The figure is
# "share of CJK characters whose central 44% box carries ink at 200 DPI":
#
#     deckcraft export/deck.pdf, 24 real pages ... min 0.96, median 1.00
#     synthetic CJK prose (china-s 16pt) ......... 1.00
#     adversarial hollow-centre chars 口囗田日 .... 0.79
#     tofu / .notdef boxes ....................... 0.00
#
# 0.50 sits inside the 0.00-0.79 band. The first cut of this ran at GLYPH_INK 140
# and called page 5 of that real deck tofu at 0.47 — the calibration sample had
# been a single heavy 16pt synthetic line, and lighter antialiased strokes never
# reached the cutoff. At 200 the numbers are flat across cut 200/225/240 and DPI
# 200/300, so this is a plateau rather than a value tuned to sit just past one
# failing sample. Recalibrate against real output, never against a synthetic line.
TOFU_CENTER_INK = 0.05      # per-char: ink share that counts as "has strokes"
TOFU_MIN_FRACTION = 0.50    # per-page: share of CJK chars that must have them
TOFU_DPI = 200
BLANK_INK = 0.0005          # page ink share below this = blank
BLANK_DPI = 100

# byte -> 1 when inked. bytes.translate + count is C-speed; a Python pixel loop over
# a full page is not. Two cutoffs on purpose: blankness stays strict so a tinted but
# empty page still reads as blank, while glyph interiors count antialiased grey,
# which is most of a thin CJK stroke at any sane render resolution.
DARK = bytes(1 if v < 140 else 0 for v in range(256))
GLYPH_INK = bytes(1 if v < 200 else 0 for v in range(256))

ERROR_TOKENS = ("#REF!", "#DIV/0!", "#VALUE!", "#N/A", "#NAME?", "#NULL!", "#NUM!")

FIN_BLUE, FIN_BLACK, FIN_GREEN = "0000FF", "000000", "008000"

# A cross-sheet reference is "SheetName!A1", not "contains a bang". Testing for the
# bang alone marks =#REF!*2 as a link and demands green — caught by the flaw matrix,
# where the X2 control fired X4 as well. Error tokens are stripped before matching
# because they carry a bang of their own.
SHEET_REF = re.compile(r"(?:'[^']+'|\w+)!\$?[A-Za-z]{1,3}\$?\d+")


def is_cross_sheet(formula: str) -> bool:
    probe = formula
    for tok in ERROR_TOKENS:
        probe = probe.replace(tok, "")
    return bool(SHEET_REF.search(probe))

CJK_RANGES = ((0x3400, 0x4DBF), (0x4E00, 0x9FFF), (0xF900, 0xFAFF),
              (0x20000, 0x2A6DF))


def has_cjk(s: str) -> bool:
    return any(any(a <= ord(c) <= b for a, b in CJK_RANGES) for c in s)


def display_width(s: str) -> int:
    """Excel column-width units: a wide/fullwidth char occupies two."""
    return sum(2 if unicodedata.east_asian_width(c) in "WF" else 1 for c in s)


# ── external tools ────────────────────────────────────────────────────────────
def find_soffice() -> str | None:
    """LibreOffice, wherever this platform hides it (CLAUDE.md §13: no unix-only paths)."""
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    system = platform.system()
    if system == "Darwin":
        cand = [Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
                Path.home() / "Applications/LibreOffice.app/Contents/MacOS/soffice"]
    elif system == "Windows":
        cand = [Path(p) / "LibreOffice" / "program" / "soffice.exe"
                for p in (os.environ.get("ProgramFiles", ""),
                          os.environ.get("ProgramFiles(x86)", "")) if p]
    else:
        cand = [Path("/usr/bin/soffice"), Path("/usr/lib/libreoffice/program/soffice"),
                Path("/snap/bin/libreoffice")]
    return next((str(p) for p in cand if p.is_file()), None)


def find_xsd_dir() -> Path | None:
    """ECMA-376 schemas are publicly downloadable but not redistributable-by-default,
    so they are not vendored. Point at an unpacked copy to light up D2."""
    for cand in (os.environ.get("ECMA376_XSD_DIR"), REPO / "scripts" / "schemas" / "ecma376"):
        if cand and Path(cand).is_dir():
            return Path(cand)
    return None


SOFFICE = find_soffice()
XSD_DIR = find_xsd_dir()
PDFTOPPM = shutil.which("pdftoppm")


def soffice_convert(src: Path, fmt: str, outdir: Path) -> tuple[Path | None, str]:
    """Convert via LibreOffice into an isolated profile.

    -env:UserInstallation is not optional: without it soffice writes into the real
    ~/.config profile and a test run mutates the user's machine.
    """
    if not SOFFICE:
        return None, "soffice not found"
    profile = outdir / "lo-profile"
    cmd = [SOFFICE, f"-env:UserInstallation={profile.as_uri()}", "--headless",
           "--norestore", "--convert-to", fmt, "--outdir", str(outdir), str(src)]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8",
                           errors="replace", timeout=180)
    except (subprocess.TimeoutExpired, OSError) as e:
        return None, f"soffice failed: {e}"
    ext = fmt.split(":")[0]
    out = outdir / (src.stem + "." + ext)
    if r.returncode != 0 or not out.is_file():
        return None, f"soffice exited {r.returncode}: {(r.stdout + r.stderr).strip()[:400]}"
    return out, ""


# ── raster helpers (shared by the PDF checks and the docx→pdf tier) ───────────
def page_ink(page, dpi: int = BLANK_DPI) -> float:
    import fitz
    pix = page.get_pixmap(dpi=dpi, colorspace=fitz.csGRAY)
    data = pix.samples
    return (data.translate(DARK).count(1) / len(data)) if data else 0.0


def cjk_center_ink_fraction(page) -> tuple[float, int]:
    """Share of CJK glyphs on the page whose central box carries ink.

    A .notdef box is hollow: ink on the perimeter, nothing inside. Real characters
    almost always cross their own centre. Returns (fraction, cjk_char_count).

    The rotation_matrix is not optional. get_text reports character boxes in PAGE
    space (unrotated) while get_pixmap renders DISPLAY space; on a page carrying a
    /Rotate the two disagree and every character gets measured against a patch of
    empty paper. Found 2026-08-01 by the first artifact the pdf skill produced: a
    perfectly rendered 90-degree page scored 0.04 here and 1.00 once mapped. The
    pdf skill has the same trap in its own bbox output (bbox vs bbox_display).
    """
    import fitz
    chars = [c for b in page.get_text("rawdict")["blocks"] if b.get("lines")
             for l in b["lines"] for sp in l["spans"] for c in sp.get("chars", [])
             if has_cjk(c.get("c", ""))]
    if not chars:
        return 0.0, 0
    pix = page.get_pixmap(dpi=TOFU_DPI, colorspace=fitz.csGRAY)
    scale, width, data = TOFU_DPI / 72.0, pix.width, pix.samples
    rotation = page.rotation_matrix
    inked = 0
    for c in chars:
        box = fitz.Rect(c["bbox"]) * rotation
        x0, y0, x1, y1 = (v * scale for v in box)
        bw, bh = x1 - x0, y1 - y0
        if bw < 3 or bh < 3:
            continue
        left, right = int(x0 + bw * 0.28), int(x1 - bw * 0.28)
        top, bottom = int(y0 + bh * 0.28), int(y1 - bh * 0.28)
        left, top = max(left, 0), max(top, 0)
        right, bottom = min(right, width), min(bottom, pix.height)
        dark = total = 0
        for yy in range(top, bottom):
            row = data[yy * width + left: yy * width + right]
            dark += row.translate(GLYPH_INK).count(1)
            total += len(row)
        if total and dark / total > TOFU_CENTER_INK:
            inked += 1
    return inked / len(chars), len(chars)


def render_findings(pdf: Path, tag: str, allow_blank: bool, want_cjk: bool) -> list[str]:
    """Rasterize + assert non-blank and non-tofu. Used for real PDFs and for the
    PDF LibreOffice produces from a .docx."""
    import fitz
    out: list[str] = []
    try:
        doc = fitz.open(pdf)
    except Exception as e:  # noqa: BLE001 - any malformed PDF lands here
        return [f"{tag} cannot open rendered PDF: {e}"]
    with doc:
        if doc.page_count == 0:
            return [f"{tag} rendered PDF has 0 pages"]
        for i, page in enumerate(doc, 1):
            try:
                ink = page_ink(page)
            except Exception as e:  # noqa: BLE001
                out.append(f"{tag} page {i} failed to rasterize: {e}")
                continue
            if ink < BLANK_INK and not allow_blank:
                out.append(f"{tag} page {i} renders blank (ink {ink:.5f} < {BLANK_INK})")
            frac, n = cjk_center_ink_fraction(page)
            if n and frac < TOFU_MIN_FRACTION:
                out.append(f"{tag} page {i} renders as tofu: only {frac:.0%} of {n} CJK "
                           f"glyphs have interior strokes (need {TOFU_MIN_FRACTION:.0%})")
            if want_cjk and i == 1 and n == 0:
                out.append(f"{tag} page 1 has no CJK glyphs in the text layer")
    if PDFTOPPM:  # second opinion when poppler is around; PyMuPDF is the baseline
        r = subprocess.run([PDFTOPPM, "-png", "-r", "50", "-f", "1", "-l", "1",
                            str(pdf), str(pdf.with_suffix(""))],
                           capture_output=True, text=True, errors="replace", timeout=120)
        if r.returncode != 0:
            out.append(f"{tag} pdftoppm exited {r.returncode}: {r.stderr.strip()[:200]}")
    return out


# ── docx plumbing ─────────────────────────────────────────────────────────────
class DocxUnreadable(Exception):
    """The main part could not be reached at all — D1's business, not D4-D6's."""


def docx_parts(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as z:
        return {n: z.read(n) for n in z.namelist()}


def docx_tree(path: Path):
    """Parse word/document.xml, or raise DocxUnreadable with the reason.

    Everything downstream of D1 needs this to fail as a finding rather than a
    traceback: a package broken enough to hide document.xml must still produce a
    report, not a crash.
    """
    from lxml import etree
    try:
        parts = docx_parts(path)
    except (zipfile.BadZipFile, OSError) as e:
        raise DocxUnreadable(f"not a readable zip: {e}") from e
    if "word/document.xml" not in parts:
        raise DocxUnreadable("package has no word/document.xml")
    try:
        return etree.fromstring(parts["word/document.xml"])
    except etree.XMLSyntaxError as e:
        raise DocxUnreadable(f"document.xml will not parse: {e}") from e


def rewrite_zip(src: Path, dst: Path, mutate) -> None:
    """Copy a zip, passing each (name, bytes) through mutate; None drops the entry."""
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = mutate(item.filename, zin.read(item.filename))
            if data is not None:
                zout.writestr(item, data)


def run_text(run) -> str:
    return "".join(t.text or "" for t in run
                   if t.tag in (w("t"), w("delText"), w("noBreakHyphen")))


# ── the assertions ────────────────────────────────────────────────────────────
CHECKS: dict[str, dict] = {}


def check(cid: str, kind: str | tuple[str, ...], title: str, tier: str = "core"):
    kinds = (kind,) if isinstance(kind, str) else tuple(kind)

    def deco(fn):
        CHECKS[cid] = {"id": cid, "kind": kinds, "title": title, "tier": tier, "fn": fn}
        return fn
    return deco


@check("D1", "docx", "package integrity: required parts, well-formed XML, resolvable r:id")
def d1_package(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    out: list[str] = []
    try:
        parts = docx_parts(path)
    except zipfile.BadZipFile as e:
        return [f"D1 not a readable zip: {e}"]
    for required in ("[Content_Types].xml", "word/document.xml", "_rels/.rels"):
        if required not in parts:
            out.append(f"D1 missing required part: {required}")
    trees = {}
    for name, data in parts.items():
        if not name.endswith(".xml") and not name.endswith(".rels"):
            continue
        try:
            trees[name] = etree.fromstring(data)
        except etree.XMLSyntaxError as e:
            out.append(f"D1 malformed XML in {name}: {e}")
    doc = trees.get("word/document.xml")
    rels = trees.get("word/_rels/document.xml.rels")
    if doc is not None:
        known = {e.get("Id") for e in rels} if rels is not None else set()
        used = {v for el in doc.iter() for k, v in el.attrib.items()
                if k.startswith(f"{{{R}}}") and k.endswith(("id", "embed", "link"))}
        for rid in sorted(used - known):
            out.append(f"D1 document.xml references r:id={rid} with no matching relationship")
    return out


@check("D2", "docx", "ECMA-376 XSD schema validation", tier="xsd")
def d2_xsd(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    schema_file = next((p for p in (XSD_DIR / "wml.xsd", XSD_DIR / "wml.xsd".upper())
                        if p.is_file()), None) if XSD_DIR else None
    if schema_file is None:
        return [f"D2 no wml.xsd under {XSD_DIR}"]
    schema = etree.XMLSchema(etree.parse(str(schema_file)))
    try:
        tree = etree.ElementTree(docx_tree(path))
    except DocxUnreadable as e:
        return [f"D2 {e}"]
    if schema.validate(tree):
        return []
    return [f"D2 schema violation: {e.message} (line {e.line})" for e in schema.error_log]


@check("D3", "docx", "python-docx round-trip: opens, counts and text match expectation")
def d3_roundtrip(path: Path, expect: dict) -> list[str]:
    import docx as pydocx
    try:
        doc = pydocx.Document(str(path))
    except Exception as e:  # noqa: BLE001 - python-docx raises a zoo of types
        return [f"D3 python-docx cannot open the file: {type(e).__name__}: {e}"]
    out: list[str] = []
    if "paragraphs" in expect and len(doc.paragraphs) != expect["paragraphs"]:
        out.append(f"D3 paragraph count {len(doc.paragraphs)} != expected {expect['paragraphs']}")
    if "tables" in expect and len(doc.tables) != expect["tables"]:
        out.append(f"D3 table count {len(doc.tables)} != expected {expect['tables']}")
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    for needle in expect.get("contains", []):
        if needle not in text:
            out.append(f"D3 expected text not found after round-trip: {needle!r}")
    return out


@check("D4", "docx", "ECMA-376 element order: w:pPr first in w:p, w:rPr first in w:r")
def d4_order(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    try:
        doc = docx_tree(path)
    except DocxUnreadable as e:
        return [f"D4 {e}"]
    out: list[str] = []
    # CT_P / CT_R / CT_Tbl put their property element first; CT_Body puts sectPr last.
    # Word repairs some of this silently, LibreOffice does not, and a "repaired"
    # prompt on open is exactly the failure this gate exists to stop.
    for tag, prop in ((w("p"), w("pPr")), (w("r"), w("rPr")),
                      (w("tbl"), w("tblPr")), (w("tc"), w("tcPr")),
                      (w("tr"), w("trPr"))):
        for i, el in enumerate(doc.iter(tag)):
            kids = [k for k in el if isinstance(k.tag, str)]
            at = [n for n, k in enumerate(kids) if k.tag == prop]
            if at and at[0] != 0:
                out.append(f"D4 <{etree.QName(tag).localname}> #{i}: "
                           f"<{etree.QName(prop).localname}> at position {at[0]}, must be first")
            if len(at) > 1:
                out.append(f"D4 <{etree.QName(tag).localname}> #{i}: "
                           f"{len(at)} <{etree.QName(prop).localname}> elements, max 1")
    for body in doc.iter(w("body")):
        kids = [k for k in body if isinstance(k.tag, str)]
        for n, k in enumerate(kids):
            if k.tag == w("sectPr") and n != len(kids) - 1:
                out.append(f"D4 <w:sectPr> at position {n} of {len(kids)}, must be last in body")
    return out


@check("D5", "docx", "revision markup: w:del holds only w:delText, w:ins holds no w:delText")
def d5_revisions(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    try:
        doc = docx_tree(path)
    except DocxUnreadable as e:
        return [f"D5 {e}"]
    out: list[str] = []
    # The single most common way a hand-edited .docx becomes unopenable: text kept as
    # <w:t> inside a deletion, or delText left behind inside an insertion.
    for i, dele in enumerate(doc.iter(w("del"))):
        for t in dele.iter(w("t")):
            out.append(f"D5 <w:del> #{i}: contains <w:t>{(t.text or '')[:24]!r}, "
                       f"must be <w:delText>")
    for i, ins in enumerate(doc.iter(w("ins"))):
        for t in ins.iter(w("delText")):
            out.append(f"D5 <w:ins> #{i}: contains <w:delText>{(t.text or '')[:24]!r}, "
                       f"deleted text cannot live in an insertion")
    for tag in (w("del"), w("ins")):
        name = etree.QName(tag).localname
        for i, el in enumerate(doc.iter(tag)):
            for attr in ("id", "author"):
                if el.get(w(attr)) is None:
                    out.append(f"D5 <w:{name}> #{i}: missing required w:{attr}")
    return out


@check("D6", "docx", "CJK runs carry both w:rFonts/@w:ascii and @w:eastAsia")
def d6_cjk_fonts(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    try:
        doc = docx_tree(path)
    except DocxUnreadable as e:
        return [f"D6 {e}"]
    out: list[str] = []
    # Word picks the CJK face from @eastAsia and the latin face from @ascii. Set only
    # one and mixed "2026 年营收" text renders in two unrelated typefaces — or, when
    # the theme font has no CJK coverage, not at all.
    for i, run in enumerate(doc.iter(w("r"))):
        text = run_text(run)
        if not has_cjk(text):
            continue
        rpr = run.find(w("rPr"))
        rfonts = rpr.find(w("rFonts")) if rpr is not None else None
        if rfonts is None:
            out.append(f"D6 run #{i} {text[:16]!r}: no <w:rFonts>")
            continue
        for attr in ("ascii", "eastAsia"):
            if not rfonts.get(w(attr)) and not rfonts.get(w(attr + "Theme")):
                out.append(f"D6 run #{i} {text[:16]!r}: <w:rFonts> has no @w:{attr}")
    return out


@check("D7", "docx", "LibreOffice renders it: convert to PDF, non-blank, non-tofu",
       tier="soffice")
def d7_render(path: Path, expect: dict) -> list[str]:
    with tempfile.TemporaryDirectory() as td:
        pdf, err = soffice_convert(path, "pdf", Path(td))
        if pdf is None:
            return [f"D7 {err}"]
        return render_findings(pdf, "D7", expect.get("allow_blank_pages", False),
                               want_cjk=any(has_cjk(s) for s in expect.get("contains", [])))


@check("X1", "xlsx", "openpyxl round-trip: values and formulas read back as written")
def x1_roundtrip(path: Path, expect: dict) -> list[str]:
    import openpyxl
    try:
        wb = openpyxl.load_workbook(path)
    except Exception as e:  # noqa: BLE001
        return [f"X1 openpyxl cannot open the file: {type(e).__name__}: {e}"]
    out: list[str] = []
    with closing(wb):
        for sheet, cells in (expect.get("sheets") or {}).items():
            if sheet not in wb.sheetnames:
                out.append(f"X1 missing sheet {sheet!r} (have {wb.sheetnames})")
                continue
            ws = wb[sheet]
            for ref, want in cells.items():
                got = ws[ref].value
                if got != want:
                    out.append(f"X1 {sheet}!{ref} = {got!r}, expected {want!r}")
    return out


@check("X2", "xlsx", "no formula errors anywhere (#REF! #DIV/0! #VALUE! #N/A #NAME? ...)")
def x2_errors(path: Path, expect: dict) -> list[str]:
    import openpyxl
    out: list[str] = []
    # Scan both views: the formula text (a broken reference is baked into the string)
    # and the cached value (what the last calculating app left behind).
    for data_only in (False, True):
        try:
            wb = openpyxl.load_workbook(path, data_only=data_only)
        except Exception as e:  # noqa: BLE001
            return [f"X2 openpyxl cannot open the file: {type(e).__name__}: {e}"]
        with closing(wb):
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        v = cell.value
                        if not isinstance(v, str):
                            continue
                        for tok in ERROR_TOKENS:
                            if tok in v:
                                where = "formula" if not data_only else "cached value"
                                msg = f"X2 {ws.title}!{cell.coordinate} {where} carries {tok}"
                                if msg not in out:
                                    out.append(msg)
    return out


@check("X3", "xlsx", "LibreOffice recalculation agrees with the expected values",
       tier="soffice")
def x3_recalc(path: Path, expect: dict) -> list[str]:
    want = expect.get("recalc") or {}
    if not want:
        return []
    with tempfile.TemporaryDirectory() as td:
        # csv forces a real recalculation on load and writes values, not formulas.
        # It only covers the first sheet, which is why "recalc" is a flat cell map.
        out_csv, err = soffice_convert(path, "csv", Path(td))
        if out_csv is None:
            return [f"X3 {err}"]
        import csv as csvmod
        with out_csv.open(encoding="utf-8", newline="") as fh:
            grid = list(csvmod.reader(fh))
    from openpyxl.utils.cell import coordinate_from_string, column_index_from_string
    findings: list[str] = []
    for ref, expected in want.items():
        col, row = coordinate_from_string(ref)
        ci, ri = column_index_from_string(col) - 1, row - 1
        got = grid[ri][ci] if ri < len(grid) and ci < len(grid[ri]) else None
        try:
            same = got is not None and abs(float(got) - float(expected)) < 1e-6
        except (TypeError, ValueError):
            same = str(got) == str(expected)
        if not same:
            findings.append(f"X3 recalculated {ref} = {got!r}, expected {expected!r}")
    return findings


@check("X4", "xlsx", "financial colour convention: blue inputs, black formulas, green links")
def x4_finance_colors(path: Path, expect: dict) -> list[str]:
    import openpyxl
    # Opt-in: the blue/black/green convention belongs to financial models. Defaulting
    # it on reds every ordinary sheet — the repo's own sample.xlsx fixture included.
    if not expect.get("finance_colors", False):
        return []
    try:
        wb = openpyxl.load_workbook(path)
    except Exception as e:  # noqa: BLE001
        return [f"X4 openpyxl cannot open the file: {type(e).__name__}: {e}"]
    out: list[str] = []

    def rgb(cell) -> str:
        c = cell.font.color if cell.font else None
        # theme/indexed colours carry no rgb; treat them as the default black.
        if c is None or c.type != "rgb" or not isinstance(c.rgb, str):
            return FIN_BLACK
        return c.rgb[-6:].upper()

    with closing(wb):
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    v = cell.value
                    if isinstance(v, str) and v.startswith("="):
                        cross = is_cross_sheet(v)
                        want = FIN_GREEN if cross else FIN_BLACK
                        kind = "cross-sheet link" if cross else "formula"
                    elif isinstance(v, (int, float)) and not isinstance(v, bool):
                        want, kind = FIN_BLUE, "hard-coded input"
                    else:
                        continue  # labels, dates, blanks are not part of the convention
                    if rgb(cell) != want:
                        out.append(f"X4 {ws.title}!{cell.coordinate} is a {kind}: "
                                   f"font {rgb(cell)}, convention wants {want}")
    return out


def horizontally_merged(ws) -> set[str]:
    """Cells whose text is displayed across more than one column."""
    out: set[str] = set()
    for rng in ws.merged_cells.ranges:
        if rng.max_col > rng.min_col:
            out.update(c.coordinate for row in ws[rng.coord] for c in row)
    return out


@check("X5", "xlsx", "CJK column widths counted in wide characters (no #### truncation)")
def x5_cjk_width(path: Path, expect: dict) -> list[str]:
    import openpyxl
    from openpyxl.utils import get_column_letter
    try:
        wb = openpyxl.load_workbook(path)
    except Exception as e:  # noqa: BLE001
        return [f"X5 openpyxl cannot open the file: {type(e).__name__}: {e}"]
    out: list[str] = []
    with closing(wb):
        for ws in wb.worksheets:
            spanned = horizontally_merged(ws)
            needed: dict[int, tuple[int, str]] = {}
            for row in ws.iter_rows():
                for cell in row:
                    v = cell.value
                    # A formula's own text is never displayed, so it must not drive width.
                    if not isinstance(v, str) or v.startswith("=") or not has_cjk(v):
                        continue
                    # A title merged across A1:F1 is DISPLAYED across all six columns,
                    # so demanding that column A alone fit it reds a workbook that
                    # renders perfectly — verified by converting one to PDF and
                    # reading all 20 characters back. A merge inside a single column
                    # (a vertical merge) gives no extra room and still counts.
                    if cell.coordinate in spanned:
                        continue
                    want = min(display_width(v) + 2, 60)
                    if want > needed.get(cell.column, (0, ""))[0]:
                        needed[cell.column] = (want, v)
            for col, (want, sample) in sorted(needed.items()):
                letter = get_column_letter(col)
                dim = ws.column_dimensions.get(letter)
                width = dim.width if dim is not None else None
                if width is None:
                    out.append(f"X5 {ws.title}!{letter} holds CJK ({sample[:12]!r}) but has "
                               f"no explicit width; the 8.43 default truncates it")
                elif width + 1e-6 < want:
                    out.append(f"X5 {ws.title}!{letter} width {width:g} < {want} needed for "
                               f"{sample[:12]!r} counted in wide characters")
    return out


@check("P1", "pdf", "rasterizes, and no page comes out blank")
def p1_render(path: Path, expect: dict) -> list[str]:
    return render_findings(path, "P1", expect.get("allow_blank_pages", False),
                           want_cjk=False)


@check("P2", "pdf", "text round-trip: what was written extracts back out")
def p2_text(path: Path, expect: dict) -> list[str]:
    import fitz
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        return [f"P2 cannot open PDF: {e}"]
    with doc:
        text = "\n".join(p.get_text() for p in doc)
    return [f"P2 expected text not extractable: {n!r}" for n in expect.get("contains", [])
            if n not in text]


@check("P3", "pdf", "filled form values stay inside their field boxes")
def p3_field_bbox(path: Path, expect: dict) -> list[str]:
    import fitz
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        return [f"P3 cannot open PDF: {e}"]
    out: list[str] = []
    with doc:
        for pno, page in enumerate(doc, 1):
            widgets = list(page.widgets() or [])
            spans = [(s["text"], fitz.Rect(s["bbox"]))
                     for b in page.get_text("dict")["blocks"] if b.get("lines")
                     for l in b["lines"] for s in l["spans"]]
            for wd in widgets:
                val = wd.field_value
                if not isinstance(val, str) or not val.strip():
                    continue
                rect = wd.rect
                size = wd.text_fontsize or 0
                if size:
                    font = "china-s" if has_cjk(val) else \
                        (wd.text_font or "helv").lower().replace(" ", "")
                    try:
                        natural = fitz.get_text_length(val, fontname=font, fontsize=size)
                    except Exception:  # noqa: BLE001 - unmapped font name
                        try:
                            natural = fitz.get_text_length(val, fontname="helv", fontsize=size)
                        except Exception:  # noqa: BLE001
                            natural = None
                    # A viewer silently clips the overflow, so the damage is invisible
                    # in the raster: the value is simply gone. Measure the text instead.
                    if natural is not None and natural > rect.width - 2:
                        out.append(f"P3 page {pno} field {wd.field_name!r}: value needs "
                                   f"{natural:.1f}pt, box is {rect.width:.1f}pt wide")
                for text, bbox in spans:
                    if not bbox.intersects(rect) or not text.strip():
                        continue
                    if bbox.x1 > rect.x1 + 0.5 or bbox.x0 < rect.x0 - 0.5:
                        out.append(f"P3 page {pno} field {wd.field_name!r}: rendered text "
                                   f"{text[:20]!r} spans {bbox.x0:.1f}-{bbox.x1:.1f}, "
                                   f"box is {rect.x0:.1f}-{rect.x1:.1f}")
    return out


@check("P4", "pdf", "CJK survives the round-trip and renders as glyphs, not tofu")
def p4_cjk(path: Path, expect: dict) -> list[str]:
    import fitz
    wanted = [s for s in expect.get("contains", []) if has_cjk(s)]
    if not wanted:
        return []
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        return [f"P4 cannot open PDF: {e}"]
    out: list[str] = []
    with doc:
        text = "\n".join(p.get_text() for p in doc)
        for s in wanted:
            if s not in text:
                out.append(f"P4 CJK does not extract back: {s!r} "
                           f"(got {''.join(ch for ch in text if has_cjk(ch))[:24]!r})")
        for i, page in enumerate(doc, 1):
            frac, n = cjk_center_ink_fraction(page)
            if n and frac < TOFU_MIN_FRACTION:
                out.append(f"P4 page {i} renders as tofu: only {frac:.0%} of {n} CJK glyphs "
                           f"have interior strokes (need {TOFU_MIN_FRACTION:.0%})")
    return out


# ── fidelity: what the edit was NOT supposed to touch ─────────────────────────
# Everything above asks "is this artifact legal". Nothing above asks "is this still
# the user's document". A no-op round-trip through python-docx drops
# word/_rels/document.xml.rels and through openpyxl drops xl/metadata.xml — both
# files still open, both pass all sixteen assertions above. For a skill family whose
# core verb is "edit", losing a chart or a macro the user never mentioned is the
# failure that matters most, and legality cannot see it.
#
# These need expect["baseline"]: the path the artifact was edited FROM.

KEY_DOCX_PART = re.compile(
    r"^word/(styles|numbering|settings|fontTable|footnotes|endnotes"
    r"|header\d*|footer\d*)\.xml$")


def zip_parts(path: Path) -> set[str]:
    with zipfile.ZipFile(path) as z:
        return {n for n in z.namelist() if not n.endswith("/")}


def part_is_inert(name: str, data: bytes) -> bool:
    """True when dropping this part loses nothing.

    Counting parts alone is not a fidelity test. sample.docx carries a
    word/_rels/document.xml.rels holding zero Relationship elements — python-docx
    declines to write it back, and that is correct, not data loss. sample.xlsx
    carries a 904-byte xl/metadata.xml of real dynamic-array metadata, and losing
    that IS data loss. Same symptom, opposite verdicts, and only the content tells
    them apart.
    """
    from lxml import etree
    if not data.strip():
        return True
    if not name.endswith((".xml", ".rels")):
        return False
    try:
        root = etree.fromstring(data)
    except etree.XMLSyntaxError:
        return False
    return len(root) == 0 and not (root.text or "").strip()


def baselines_of(expect: dict) -> list[Path]:
    """`baseline` is one path, or several when the artifact was merged from several."""
    b = expect.get("baseline")
    if not b:
        return []
    return [Path(b)] if isinstance(b, (str, Path)) else [Path(x) for x in b]


def baseline_of(expect: dict) -> Path | None:
    """The single input an edit came from; None when there is no baseline at all.

    A list means the caller merged several files. The docx/xlsx fidelity checks have
    no meaning for that, so they say so instead of silently examining the first one.
    """
    paths = baselines_of(expect)
    return paths[0] if len(paths) == 1 else None


def multi_baseline_note(cid: str, expect: dict) -> list[str]:
    n = len(baselines_of(expect))
    if n > 1:
        return [f"{cid} takes a single baseline; {n} were given — a merged artifact "
                f"has no single input to compare against"]
    return []


@check("F1", ("docx", "xlsx"), "fidelity: every part of the input survives the edit")
def f1_parts(path: Path, expect: dict) -> list[str]:
    base = baseline_of(expect)
    if base is None:
        return multi_baseline_note("F1", expect)
    if not base.is_file():
        return [f"F1 baseline not found: {base}"]
    try:
        before, after = zip_parts(base), zip_parts(path)
    except zipfile.BadZipFile as e:
        return [f"F1 cannot compare packages: {e}"]
    allowed = set(expect.get("may_drop", []))
    out: list[str] = []
    with zipfile.ZipFile(base) as z:
        for n in sorted(before - after - allowed):
            if part_is_inert(n, z.read(n)):
                continue
            out.append(f"F1 edit dropped {n} ({z.getinfo(n).file_size} bytes) — "
                       f"present in the input, gone from the output")
    return out


@check("F2", "docx", "fidelity: styles, numbering, headers and relationships survive")
def f2_docx_structure(path: Path, expect: dict) -> list[str]:
    from lxml import etree
    base = baseline_of(expect)
    if base is None:
        return multi_baseline_note("F2", expect)
    if not base.is_file():
        return [f"F2 baseline not found: {base}"]
    out: list[str] = []
    try:
        before, after = zip_parts(base), zip_parts(path)
    except zipfile.BadZipFile as e:
        return [f"F2 cannot compare packages: {e}"]
    # Named separately from F1 so the report says WHICH capability was silently
    # dropped — "numbering.xml is gone" means every list in the document is now flat.
    for n in sorted(p for p in before if KEY_DOCX_PART.match(p)):
        if n not in after:
            out.append(f"F2 {n} was in the input and is not in the output")

    def rel_count(pkg: Path) -> int | None:
        try:
            with zipfile.ZipFile(pkg) as z:
                if "word/_rels/document.xml.rels" not in z.namelist():
                    return 0
                return len(etree.fromstring(z.read("word/_rels/document.xml.rels")))
        except (zipfile.BadZipFile, etree.XMLSyntaxError):
            return None

    b, a = rel_count(base), rel_count(path)
    if b is not None and a is not None and a < b:
        out.append(f"F2 document relationships dropped from {b} to {a} — images, "
                   f"hyperlinks or headers lost their wiring")
    return out


@check("F3", "xlsx", "fidelity: sheets, charts, conditional formats and panes survive")
def f3_xlsx_features(path: Path, expect: dict) -> list[str]:
    import openpyxl
    base = baseline_of(expect)
    if base is None:
        return multi_baseline_note("F3", expect)
    if not base.is_file():
        return [f"F3 baseline not found: {base}"]

    def features(pkg: Path) -> dict:
        # Charts are counted from the package, not from openpyxl: a chart it failed
        # to parse would read as "no chart" on both sides and cancel out.
        charts = sum(1 for n in zip_parts(pkg) if n.startswith("xl/charts/chart"))
        wb = openpyxl.load_workbook(pkg)
        with closing(wb):
            per = {ws.title: {
                "cf": len(ws.conditional_formatting._cf_rules),
                "dv": len(ws.data_validations.dataValidation),
                "merged": len(ws.merged_cells.ranges),
                "widths": len(ws.column_dimensions),
                "freeze": ws.freeze_panes,
                "filter": ws.auto_filter.ref,
            } for ws in wb.worksheets}
        return {"sheets": list(per), "charts": charts, "per": per}

    try:
        b, a = features(base), features(path)
    except Exception as e:  # noqa: BLE001
        return [f"F3 cannot compare workbooks: {type(e).__name__}: {e}"]
    out: list[str] = []
    for name in b["sheets"]:
        if name not in a["sheets"]:
            out.append(f"F3 sheet {name!r} was in the input and is not in the output")
    if a["charts"] < b["charts"]:
        out.append(f"F3 chart count dropped from {b['charts']} to {a['charts']}")
    for name, bf in b["per"].items():
        af = a["per"].get(name)
        if af is None:
            continue
        for key, label in (("cf", "conditional-format rules"), ("dv", "data validations"),
                           ("merged", "merged ranges"), ("widths", "explicit column widths")):
            if af[key] < bf[key]:
                out.append(f"F3 {name}: {label} dropped from {bf[key]} to {af[key]}")
        for key, label in (("freeze", "frozen panes"), ("filter", "auto filter")):
            if bf[key] and not af[key]:
                out.append(f"F3 {name}: {label} ({bf[key]}) was lost")
    return out


@check("P5", "pdf", "fidelity: page count holds and untouched pages keep their text")
def p5_pdf_fidelity(path: Path, expect: dict) -> list[str]:
    """One baseline is an edit; several is a merge.

    A merge that quietly drops its second input produces a perfectly legal PDF, and
    nothing else in this file would notice — page count, text extraction and the
    raster are all happy with it. Walking the inputs in order is what catches it.
    """
    import fitz
    bases = baselines_of(expect)
    if not bases:
        return []
    missing = [str(b) for b in bases if not b.is_file()]
    if missing:
        return [f"P5 baseline not found: {', '.join(missing)}"]
    touched = set(expect.get("touched_pages", []))
    out: list[str] = []
    try:
        dst = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        return [f"P5 cannot open the artifact: {e}"]
    with dst:
        offset = 0
        for base in bases:
            try:
                src = fitz.open(base)
            except Exception as e:  # noqa: BLE001
                return [f"P5 cannot open baseline {base.name}: {e}"]
            with src:
                if dst.page_count < offset + src.page_count:
                    out.append(f"P5 {base.name} contributes {src.page_count} page(s) "
                               f"at index {offset}, but the artifact has only "
                               f"{dst.page_count}")
                for i in range(min(src.page_count, max(dst.page_count - offset, 0))):
                    if (offset + i + 1) in touched:
                        continue
                    b = src[i].get_text().strip()
                    a = dst[offset + i].get_text().strip()
                    if a != b:
                        where = (f"page {offset + i + 1}" if len(bases) == 1 else
                                 f"page {offset + i + 1} (page {i + 1} of {base.name})")
                        out.append(f"P5 {where} was not meant to be edited but its "
                                   f"text changed ({len(b)} chars -> {len(a)})")
                offset += src.page_count
    return out


KIND_BY_SUFFIX = {".docx": "docx", ".xlsx": "xlsx", ".pdf": "pdf"}


# LibreOffice is a REQUIRED dependency of the docx/xlsx skills (059 §7, decided
# 2026-08-01), so a missing soffice is a failure, not a skip — the contract says the
# machine has it. Same for the ECMA-376 schemas, which CI is expected to supply.
# --allow-missing exists for dev machines and must name the tier out loud; it never
# turns into a silent pass.
REQUIRED_TIERS = ("soffice", "xsd")
ALL_TIERS = frozenset(REQUIRED_TIERS)

# What still covers a tier when the tool behind it is missing. Saying "skipped" and
# leaving it there invites reading the run as fully green; these say exactly how much
# of each unavailable assertion has actually been watched work.
TIER_RESIDUAL = {
    "soffice": ("its failure path IS covered (fault injection) and its raster half is "
                "shared with P1/P4; only the LibreOffice conversion itself is unrun"),
    "xsd": "NO residual coverage — neither its pass nor its fail path has ever run here",
}


def tier_available(tier: str) -> tuple[bool, str]:
    if tier == "soffice":
        return bool(SOFFICE), "LibreOffice (soffice) not installed"
    if tier == "xsd":
        return bool(XSD_DIR), "ECMA-376 schemas not present (set $ECMA376_XSD_DIR)"
    return True, ""


# Assertions that have nothing to compare against unless the caller supplies it.
# Without this the empty-expectations case reads as a clean pass — which is how L1
# shipped a green C4 while not one capability had been declared.
# Fidelity assertions only apply to an EDIT. A capability that generates a document
# from nothing has no baseline, and reporting that as a missing expectation would
# push authors to fake one. Kept separate so callers can distinguish "not
# applicable here" from "you forgot to say what you expect".
FIDELITY_CHECKS = ("F1", "F2", "F3", "P5")

NEEDS_EXPECTATION = {
    "F1": ("baseline",), "F2": ("baseline",), "F3": ("baseline",), "P5": ("baseline",),
    "D3": ("paragraphs", "tables", "contains"),
    "X1": ("sheets",),
    "X3": ("recalc",),
    "X4": ("finance_colors",),
    "P2": ("contains",),
    "P4": ("contains",),
}


# For most keys the VALUE is the thing to compare against, so an empty one really is
# nothing to assert. `finance_colors` is a boolean opt-in, where `false` is an answer
# — "this is not a financial model" — and only an absent key means nobody said.
# Without this distinction an ordinary spreadsheet can never be a verified artifact:
# omitting the key is INERT (which L1 rejects) and setting it true reds every sheet
# that is not a financial model, which is the default-on problem coming back through
# the other door.
PRESENCE_IS_AN_ANSWER = {"finance_colors"}


def inert_reason(cid: str, expect: dict) -> str | None:
    keys = NEEDS_EXPECTATION.get(cid)
    if not keys:
        return None
    for k in keys:
        if (k in expect) if k in PRESENCE_IS_AN_ANSWER else expect.get(k):
            return None
    return f"nothing to assert — expectations carry none of: {', '.join(keys)}"


def readable_pdf(path: Path, expect: dict, tmpdir: str) -> tuple[Path | None, list[str]]:
    """An encrypted artifact, decrypted into `tmpdir` so the checks can read it.

    Added 2026-08-01 for P12: before this, handing run_checks an encrypted PDF blew
    up with an uncaught `ValueError: document closed or encrypted` from the first
    page access. A gate that CRASHES on an artifact type a skill legitimately
    produces reports nothing at all — which is worse than reporting a failure.

    A locked artifact with no password in `expect` is a finding, not a pass: nobody
    has checked it, and that is exactly what "not a verified artifact" means.

    These findings carry the `A0` id — they come from before any assertion runs, so
    they belong to no check, and every other finding in this file starts with the id
    of the check that raised it.
    """
    import fitz
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001 - a malformed file raises several types
        return None, [f"A0 cannot open {path.name} as a PDF: {type(e).__name__}: {e}"]
    with doc:
        if not doc.needs_pass:
            return path, []
        password = expect.get("password")
        if password is None:
            return None, [f"A0 {path.name} is encrypted and expectations carry no "
                          f"`password`, so not one assertion could run on it"]
        if not doc.authenticate(password):
            return None, [f"A0 the `password` in expectations was rejected by {path.name}"]
        plain = Path(tmpdir) / path.name
        doc.save(str(plain), encryption=fitz.PDF_ENCRYPT_NONE)
        return plain, []


def run_checks(path: Path, expect: dict, only: set[str] | None = None,
               allow_missing: frozenset[str] = frozenset()
               ) -> tuple[list[str], list[str], list[str]]:
    """Returns (findings, skipped, inert)."""
    kind = KIND_BY_SUFFIX.get(path.suffix.lower())
    if kind is None:
        return [f"unsupported artifact type: {path.name}"], [], []
    if kind == "pdf":
        with tempfile.TemporaryDirectory(prefix="l2-decrypt-") as td:
            usable, errors = readable_pdf(path, expect, td)
            if usable is None:
                return errors, [], []
            return _run_checks(usable, expect, kind, only, allow_missing)
    return _run_checks(path, expect, kind, only, allow_missing)


def _run_checks(path: Path, expect: dict, kind: str, only: set[str] | None,
                allow_missing: frozenset[str]) -> tuple[list[str], list[str], list[str]]:
    findings, skipped, inert = [], [], []
    for cid, c in CHECKS.items():
        if kind not in c["kind"] or (only and cid not in only):
            continue
        ok, why = tier_available(c["tier"])
        if not ok:
            if c["tier"] in REQUIRED_TIERS and c["tier"] not in allow_missing:
                findings.append(f"{cid} cannot run — {why}. This tier is required; "
                                f"pass --allow-missing {c['tier']} to run without it.")
                continue
            skipped.append(f"{cid} ({c['title']}) — {why}\n      "
                           f"residual coverage: {TIER_RESIDUAL.get(c['tier'], 'none')}")
            continue
        why_inert = inert_reason(cid, expect)
        if why_inert:
            inert.append(f"{cid} ({c['title']}) — {why_inert}")
        findings += c["fn"](path, expect)
    return findings, skipped, inert


# ══ fixtures ══════════════════════════════════════════════════════════════════
# Every flaw below is a defect one of the assertions above claims to catch. Building
# them here is the whole point: a positive sample only proves the checker can stay
# quiet.

CJK_TITLE = "季度经营分析报告"
CJK_BODY = "2026 年营收 Revenue 同比增长，毛利率保持稳定，费用结构持续优化。"
CJK_KEEP = "保留段落用于验证表格与正文的中文排版效果"


def build_docx(path: Path, flaw: str | None = None) -> dict:
    import docx as pydocx
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from lxml import etree

    doc = pydocx.Document()

    def cjk_run(par, text: str):
        run = par.add_run(text)
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), "宋体")
        return run

    # Centre the title so the paragraph actually carries a <w:pPr> — the element the
    # D4 order check is about only exists once a paragraph has properties.
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cjk_run(title, CJK_TITLE)
    cjk_run(doc.add_paragraph(), CJK_BODY)
    cjk_run(doc.add_paragraph(), CJK_KEEP)
    doc.add_paragraph().add_run("Latin-only run needs no eastAsia binding.")

    rev = doc.add_paragraph()
    fonts = ('<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="宋体"/></w:rPr>')
    rev._p.append(parse_xml(
        f'<w:ins {nsdecls("w")} w:id="101" w:author="ultrawork" '
        f'w:date="2026-08-01T00:00:00Z"><w:r>{fonts}'
        f'<w:t xml:space="preserve">新增的修订内容</w:t></w:r></w:ins>'))
    rev._p.append(parse_xml(
        f'<w:del {nsdecls("w")} w:id="102" w:author="ultrawork" '
        f'w:date="2026-08-01T00:00:00Z"><w:r>{fonts}'
        f'<w:delText xml:space="preserve">被删除的旧内容</w:delText></w:r></w:del>'))

    table = doc.add_table(rows=2, cols=2)
    for (r, c), text in {(0, 0): "科目", (0, 1): "金额",
                         (1, 0): "营业收入", (1, 1): "1200"}.items():
        cjk_run(table.cell(r, c).paragraphs[0], text)

    meta = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables),
            "contains": [CJK_TITLE, CJK_KEEP, "营业收入"]}
    doc.save(str(path))
    if flaw is None:
        return meta

    def edit(name: str, data: bytes):
        if name != "word/document.xml":
            return data
        if flaw == "missing-part":
            return None
        if flaw == "xml-broken":
            return data[: len(data) // 2]
        tree = etree.fromstring(data)
        body = tree.find(w("body"))
        if flaw == "dangling-rel":
            body.append(parse_xml(
                f'<w:p {nsdecls("w", "r")}><w:hyperlink r:id="rIdNope">'
                f'<w:r><w:t>dangling</w:t></w:r></w:hyperlink></w:p>'))
        elif flaw == "drop-paragraph":
            body.remove(next(p for p in body.iter(w("p"))
                             if CJK_KEEP in "".join(t.text or "" for t in p.iter(w("t")))))
        elif flaw == "order-ppr":
            par = next(p for p in tree.iter(w("p")) if p.find(w("pPr")) is not None)
            ppr = par.find(w("pPr"))
            par.remove(ppr)
            par.append(ppr)
        elif flaw == "order-rpr":
            run = next(r for r in tree.iter(w("r")) if r.find(w("rPr")) is not None)
            rpr = run.find(w("rPr"))
            run.remove(rpr)
            run.append(rpr)
        elif flaw == "sectpr-not-last":
            sect = body.find(w("sectPr"))
            body.remove(sect)
            body.insert(0, sect)
        elif flaw == "del-holds-wt":
            dt = next(tree.iter(w("delText")))
            dt.tag = w("t")
        elif flaw == "ins-holds-deltext":
            run = next(tree.iter(w("ins"))).find(w("r"))
            run.append(parse_xml(f'<w:delText {nsdecls("w")}>不该在这里</w:delText>'))
        elif flaw == "del-no-author":
            del next(tree.iter(w("del"))).attrib[w("author")]
        elif flaw == "cjk-no-eastasia":
            rf = next(r.find(w("rPr")).find(w("rFonts")) for r in tree.iter(w("r"))
                      if has_cjk(run_text(r)) and r.find(w("rPr")) is not None)
            del rf.attrib[w("eastAsia")]
        elif flaw == "cjk-no-eastasia-in-table":
            # The first CJK run is the title; breaking it proves nothing about whether
            # the walk reaches table cells. This one only lives inside a <w:tc>.
            cell_run = next(r for tc in tree.iter(w("tc")) for r in tc.iter(w("r"))
                            if has_cjk(run_text(r)))
            del cell_run.find(w("rPr")).find(w("rFonts")).attrib[w("eastAsia")]
        elif flaw == "cjk-no-eastasia-in-revision":
            ins_run = next(r for el in tree.iter(w("ins")) for r in el.iter(w("r"))
                           if has_cjk(run_text(r)))
            del ins_run.find(w("rPr")).find(w("rFonts")).attrib[w("eastAsia")]
        elif flaw == "order-tcpr":
            tc = next(t for t in tree.iter(w("tc")) if t.find(w("tcPr")) is not None)
            tcpr = tc.find(w("tcPr"))
            tc.remove(tcpr)
            tc.append(tcpr)
        elif flaw == "cjk-no-rfonts":
            run = next(r for r in tree.iter(w("r")) if has_cjk(run_text(r)))
            rpr = run.find(w("rPr"))
            rpr.remove(rpr.find(w("rFonts")))
        else:
            raise ValueError(f"unknown docx flaw {flaw!r}")
        return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    tmp = path.with_suffix(".good.docx")
    path.replace(tmp)
    rewrite_zip(tmp, path, edit)
    tmp.unlink()
    return meta


def build_xlsx(path: Path, flaw: str | None = None) -> dict:
    import openpyxl
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter

    blue, black, green = Font(color="FF0000FF"), Font(color="FF000000"), Font(color="FF008000")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "利润表"
    for col, head in enumerate(["科目", "预算", "实际", "差异"], start=1):
        ws.cell(row=1, column=col, value=head)
    rows = [("营业收入", 1000, 1200), ("营业成本", 600, 700)]
    for i, (label, budget, actual) in enumerate(rows, start=2):
        ws.cell(row=i, column=1, value=label)
        ws.cell(row=i, column=2, value=budget).font = blue
        ws.cell(row=i, column=3, value=actual).font = blue
        ws.cell(row=i, column=4, value=f"=C{i}-B{i}").font = black
    ws.cell(row=4, column=1, value="毛利")
    for col in (2, 3):
        c = get_column_letter(col)
        ws.cell(row=4, column=col, value=f"={c}2-{c}3").font = black
    ws.cell(row=4, column=4, value="=C4-B4").font = black

    sm = wb.create_sheet("汇总")
    sm["A1"] = "跨表汇总"
    sm["B1"] = "=利润表!D4"
    sm["B1"].font = green

    def autofit(sheet):
        widest: dict[int, int] = {}
        for row in sheet.iter_rows():
            for cell in row:
                v = cell.value
                if isinstance(v, str) and not v.startswith("="):
                    widest[cell.column] = max(widest.get(cell.column, 0),
                                              min(display_width(v) + 2, 60))
        for col, wd in widest.items():
            sheet.column_dimensions[get_column_letter(col)].width = max(wd, 10)

    autofit(ws)
    autofit(sm)

    if flaw == "err-cached-value":
        ws["F1"] = "#REF!"
    elif flaw == "err-in-formula":
        ws.cell(row=5, column=4, value="=#REF!*2").font = black
    elif flaw == "input-not-blue":
        ws["B2"].font = black
    elif flaw == "formula-not-black":
        ws["D2"].font = blue
    elif flaw == "link-not-green":
        sm["B1"].font = black
    elif flaw == "narrow-cjk-col":
        ws.column_dimensions["A"].width = 6
    elif flaw == "no-width-cjk-col":
        del ws.column_dimensions["A"]
    elif flaw == "merged-cjk-title":
        # NOT a flaw — the false positive this check shipped with. A title merged
        # across four columns is DISPLAYED across all four, so requiring column A
        # alone to fit it reds a workbook that renders in full (confirmed by
        # converting one to PDF and reading all 20 characters back out). Found
        # 2026-08-01 by the xlsx skill's autofit disagreeing with the gate.
        ws["A6"] = "二零二六年第三季度经营分析报告与附注说明"
        ws.merge_cells("A6:D6")
    elif flaw == "merged-cjk-vertical":
        # The control for the line above. A merge inside ONE column buys no
        # horizontal room, so the width still has to fit. Without this case,
        # "ignore merged cells" could have been implemented as "ignore every merge"
        # and looked exactly as green.
        ws["A6"] = "二零二六年第三季度经营分析报告与附注说明"
        ws.merge_cells("A6:A8")
    elif flaw == "wrong-value":
        ws["B2"] = 999
        ws["B2"].font = blue
    elif flaw == "recalc-drift":
        # B4 feeds D4 (=C4-B4) but is not in expect["sheets"], so ONLY X3 can see
        # this: every stored formula still reads back correctly, the recalculated
        # number does not. X3's only other control is "soffice is broken", which
        # says nothing about whether it can spot a wrong number.
        ws["B4"] = "=B2+B3"
    elif flaw is not None and flaw != "err-typed-cell":
        raise ValueError(f"unknown xlsx flaw {flaw!r}")

    wb.save(str(path))
    if flaw == "err-typed-cell":
        # openpyxl never writes cached values, so the data_only half of X2 would
        # otherwise never see a real error. This is what Excel actually stores when
        # a formula has blown up: <c t="e"><v>#DIV/0!</v></c>.
        def inject(name: str, data: bytes):
            if name != "xl/worksheets/sheet1.xml":
                return data
            text = data.decode("utf-8")
            return text.replace("</row>", '<c r="F1" t="e"><v>#DIV/0!</v></c></row>',
                                1).encode("utf-8")
        tmp = path.with_suffix(".good.xlsx")
        path.replace(tmp)
        rewrite_zip(tmp, path, inject)
        tmp.unlink()
    return {"sheets": {"利润表": {"B2": 1000, "C2": 1200, "D2": "=C2-B2"},
                       "汇总": {"B1": "=利润表!D4"}},
            "finance_colors": True,          # this fixture IS a financial model
            # D4 is =C4-B4 where B4 is =B2-B3 (400) and C4 is =C2-C3 (500), so 100.
            # This said 300 — the value of D2+D3 — from S1 until 2026-08-01, and it
            # went unnoticed because the machine that could disprove it had no
            # LibreOffice: the assertion was SKIPPED, and skipped reads as green at a
            # glance. The first real run of X3 caught it. This file prints "skipped is
            # not green"; it turns out that was literal.
            "recalc": {"D2": 200, "D4": 100}}


PDF_CJK = "季度经营分析报告与中文排版验证内容"
PDF_LATIN = "Ultrawork office skill artifact"


def build_pdf(path: Path, flaw: str | None = None) -> dict:
    import fitz
    # A /Rotate page is built by producing the ordinary content and rotating at the
    # end, so "rotated" and "rotated-tofu" differ from their unrotated twins in
    # exactly one property. Anything else and a finding could not be attributed to
    # the rotation.
    rotate = flaw in ("rotated", "rotated-tofu")
    flaw = {"rotated": None, "rotated-tofu": "tofu"}.get(flaw, flaw)
    doc = fitz.open()
    page = doc.new_page(width=420, height=260)

    if flaw != "blank":
        page.insert_text((40, 60), PDF_LATIN, fontsize=13)
        cjk = "????????" if flaw == "cjk-mojibake" else PDF_CJK
        font = "helv" if flaw == "cjk-mojibake" else "china-s"
        if flaw == "tofu":
            # Text layer intact, glyphs missing: exactly what a PDF written with a
            # font that has no CJK coverage looks like. Draw the .notdef boxes on the
            # character cells and hide the real text behind them.
            page.insert_text((40, 110), PDF_CJK, fontname="china-s", fontsize=16)
            boxes = [fitz.Rect(c["bbox"]) for b in page.get_text("rawdict")["blocks"]
                     for l in b["lines"] for sp in l["spans"] for c in sp["chars"]
                     if has_cjk(c["c"])]
            doc.close()
            doc = fitz.open()
            page = doc.new_page(width=420, height=260)
            page.insert_text((40, 60), PDF_LATIN, fontsize=13)
            page.insert_text((40, 110), PDF_CJK, fontname="china-s", fontsize=16,
                             render_mode=3)
            for r in boxes:
                page.draw_rect(fitz.Rect(r.x0 + 0.6, r.y0 + 0.6, r.x1 - 0.6, r.y1 - 0.6),
                               color=(0, 0, 0), width=0.8)
        elif flaw != "missing-text":
            page.insert_text((40, 110), cjk, fontname=font, fontsize=16)

    if flaw != "blank":
        wd = fitz.Widget()
        wd.field_name = "applicant"
        wd.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        wd.rect = fitz.Rect(40, 170, 200, 192)
        wd.field_value = ("This applicant name is far too long to fit in the printed box"
                          if flaw == "field-overflow" else "Zhang")
        wd.text_fontsize = 11
        page.add_widget(wd)

    if rotate:
        for pg in doc:
            pg.set_rotation(90)
    doc.save(str(path))
    doc.close()
    if flaw not in (None, "blank", "missing-text", "cjk-mojibake", "tofu",
                    "field-overflow"):
        raise ValueError(f"unknown pdf flaw {flaw!r}")
    return {"contains": [PDF_LATIN, PDF_CJK]}


# ── fidelity fixtures ─────────────────────────────────────────────────────────
# The negative controls here are not invented damage: they are what the obvious
# implementation actually does. "Load it with the library, save it back" is the
# shape every one of these skills will reach for first, and it is what silently
# drops xl/metadata.xml.

FIXTURES = REPO / "packages" / "knowledge" / "sidecar" / "src" / "__fixtures__"

OPAQUE_METADATA = (b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                   b'<metadata xmlns="http://schemas.openxmlformats.org/spreadsheetml'
                   b'/2006/main"><metadataTypes count="1"><metadataType name="XLDAPR"'
                   b' minSupportedVersion="120000"/></metadataTypes></metadata>')


def build_rich_xlsx(path: Path) -> None:
    """A workbook carrying the features an edit must not quietly discard."""
    import openpyxl
    from openpyxl.chart import BarChart, Reference
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "明细"
    for i in range(1, 6):
        ws.cell(i, 1, f"项目{i}")
        ws.cell(i, 2, i * 10)
    chart = BarChart()
    chart.add_data(Reference(ws, min_col=2, min_row=1, max_row=5))
    ws.add_chart(chart, "E2")
    ws.conditional_formatting.add("B1:B5", CellIsRule(
        operator="greaterThan", formula=["20"],
        fill=PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")))
    ws.add_data_validation(DataValidation(type="list", formula1='"甲,乙,丙"', sqref="C1:C5"))
    ws.freeze_panes = "B2"
    ws.auto_filter.ref = "A1:B5"
    ws.merge_cells("A7:B7")
    ws.column_dimensions["A"].width = 14
    summary = wb.create_sheet("汇总")
    summary["A1"] = "汇总页"
    summary.column_dimensions["A"].width = 12   # X5: CJK columns need an explicit width
    wb.save(str(path))

    # Two parts openpyxl has no model for. Real workbooks are full of these
    # (rich data, custom XML, pivot caches); dropping them is invisible until the
    # user reopens the file in Excel and something is gone.
    def declare(name: str, data: bytes):
        if name != "[Content_Types].xml":
            return data
        return data.decode("utf-8").replace(
            "</Types>",
            '<Override PartName="/xl/metadata.xml" ContentType="application/vnd.'
            'openxmlformats-officedocument.spreadsheetml.sheetMetadata+xml"/>'
            '<Override PartName="/customXml/item1.xml" ContentType="application/xml"/>'
            "</Types>").encode("utf-8")

    tmp = path.with_suffix(".plain.xlsx")
    path.replace(tmp)
    rewrite_zip(tmp, path, declare)
    tmp.unlink()
    with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as z:
        z.writestr("xl/metadata.xml", OPAQUE_METADATA)
        z.writestr("customXml/item1.xml",
                   b'<?xml version="1.0"?><props><p n="owner">finance</p></props>')


def build_pdf_pages(path: Path, pages: int = 3) -> None:
    import fitz
    doc = fitz.open()
    for i in range(1, pages + 1):
        page = doc.new_page(width=380, height=240)
        page.insert_text((40, 60), f"Page {i} of {pages}", fontsize=14)
        page.insert_text((40, 100), f"第 {i} 页正文内容", fontname="china-s", fontsize=15)
    doc.save(str(path))
    doc.close()


def byte_edit(src: Path, dst: Path, part: str, old: bytes, new: bytes) -> None:
    """A preserving edit: rewrite one part's bytes, copy every other part verbatim.
    This is the positive control — it is what a fidelity-respecting skill does."""
    hit = False

    def mutate(name: str, data: bytes):
        nonlocal hit
        if name != part or old not in data:
            return data
        hit = True
        return data.replace(old, new, 1)

    rewrite_zip(src, dst, mutate)
    if not hit:
        raise ValueError(f"byte_edit found no {old!r} in {part}")


# Each returns (baseline, edited, expectations).
def fid_xlsx_preserving(root: Path) -> tuple[Path, Path, dict]:
    base, out = root / "rich.xlsx", root / "rich-edited.xlsx"
    build_rich_xlsx(base)
    byte_edit(base, out, "xl/worksheets/sheet1.xml", b"<v>10</v>", b"<v>99</v>")
    return base, out, {"baseline": str(base)}


def fid_xlsx_naive(root: Path) -> tuple[Path, Path, dict]:
    import openpyxl
    base, out = root / "naive.xlsx", root / "naive-edited.xlsx"
    build_rich_xlsx(base)
    wb = openpyxl.load_workbook(base)
    with closing(wb):
        wb["明细"]["A1"] = "改过了"
        wb.save(str(out))
    return base, out, {"baseline": str(base)}


def fid_xlsx_drop_sheet(root: Path) -> tuple[Path, Path, dict]:
    import openpyxl
    base, out = root / "ds.xlsx", root / "ds-edited.xlsx"
    build_rich_xlsx(base)
    wb = openpyxl.load_workbook(base)
    with closing(wb):
        del wb["汇总"]
        wb.save(str(out))
    return base, out, {"baseline": str(base)}


def fid_xlsx_kill_features(root: Path) -> tuple[Path, Path, dict]:
    import openpyxl
    base, out = root / "kf.xlsx", root / "kf-edited.xlsx"
    build_rich_xlsx(base)
    wb = openpyxl.load_workbook(base)
    with closing(wb):
        ws = wb["明细"]
        ws.freeze_panes = None
        ws.auto_filter.ref = None
        ws.conditional_formatting._cf_rules.clear()
        wb.save(str(out))
    return base, out, {"baseline": str(base)}


def fid_docx_preserving(root: Path) -> tuple[Path, Path, dict]:
    base, out = root / "f.docx", root / "f-edited.docx"
    build_docx(base, None)
    byte_edit(base, out, "word/document.xml",
              CJK_TITLE.encode("utf-8"), "年度经营分析报告".encode("utf-8"))
    return base, out, {"baseline": str(base)}


def fid_docx_drop_styles(root: Path) -> tuple[Path, Path, dict]:
    base, out = root / "fs.docx", root / "fs-edited.docx"
    build_docx(base, None)
    rewrite_zip(base, out, lambda n, d: None if n == "word/styles.xml" else d)
    return base, out, {"baseline": str(base)}


def fid_docx_drop_relationships(root: Path) -> tuple[Path, Path, dict]:
    from lxml import etree
    base, out = root / "fr.docx", root / "fr-edited.docx"
    build_docx(base, None)

    def strip(name: str, data: bytes):
        if name != "word/_rels/document.xml.rels":
            return data
        tree = etree.fromstring(data)
        for child in list(tree)[:1]:
            tree.remove(child)
        return etree.tostring(tree, xml_declaration=True, encoding="UTF-8",
                              standalone=True)

    rewrite_zip(base, out, strip)
    return base, out, {"baseline": str(base)}


def fid_pdf_preserving(root: Path) -> tuple[Path, Path, dict]:
    import fitz
    base, out = root / "m.pdf", root / "m-edited.pdf"
    build_pdf_pages(base)
    doc = fitz.open(base)
    doc[0].insert_text((40, 140), "Added on page one", fontsize=12)
    doc.save(str(out))
    doc.close()
    return base, out, {"baseline": str(base), "touched_pages": [1]}


def fid_pdf_drop_page(root: Path) -> tuple[Path, Path, dict]:
    import fitz
    base, out = root / "mp.pdf", root / "mp-edited.pdf"
    build_pdf_pages(base)
    doc = fitz.open(base)
    doc.delete_page(1)
    doc.save(str(out))
    doc.close()
    return base, out, {"baseline": str(base), "touched_pages": [1]}


def fid_pdf_touch_untouched(root: Path) -> tuple[Path, Path, dict]:
    import fitz
    base, out = root / "mt.pdf", root / "mt-edited.pdf"
    build_pdf_pages(base)
    doc = fitz.open(base)
    doc[1].insert_text((40, 160), "silently rewritten", fontsize=12)
    doc.save(str(out))
    doc.close()
    return base, out, {"baseline": str(base), "touched_pages": [1]}


def _real_roundtrip(root: Path, name: str, loader) -> tuple[Path, Path, dict] | None:
    src = FIXTURES / name
    if not src.is_file():
        return None
    base = root / name
    shutil.copyfile(src, base)
    out = root / f"rt-{name}"
    loader(base, out)
    return base, out, {"baseline": str(base)}


def fid_real_docx(root: Path):
    import docx as pydocx
    return _real_roundtrip(root, "sample.docx",
                           lambda b, o: pydocx.Document(str(b)).save(str(o)))


def fid_real_xlsx(root: Path):
    import openpyxl

    def rt(b: Path, o: Path) -> None:
        wb = openpyxl.load_workbook(b)
        with closing(wb):
            wb.save(str(o))

    return _real_roundtrip(root, "sample.xlsx", rt)


# (case name, builder, marker, should_fire)
def fid_pdf_merge(root: Path) -> tuple[Path, Path, dict]:
    """A merge that keeps everything. `baseline` is a LIST — the extension P11 needed."""
    import fitz
    a, b, out = root / "ma.pdf", root / "mb.pdf", root / "merged.pdf"
    build_pdf_pages(a, pages=2)
    build_pdf_pages(b, pages=3)
    doc = fitz.open(a)
    with doc:
        with fitz.open(b) as second:
            doc.insert_pdf(second)
        doc.save(str(out))
    return a, out, {"baseline": [str(a), str(b)]}


def fid_pdf_merge_drops_input(root: Path) -> tuple[Path, Path, dict]:
    """The merge that forgot its second input.

    The result opens, renders, extracts text and passes every other assertion here.
    Only walking the declared inputs in order shows that half the document is gone.
    """
    import fitz
    a, b, out = root / "da.pdf", root / "db.pdf", root / "half.pdf"
    build_pdf_pages(a, pages=2)
    build_pdf_pages(b, pages=3)
    doc = fitz.open(a)
    with doc:
        doc.save(str(out))
    return a, out, {"baseline": [str(a), str(b)]}


def fid_pdf_merge_wrong_order(root: Path) -> tuple[Path, Path, dict]:
    """All the pages are present — in the wrong order. Page COUNT cannot see this."""
    import fitz
    a, b, out = root / "oa.pdf", root / "ob.pdf", root / "swapped.pdf"
    build_pdf_pages(a, pages=2)
    build_pdf_pages(b, pages=3)
    doc = fitz.open(b)
    with doc:
        with fitz.open(a) as first:
            doc.insert_pdf(first)
        doc.save(str(out))
    return a, out, {"baseline": [str(a), str(b)]}


def fid_pdf_encrypted(root: Path) -> tuple[Path, Path, dict]:
    """An encrypted artifact WITH the password: every assertion must run normally."""
    import fitz
    base, out = root / "ea.pdf", root / "enc.pdf"
    build_pdf_pages(base, pages=2)
    doc = fitz.open(base)
    with doc:
        doc.save(str(out), encryption=fitz.PDF_ENCRYPT_AES_256,
                 user_pw="s3cret", owner_pw="s3cret-owner")
    return base, out, {"baseline": str(base), "password": "s3cret",
                       "contains": ["Page 1 of 2"]}


def fid_pdf_encrypted_no_password(root: Path) -> tuple[Path, Path, dict]:
    """The same artifact with no password in expectations.

    Before this path existed the gate raised an uncaught ValueError from the first
    page access — a crash reports nothing, which is worse than a failure.
    """
    import fitz
    base, out = root / "eb.pdf", root / "enc2.pdf"
    build_pdf_pages(base, pages=2)
    doc = fitz.open(base)
    with doc:
        doc.save(str(out), encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="s3cret")
    return base, out, {"baseline": str(base), "contains": ["Page 1 of 2"]}


FIDELITY_CASES = [
    ("preserving xlsx edit keeps every feature", fid_xlsx_preserving, "", False),
    ("F1 library round-trip drops parts it has no model for", fid_xlsx_naive, "F1", True),
    ("F3 a sheet disappears", fid_xlsx_drop_sheet, "F3", True),
    ("F3 frozen panes, auto filter and conditional formats wiped",
     fid_xlsx_kill_features, "F3", True),
    ("preserving docx edit keeps every part", fid_docx_preserving, "", False),
    ("F2 word/styles.xml dropped", fid_docx_drop_styles, "F2", True),
    ("F2 a document relationship dropped", fid_docx_drop_relationships, "F2", True),
    ("preserving pdf edit leaves other pages alone", fid_pdf_preserving, "", False),
    ("P5 a page disappears", fid_pdf_drop_page, "P5", True),
    ("P5 an undeclared page is rewritten", fid_pdf_touch_untouched, "P5", True),
    # Real-file regressions. These two are the pair that motivated the emptiness
    # test in part_is_inert: same symptom, opposite verdicts.
    ("real sample.docx: empty rels omitted is NOT data loss", fid_real_docx, "", False),
    ("real sample.xlsx: xl/metadata.xml lost IS data loss", fid_real_xlsx, "F1", True),
    # P11 needed `baseline` to accept several inputs; these are what make that real.
    ("merge keeping every input stays silent", fid_pdf_merge, "", False),
    ("P5 a merge silently drops an input", fid_pdf_merge_drops_input, "P5", True),
    ("P5 a merge keeps every page but reorders them", fid_pdf_merge_wrong_order,
     "P5", True),
    # P12 needed encrypted artifacts to be checkable at all.
    ("an encrypted artifact with its password checks normally", fid_pdf_encrypted,
     "", False),
    ("an encrypted artifact with no password is reported, not crashed on",
     fid_pdf_encrypted_no_password, "A0", True),
]


BUILDERS = {"docx": (build_docx, ".docx"), "xlsx": (build_xlsx, ".xlsx"),
            "pdf": (build_pdf, ".pdf")}

# (case name, kind, flaw, marker, should_fire). marker "" on a positive control means
# "no finding at all"; on a negative control it is the check id that must speak up.
CASES: list[tuple[str, str, str | None, str, bool]] = [
    ("docx clean artifact stays silent", "docx", None, "", False),
    ("D1 truncated document.xml", "docx", "xml-broken", "D1", True),
    ("D1 document.xml removed", "docx", "missing-part", "D1", True),
    ("D1 r:id with no relationship", "docx", "dangling-rel", "D1", True),
    ("D3 paragraph count drifts", "docx", "drop-paragraph", "D3", True),
    ("D4 w:pPr moved off the front", "docx", "order-ppr", "D4", True),
    ("D4 w:rPr moved off the front", "docx", "order-rpr", "D4", True),
    ("D4 w:sectPr not last in body", "docx", "sectpr-not-last", "D4", True),
    ("D5 w:t inside w:del", "docx", "del-holds-wt", "D5", True),
    ("D5 w:delText inside w:ins", "docx", "ins-holds-deltext", "D5", True),
    ("D5 revision without w:author", "docx", "del-no-author", "D5", True),
    ("D6 CJK run loses @eastAsia", "docx", "cjk-no-eastasia", "D6", True),
    ("D6 CJK run loses w:rFonts", "docx", "cjk-no-rfonts", "D6", True),
    ("D6 reaches CJK runs inside a table cell", "docx",
     "cjk-no-eastasia-in-table", "D6", True),
    ("D6 reaches CJK runs inside a w:ins revision", "docx",
     "cjk-no-eastasia-in-revision", "D6", True),
    ("D4 w:tcPr moved off the front", "docx", "order-tcpr", "D4", True),

    ("xlsx clean artifact stays silent", "xlsx", None, "", False),
    ("X1 cell value drifts from expectation", "xlsx", "wrong-value", "X1", True),
    ("X2 cached #REF! value", "xlsx", "err-cached-value", "X2", True),
    ("X2 #REF! baked into a formula", "xlsx", "err-in-formula", "X2", True),
    ("X2 error-typed cell with a cached #DIV/0!", "xlsx", "err-typed-cell", "X2", True),
    ("X3 a recalculated value drifts from the expectation", "xlsx", "recalc-drift",
     "X3", True),
    ("X4 hard-coded input not blue", "xlsx", "input-not-blue", "X4", True),
    ("X4 formula not black", "xlsx", "formula-not-black", "X4", True),
    ("X4 cross-sheet link not green", "xlsx", "link-not-green", "X4", True),
    ("X5 CJK column too narrow", "xlsx", "narrow-cjk-col", "X5", True),
    ("X5 CJK column left at default width", "xlsx", "no-width-cjk-col", "X5", True),
    # The pair that pins X5's merge handling, same shape as the rotated-page pair
    # below: a fix that buys silence by making the check blind is not a fix.
    ("X5 CJK title merged across columns stays silent", "xlsx", "merged-cjk-title",
     "", False),
    ("X5 CJK cell merged within one column still fires", "xlsx",
     "merged-cjk-vertical", "X5", True),

    ("pdf clean artifact stays silent", "pdf", None, "", False),
    ("P1 page renders blank", "pdf", "blank", "P1", True),
    ("P2 written text is not extractable", "pdf", "missing-text", "P2", True),
    ("P3 field value overflows its box", "pdf", "field-overflow", "P3", True),
    ("P4 CJK comes back as mojibake", "pdf", "cjk-mojibake", "P4", True),
    ("P4 CJK renders as tofu boxes", "pdf", "tofu", "P4", True),
    # The pair that pins cjk_center_ink_fraction's rotation handling. Before the fix
    # the first of these failed — a correctly rendered rotated page scored 0.04
    # because the character boxes were read in page space and the raster is in
    # display space. The second is the control: the fix must not have bought silence
    # by making the check blind on rotated pages.
    ("P4 rotated page with real glyphs stays silent", "pdf", "rotated", "", False),
    ("P4 rotated page with tofu boxes still fires", "pdf", "rotated-tofu", "P4", True),
]


def fidelity_cases(root: Path) -> list[dict]:
    """Edit-fidelity: did the edit keep everything it was not asked to change."""
    results = []
    for i, (name, build, marker, should_fire) in enumerate(FIDELITY_CASES):
        work = root / f"fid{i:02d}"
        work.mkdir()
        made = build(work)
        if made is None:  # a real-file regression whose fixture is not in the tree
            results.append({"case": name, "kind": "fidelity", "flaw": None,
                            "marker": marker, "ok": True, "skipped_missing": True,
                            "detail": "fixture not present in this checkout",
                            "findings": []})
            continue
        base, out, expect = made
        findings, _, _ = run_checks(out, expect, allow_missing=ALL_TIERS)
        fired = [f for f in findings if f.startswith(marker)] if marker else []
        ok = bool(fired) if should_fire else not findings
        results.append({"case": name, "kind": "fidelity", "flaw": base.name,
                        "marker": marker, "ok": ok,
                        "detail": "" if ok else ("no finding carried this id"
                                                 if should_fire
                                                 else "; ".join(findings[:4])),
                        "findings": findings})
    return results


def inert_cases(root: Path) -> list[dict]:
    """"Nobody said" and "the answer is no" must not look the same.

    L1 rejects an artifact whose non-fidelity assertions were all INERT, so if a
    declared `finance_colors: false` still counted as INERT, no ordinary spreadsheet
    could ever be a verified artifact — the only way through would be to claim every
    workbook is a financial model. The third case is the control: without it,
    "false silences X4" would be indistinguishable from "X4 is silent anyway".
    """
    art = root / "inert-probe.xlsx"
    expect = build_xlsx(art, None)
    plain = {k: v for k, v in expect.items() if k != "finance_colors"}
    # Same workbook, colours deliberately wrong for the convention.
    off = root / "inert-probe-offcolour.xlsx"
    build_xlsx(off, "input-not-blue")

    def case(name: str, ok: bool, detail: str = "") -> dict:
        return {"case": name, "kind": "xlsx", "flaw": "inert-semantics", "marker": "",
                "arrow": "expectation semantics", "ok": ok, "detail": detail,
                "findings": []}

    absent = inert_reason("X4", plain)
    declared_false = inert_reason("X4", {**plain, "finance_colors": False})
    empty_list = inert_reason("P2", {"contains": []})
    fires = CHECKS["X4"]["fn"](off, {"finance_colors": True})
    silent = CHECKS["X4"]["fn"](off, {"finance_colors": False})
    return [
        case("X4 with no finance_colors key at all is INERT, not a pass",
             bool(absent), "an unstated opt-in read as an executed assertion"),
        case("X4 with finance_colors:false is an answer, not an absence",
             declared_false is None, f"still INERT: {declared_false}"),
        case("CONTROL: finance_colors:true on the same file does fire",
             bool(fires) and not silent,
             "opting out is indistinguishable from the check never having teeth"),
        case("a value-bearing key that is empty is still INERT",
             bool(empty_list), "an empty `contains` was accepted as an expectation"),
    ]


def tier_fault_cases(root: Path) -> list[dict]:
    """A skipped tier is honest. A tier that runs and finds nothing because the tool
    beneath it is broken is not — that is the shape of a guard that always passes.

    Force SOFFICE to a binary that cannot work and assert the LibreOffice checks
    report the failure instead of returning clean. This exercises the degradation
    path on machines that have no LibreOffice at all, which is where the risk lives.
    """
    global SOFFICE
    saved, results = SOFFICE, []
    SOFFICE = str(root / "no-such-soffice-binary")
    try:
        for cid, kind in (("D7", "docx"), ("X3", "xlsx")):
            build, suffix = BUILDERS[kind]
            art = root / f"fault-{cid}{suffix}"
            expect = build(art, None)
            findings = CHECKS[cid]["fn"](art, expect)
            results.append({
                "case": f"{cid} reports a broken LibreOffice instead of passing quietly",
                "kind": kind, "flaw": "soffice-binary-missing", "marker": cid,
                "ok": any(f.startswith(cid) for f in findings),
                "detail": "" if findings else "returned no findings at all",
                "findings": findings})
    finally:
        SOFFICE = saved
    return results


def selftest(as_json: bool, only_kind: str | None) -> int:
    results, skipped_all = [], []
    with tempfile.TemporaryDirectory(prefix="office-selftest-") as td:
        root = Path(td)
        for i, (name, kind, flaw, marker, should_fire) in enumerate(CASES):
            if only_kind and kind != only_kind:
                continue
            # A NEGATIVE control whose check needs a tier this host does not have
            # cannot fire, and calling that a failure makes the whole suite red for a
            # reason that has nothing to do with the code under test. It is reported
            # as skipped and named — never folded into the pass count.
            #
            # This is only honest because the tier IS exercised somewhere: CI installs
            # LibreOffice on Linux (059 §7 option B). If that ever stops being true,
            # this stops being a skip and becomes a control nobody runs — which is
            # exactly how X3 carried a wrong expectation for a month.
            tier = CHECKS.get(marker, {}).get("tier", "core") if marker else "core"
            available, why = tier_available(tier)
            if should_fire and not available:
                results.append({
                    "case": name, "kind": kind, "flaw": flaw, "marker": marker,
                    "ok": True, "skipped_missing": True,
                    "skip_label": f"needs the {tier} tier",
                    "detail": f"{why} — this negative control cannot fire without it; "
                              f"CI runs it on Linux, where LibreOffice is installed",
                    "findings": []})
                continue
            build, suffix = BUILDERS[kind]
            art = root / f"case{i:02d}{suffix}"
            expect = build(art, flaw)
            findings, skipped, _ = run_checks(art, expect, allow_missing=ALL_TIERS)
            skipped_all += skipped
            fired = [f for f in findings if f.startswith(marker)] if marker else []
            if should_fire:
                ok = bool(fired)
                detail = "" if ok else "no finding carried this id"
            else:
                ok = not findings
                detail = "" if ok else "; ".join(findings[:4])
            results.append({"case": name, "kind": kind, "flaw": flaw, "marker": marker,
                            "ok": ok, "detail": detail, "findings": findings})

        if not only_kind:
            results += fidelity_cases(root)
            results += inert_cases(root)
            results += tier_fault_cases(root)

    failed = [r for r in results if not r["ok"]]
    seen: list[str] = []
    for s in skipped_all:
        if s not in seen:
            seen.append(s)

    if as_json:
        print(json.dumps({"results": results, "skipped": seen,
                          "passed": len(results) - len(failed), "failed": len(failed)},
                         ensure_ascii=False, indent=2))
    else:
        for r in results:
            if r.get("skipped_missing"):
                print(f"SKIP  [{r.get('skip_label', 'fixture absent')}] {r['case']} "
                      f"— {r['detail']}")
                continue
            verdict = "PASS" if r["ok"] else "FAIL"
            arrow = r.get("arrow") or ("must fire" if r["marker"] else "must stay silent")
            print(f"{verdict}  [{arrow}] {r['case']}")
            if not r["ok"]:
                print(f"      {r['detail']}")
                for f in r["findings"][:6]:
                    print(f"      · {f}")
        if seen:
            print("\n[skipped] assertions this machine cannot exercise:")
            for s in seen:
                print(f"  - {s}")
            print("  Skipped is not green. Install LibreOffice / point $ECMA376_XSD_DIR\n"
                  "  at unpacked ECMA-376 schemas to close these.")
        n_skipped = sum(1 for r in results if r.get("skipped_missing"))
        print(f"\n[office-selftest] {len(results) - len(failed) - n_skipped} passed, "
              f"{len(failed)} failed"
              + (f", {n_skipped} case(s) skipped" if n_skipped else "")
              + (f", {len(seen)} assertion(s) skipped" if seen else ""))
    return 1 if failed else 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--check", nargs="+", metavar="FILE",
                    help="gate real artifacts instead of running the selftest")
    ap.add_argument("--expect", metavar="JSON", help="expectations file for --check")
    ap.add_argument("--only", metavar="ID", nargs="+", help="restrict to these check ids")
    ap.add_argument("--kind", choices=sorted(BUILDERS), help="selftest: one format only")
    ap.add_argument("--list", action="store_true", help="list the assertions and exit")
    ap.add_argument("--allow-missing", nargs="+", default=[], metavar="TIER",
                    choices=list(REQUIRED_TIERS),
                    help="run without a required tier (dev machines); named in the report")
    ap.add_argument("--json", action="store_true")
    args = ap.parse_args()

    if args.list:
        for cid, c in CHECKS.items():
            ok, why = tier_available(c["tier"])
            state = "available" if ok else f"UNAVAILABLE — {why}"
            print(f"{cid:3} [{'+'.join(c['kind']):9}/{c['tier']:7}] {c['title']}"
                  f"\n     {state}")
        return 0

    if not args.check:
        return selftest(args.json, args.kind)

    expect = json.loads(Path(args.expect).read_text(encoding="utf-8")) if args.expect else {}
    only = set(args.only) if args.only else None
    report, bad = {}, 0
    for name in args.check:
        p = Path(name)
        if not p.is_file():
            report[name] = {"findings": [f"file not found: {name}"], "skipped": [],
                            "inert": []}
            bad += 1
            continue
        per_file = expect.get(p.name, expect) if isinstance(expect, dict) else {}
        findings, skipped, inert = run_checks(p, per_file, only,
                                              frozenset(args.allow_missing))
        report[name] = {"findings": findings, "skipped": skipped, "inert": inert}
        bad += len(findings)

    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        for name, r in report.items():
            print(f"[{'FAIL' if r['findings'] else 'OK'}] {name}")
            for f in r["findings"]:
                print(f"    {f}")
            for s in r["skipped"]:
                print(f"    SKIPPED {s}")
            for s in r["inert"]:
                print(f"    INERT   {s}")
        inert_total = sum(len(r["inert"]) for r in report.values())
        print(f"[office-selftest] {bad} finding(s) across {len(report)} artifact(s)"
              + (f"; {inert_total} assertion(s) had nothing to check — an artifact with "
                 f"no expectations is not a verified artifact" if inert_total else ""))
    return 1 if bad else 0


if __name__ == "__main__":
    sys.exit(main())
