#!/usr/bin/env python3
"""Read a .docx: dump paragraphs and tables. ultrawork doc-edit skill."""
from __future__ import annotations
import argparse, json, sys

def _require():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError:
        print("Missing dependency: python-docx (pip install python-docx)", file=sys.stderr)
        raise SystemExit(1)

def main(argv):
    ap = argparse.ArgumentParser(description="Read text/tables from a .docx")
    ap.add_argument("file")
    ap.add_argument("--json", action="store_true")
    args = ap.parse_args(argv)
    docx = _require()
    try:
        doc = docx.Document(args.file)
    except Exception as exc:  # noqa: BLE001
        print(f"Error opening {args.file}: {exc}", file=sys.stderr)
        return 1
    paragraphs = [p.text for p in doc.paragraphs]
    tables = [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]
    if args.json:
        print(json.dumps({"paragraphs": paragraphs, "tables": tables}, ensure_ascii=False, indent=2))
    else:
        for p in paragraphs:
            if p.strip():
                print(p)
        for i, t in enumerate(tables):
            print(f"\n--- table {i} ---")
            for row in t:
                print(" | ".join(row))
    return 0

if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
