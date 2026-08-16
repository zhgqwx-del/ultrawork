#!/usr/bin/env python3
"""Capability-matrix gate for the office skills (discussions/059 §3.4 / §5 L1).

"More complete than the references" is not a judgement call — §3 of 059 pins it to
three countable rules over a matrix of 49 capabilities. This script is what keeps
that matrix honest: it parses the tables straight out of the markdown, checks the
three rules, and then makes each skill PROVE its claims by running the entry point
it declares for every capability. A row can't just be ticked; something has to run.

    python3 scripts/check-office-skill-capabilities.py
    python3 scripts/check-office-skill-capabilities.py --json

Each skill declares what it implements in <skill>/capabilities.json:

    {"skill": "pdf",
     "capabilities": {"P1": {"entry": "scripts/pdf_render.py", "probe": ["--help"]}}}

Skills with no capabilities.json are reported as undeclared, not failed — the older
skills predate this contract and migrate as they get rewritten.

Exit 0 = all gates pass, 1 = a gate failed.
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path

# ── stdout must be UTF-8 on every platform, and on Windows it is not ──────────
# This gate prints ✅/❌ and Chinese. Windows encodes a CAPTURED stdout in the
# machine's ANSI code page and Python only defaults to UTF-8 from 3.15 (PEP 686);
# CI pins 3.11. Measured on CI: this script died with
# `UnicodeEncodeError: 'charmap' codec can't encode character '\u2705'` inside its
# own `print(json.dumps(...))`. The skills were fixed for this first and the GATES
# were missed — the same defect has two homes, and only one of them was product.
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        try:
            _stream.reconfigure(encoding="utf-8")
        except (ValueError, OSError):        # already detached / not reconfigurable
            pass


REPO = Path(__file__).resolve().parent.parent
DOC = REPO / "docs" / "discussions" / "059-office-skills-revamp.md"
BUILTIN = REPO / "skills" / "builtin"
PY = sys.executable  # never hardcode python3 vs python (Windows)

MIN_BEYOND = 4   # §3.4 C3: capabilities no reference has at all
YES, PARTIAL = "✅", "△"

# Column layout of the §3 tables. "current" and "target" are ours; the rest are the
# reference apps we must out-cover.
CURRENT_COL, TARGET_COL = "当前", "目标"
ROW_RE = re.compile(r"^\|\s*([PWX]\d+)\s*\|(.+)\|\s*$")

# Capability IDs are namespaced by skill: P* -> pdf, W* -> docx, X* -> xlsx.
SKILL_PREFIX = {"pdf": "P", "docx": "W", "xlsx": "X"}
BEYOND_RE = re.compile(r"超越项")


def parse_matrix(doc: Path) -> tuple[list[dict], list[str]]:
    """Pull the capability rows out of the markdown tables in §3."""
    if not doc.is_file():
        raise SystemExit(f"[error] capability matrix not found: {doc}")
    rows: list[dict] = []
    columns: list[str] = []
    for line in doc.read_text(encoding="utf-8").splitlines():
        if line.startswith("| #") and "能力" in line:
            columns = [c.strip().strip("*") for c in line.strip("|").split("|")]
            continue
        m = ROW_RE.match(line)
        if not m or not columns:
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if len(cells) != len(columns):
            continue
        row = dict(zip(columns, cells))
        row["id"] = m.group(1)
        row["beyond"] = bool(BEYOND_RE.search(cells[-1]))
        rows.append(row)
    return rows, columns


def mark(cell: str) -> str:
    """Reduce a matrix cell to its verdict symbol."""
    if YES in cell:
        return YES
    if PARTIAL in cell:
        return PARTIAL
    return "❌"


RANK = {"❌": 0, PARTIAL: 1, YES: 2}


# The L2 artifact gate, loaded as a module (the filename has hyphens, so no import).
# C4 used to accept `--help` exiting 0 as proof a capability existed. It is not:
# a script can start, do nothing, and still exit 0 — "ticked on paper, absent in
# code" one level down. A capability now proves itself by producing a real artifact
# that L2 then has to accept.
# Mirrors the L2 gate's required tiers so --allow-missing means the same thing in
# both places; a dev machine that can skip soffice for L2 can skip it for C4 too.
ALL_TIERS = frozenset({"soffice", "xsd"})


def load_l2():
    import importlib.util
    path = REPO / "scripts" / "office-skills-selftest.py"
    if not path.is_file():
        return None, f"L2 gate not found at {path}"
    try:
        spec = importlib.util.spec_from_file_location("office_selftest", path)
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        return mod, ""
    except Exception as e:  # noqa: BLE001 - a missing python-docx lands here
        return None, f"L2 gate unavailable ({type(e).__name__}: {e})"


def render(value, subs: dict[str, str]):
    """Substitute {out}/{skill}/{fixtures} anywhere in a string, list or dict.

    Lists matter: `baseline` accepts several inputs now (a merged artifact has more
    than one), and a string-only version left those placeholders literal — the L2
    gate then reported "baseline not found: {fixtures}/…", which reads like a
    missing file rather than a substitution that never happened.
    """
    if isinstance(value, str):
        for k, v in subs.items():
            value = value.replace("{" + k + "}", v)
        return value
    if isinstance(value, list):
        return [render(v, subs) for v in value]
    if isinstance(value, dict):
        return {k: render(v, subs) for k, v in value.items()}
    return value


def run_sample(skill_dir: Path, cap_id: str, spec: dict, l2, l2_err: str,
               timeout: int, allow_missing: frozenset[str]) -> tuple[dict, list[str]]:
    """Run a capability's sample, then hand every artifact it made to L2."""
    entry = skill_dir / spec.get("entry", "")
    sample = spec["sample"]
    errors: list[str] = []
    notes: list[str] = []
    with tempfile.TemporaryDirectory(prefix=f"cap-{cap_id}-") as td:
        subs = {"out": td, "skill": str(skill_dir), "fixtures": str(skill_dir / "fixtures")}
        args = [render(str(a), subs) for a in sample.get("args", [])]
        try:
            r = subprocess.run([PY, str(entry), *args], capture_output=True, text=True,
                               encoding="utf-8", errors="replace", timeout=timeout)
        except (subprocess.TimeoutExpired, OSError) as e:
            return {"entry": spec.get("entry"), "runs": False, "proof": "sample"}, [
                f"{skill_dir.name} {cap_id}: sample failed to run: {e}"]
        if r.returncode != 0:
            errors.append(f"{skill_dir.name} {cap_id}: sample exited {r.returncode}: "
                          f"{(r.stdout + r.stderr).strip()[:300]}")
        produced = [Path(td) / render(str(n), subs) for n in sample.get("produces", [])]
        checked = 0
        for art in produced:
            if not art.is_file() or art.stat().st_size == 0:
                errors.append(f"{skill_dir.name} {cap_id}: sample declared it produces "
                              f"{art.name} — it is missing or empty")
                continue
            if art.suffix.lower() not in {".docx", ".xlsx", ".pdf"}:
                continue
            if l2 is None:
                errors.append(f"{skill_dir.name} {cap_id}: cannot verify {art.name} — {l2_err}")
                continue
            expect = render(spec.get("expect") or {}, subs)
            findings, _, inert = l2.run_checks(art, expect,
                                               allow_missing=allow_missing)
            checked += 1
            errors += [f"{skill_dir.name} {cap_id}: {art.name} fails L2 — {f}"
                       for f in findings]
            fid = tuple(getattr(l2, "FIDELITY_CHECKS", ()))
            missing = [i for i in inert if not i.startswith(fid)]
            if missing and not findings:
                errors.append(f"{skill_dir.name} {cap_id}: {art.name} passed L2 but "
                              f"{len(missing)} assertion(s) had nothing to check — the "
                              f"sample needs richer `expect`: "
                              f"{', '.join(i.split()[0] for i in missing)}")
            if len(missing) != len(inert):
                # Not a failure — but an edit capability that never declares a
                # baseline is one whose fidelity nobody has ever checked.
                notes.append(f"{skill_dir.name} {cap_id}: no `baseline` in expect, so "
                             f"fidelity was not verified for {art.name}")
        runs = not errors
    return {"entry": spec.get("entry"), "runs": runs, "proof": "sample",
            "artifacts_checked": checked, "notes": notes}, errors


def parse_pending(skill_dir: Path, data: dict, errors: list[str]) -> dict[str, str]:
    """Read the manifest's acknowledged debt.

    C5 without an escape hatch forces all-or-nothing delivery: pdf has 14 target-✅
    rows, so declaring 8 of them reads as a failure and the rational move becomes
    "declare nothing", which is exactly the invisible state C5 was added to kill.
    `pending` makes the debt visible instead — but only debt that is real: an ID
    outside this skill's range, or one that is also implemented, is an error, and
    every pending row is named in the report and rejected outright under --no-pending.
    """
    raw = data.get("pending") or {}
    if isinstance(raw, list):
        pending = {str(k): "" for k in raw}
    elif isinstance(raw, dict):
        pending = {str(k): str(v) for k, v in raw.items()}
    else:
        errors.append(f"{skill_dir.name}: `pending` must be a list or an object, "
                      f"got {type(raw).__name__}")
        return {}
    prefix = SKILL_PREFIX.get(skill_dir.name)
    implemented = set(data.get("capabilities") or {})
    for cap_id in sorted(pending):
        if prefix and not cap_id.startswith(prefix):
            errors.append(f"{skill_dir.name}: pending lists {cap_id}, which is not in "
                          f"this skill's {prefix}* range")
        if cap_id in implemented:
            errors.append(f"{skill_dir.name}: {cap_id} is both implemented and pending")
    return pending


def probe_declared(skill_dir: Path, timeout: int = 60,
                   allow_missing: frozenset[str] = frozenset()
                   ) -> tuple[dict[str, dict], list[str]]:
    """Make every declared capability prove itself. Returns (capabilities, errors)."""
    manifest = skill_dir / "capabilities.json"
    if not manifest.is_file():
        return {}, [], {}
    try:
        data = json.loads(manifest.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        return {}, [f"{skill_dir.name}/capabilities.json is not valid JSON: {e}"], {}

    l2, l2_err = load_l2()
    errors: list[str] = []
    caps: dict[str, dict] = {}
    pending = parse_pending(skill_dir, data, errors)
    for cap_id, spec in (data.get("capabilities") or {}).items():
        entry = skill_dir / spec.get("entry", "")
        if not entry.is_file():
            errors.append(f"{skill_dir.name} {cap_id}: declared entry missing: {spec.get('entry')}")
            caps[cap_id] = {"entry": spec.get("entry"), "runs": False, "proof": "none"}
            continue
        if spec.get("sample"):
            caps[cap_id], errs = run_sample(skill_dir, cap_id, spec, l2, l2_err,
                                            timeout, allow_missing)
            errors += errs
            continue
        # Fallback: only proves the script starts. Not a failure — a skill mid-build
        # needs somewhere to stand — but it is counted and named in the report so
        # "all green" can never mean "all of them were --help".
        probe = spec.get("probe", ["--help"])
        try:
            r = subprocess.run([PY, str(entry), *probe], capture_output=True, text=True,
                               encoding="utf-8", errors="replace", timeout=timeout)
            runs = r.returncode == 0
            if not runs:
                errors.append(f"{skill_dir.name} {cap_id}: {entry.name} {' '.join(probe)} "
                              f"exited {r.returncode}")
        except (subprocess.TimeoutExpired, OSError) as e:
            runs = False
            errors.append(f"{skill_dir.name} {cap_id}: probing {entry.name} failed: {e}")
        caps[cap_id] = {"entry": spec.get("entry"), "runs": runs, "proof": "startup-only"}
    return caps, errors, pending


# ── selftest ──────────────────────────────────────────────────────────────────
# C4 is now the seam between L1 and L2, and a seam nobody has watched fail is not
# evidence. Each case below declares ONE capability and calls probe_declared
# directly, so C5's "cover the whole ID range" rule stays out of the way.

# ⚠️ Several paragraphs, not one. This stub used to be a single line of EIGHT
# Chinese characters on a Letter page, and D7's blank detector works on the share of
# inked pixels: measured, the same file renders at ink 0.00569 on macOS and 0.00028
# on Windows — the two platforms straddle the 0.0005 threshold, so the POSITIVE
# control was red on Windows for no defect at all. Lowering the threshold was the
# wrong repair (it blunts blank detection for every artifact); making the positive
# control unambiguous is the right one. Generated at run time, so nothing committed
# and no .builtin-version churn.
_LINES = ("季度经营分析报告",
          "本季度营业收入同比增长百分之十二，毛利率保持稳定，费用结构持续优化。",
          "华东区域贡献了增量的主要部分，华南区域受季节性因素影响略有回落。",
          "经营性现金流保持健康水平，应收账款账龄需要在下季度重点关注。",
          "综合判断，维持全年增长预期不变，并建议加强对存货周转的监控。")
_HEAD = ("import sys\nimport docx\nfrom docx.oxml import parse_xml\n"
         "from docx.oxml.ns import nsdecls\n"
         "out = sys.argv[sys.argv.index('--out') + 1]\n"
         "d = docx.Document()\n"
         f"lines = {_LINES!r}\n"
         "runs = [d.add_paragraph().add_run(t) for t in lines]\n")
# Every CJK run gets the binding, because D6 judges runs and not documents: leaving
# one unbound would make the "good" stub fail for a reason the case is not about.
_FONTS = ("for r in runs:\n"
          "    rpr = r._r.get_or_add_rPr()\n"
          "    rpr.insert(0, parse_xml('<w:rFonts %s w:ascii=\"Calibri\" "
          "w:hAnsi=\"Calibri\" w:eastAsia=\"宋体\"/>' % nsdecls('w')))\n")
# python-docx's bundled default.docx ships `<w:zoom w:val="bestFit"/>`, and
# ECMA-376 Transitional makes w:percent REQUIRED — so every document it produces
# fails L2's D2 on a defect inherited from the library's template (found 2026-08-02,
# the first time D2 ran). These stubs are POSITIVE controls, so they have to emit a
# conformant file; the finding itself belongs to S4's docx skill, which must write
# the attribute or omit the element.
_SAVE = (
    "d.save(out)\n"
    "import re, shutil, zipfile\n"
    "tmp = out + '.pre'\n"
    "shutil.move(out, tmp)\n"
    "with zipfile.ZipFile(tmp) as zin, zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zo:\n"
    "    for it in zin.infolist():\n"
    "        b = zin.read(it.filename)\n"
    "        if it.filename == 'word/settings.xml' and b'<w:zoom' in b:\n"
    "            b = re.sub(rb'<w:zoom(?![^>]*w:percent)([^>]*?)/>',\n"
    "                       rb'<w:zoom\\1 w:percent=\"100\"/>', b, count=1)\n"
    "        zo.writestr(it, b)\n"
    "import os; os.remove(tmp)\n")

STUBS = {
    "good": _HEAD + _FONTS + _SAVE,
    # Same script minus the eastAsia binding: produces a perfectly openable .docx
    # that L2's D6 must reject. This is the case `--help` could never catch.
    "bad-artifact": _HEAD + _SAVE,
    "nonzero": "import sys\nsys.exit(3)\n",
    "produces-nothing": "import sys\nsys.exit(0)\n",
}

SAMPLE_CASES = [
    # `len(_LINES)`, not a literal: the stub grew from one paragraph to five so its
    # rendered ink would clear D7's blank threshold on every platform, and a
    # hard-coded 1 here turned that into a D3 count failure. Deriving it means the
    # next change to the stub cannot desynchronise the two.
    ("good sample produces an artifact L2 accepts", "good",
     {"contains": ["季度经营分析报告"], "paragraphs": len(_LINES)}, None),
    ("sample whose artifact violates L2 is rejected", "bad-artifact",
     {"contains": ["季度经营分析报告"], "paragraphs": len(_LINES)}, "fails L2 — D6"),
    ("sample exiting non-zero is rejected", "nonzero", {"contains": ["x"]},
     "sample exited 3"),
    ("sample producing nothing is rejected", "produces-nothing",
     {"contains": ["x"]}, "missing or empty"),
    ("artifact with empty expectations is not accepted as verified", "good", {},
     "had nothing to check"),
]


def _pending_skill(root: Path, tag: str, manifest: dict) -> Path:
    skill = root / tag / "xlsx"
    (skill / "scripts").mkdir(parents=True)
    (skill / "scripts" / "gen.py").write_text("import sys\nsys.exit(0)\n", encoding="utf-8")
    (skill / "capabilities.json").write_text(json.dumps(manifest, ensure_ascii=False),
                                             encoding="utf-8")
    return skill


def pending_cases(root: Path) -> list[dict]:
    """An escape hatch that can hide debt is worse than no escape hatch.

    These pin the three ways `pending` could have become one: laundering an ID that
    is not this skill's, claiming a capability as both done and owed, and surviving
    the acceptance pass.
    """
    ok_entry = {"entry": "scripts/gen.py", "probe": ["--help"]}
    cases = []

    skill = _pending_skill(root, "p1", {"skill": "xlsx",
                                        "capabilities": {"X1": ok_entry},
                                        "pending": {"X3": "needs the formula evaluator"}})
    _, errors, pending = probe_declared(skill, allow_missing=ALL_TIERS)
    cases.append({"case": "valid pending is accepted and carries its reason",
                  "marker": None, "ok": not errors and pending.get("X3", "") != "",
                  "errors": errors or ([] if pending else ["pending was dropped"])})

    skill = _pending_skill(root, "p2", {"skill": "xlsx",
                                        "capabilities": {"X1": ok_entry},
                                        "pending": ["W6"]})
    _, errors, _ = probe_declared(skill, allow_missing=ALL_TIERS)
    cases.append({"case": "pending listing another skill's ID is rejected",
                  "marker": "not in this skill",
                  "ok": any("not in this skill" in e for e in errors), "errors": errors})

    skill = _pending_skill(root, "p3", {"skill": "xlsx",
                                        "capabilities": {"X1": ok_entry},
                                        "pending": ["X1"]})
    _, errors, _ = probe_declared(skill, allow_missing=ALL_TIERS)
    cases.append({"case": "a capability cannot be both implemented and pending",
                  "marker": "both implemented and pending",
                  "ok": any("both implemented and pending" in e for e in errors),
                  "errors": errors})

    skill = _pending_skill(root, "p4", {"skill": "xlsx",
                                        "capabilities": {"X1": ok_entry},
                                        "pending": "X3"})
    _, errors, _ = probe_declared(skill, allow_missing=ALL_TIERS)
    cases.append({"case": "a malformed pending block is rejected",
                  "marker": "must be a list or an object",
                  "ok": any("must be a list or an object" in e for e in errors),
                  "errors": errors})

    # End-to-end against the real gate. Three points, because two would not pin it:
    # asserting only "no C5 pending failure" would pass even if `pending` never
    # suppressed anything, since C5 unbacked is a different message. The C arm is
    # the control that proves the suppression is real.
    rows, _ = parse_matrix(DOC)
    xlsx_ids = [r["id"] for r in rows
                if r["id"].startswith("X") and mark(r[TARGET_COL]) == YES]
    owed, done = xlsx_ids[-2:], xlsx_ids[:-2]

    def gate(tag: str, manifest: dict, extra: list[str]) -> list[str]:
        base = root / tag
        _pending_skill(base, ".", manifest)
        r = subprocess.run(
            [PY, str(Path(__file__).resolve()), "--json", "--skill", "xlsx",
             "--builtin-root", str(base / "."), "--allow-missing", "soffice", "xsd",
             *extra],
            capture_output=True, text=True, encoding="utf-8", errors="replace",
            timeout=300)
        try:
            return json.loads(r.stdout)["failures"]
        except (json.JSONDecodeError, KeyError):
            # 2000, not 200. When the gate itself CRASHES this string is the only
            # evidence there is, and at 200 characters it stopped one line into the
            # traceback — a Windows-only failure was unreadable from the CI log and
            # had to be re-derived from scratch. A diagnostic that truncates before
            # the cause is not a diagnostic.
            return [f"<could not parse gate output>: exit={r.returncode}\n"
                    f"{(r.stdout + r.stderr)[:2000]}"]

    full = {"skill": "xlsx", "capabilities": {i: ok_entry for i in done},
            "pending": {i: "scheduled for S3" for i in owed}}
    fails = gate("e2e-a", full, [])
    cases.append({"case": f"pending clears C5 for the {len(owed)} owed rows while building",
                  "marker": None,
                  "ok": not any(f.startswith("C5") for f in fails), "errors": fails})

    fails = gate("e2e-b", full, ["--no-pending"])
    cases.append({"case": "--no-pending refuses to ship acknowledged debt",
                  "marker": "C5 pending",
                  "ok": any(f.startswith("C5 pending") for f in fails), "errors": fails})

    silent = {"skill": "xlsx", "capabilities": {i: ok_entry for i in done}}
    fails = gate("e2e-c", silent, [])
    cases.append({"case": "CONTROL: dropping the pending block brings C5 unbacked back",
                  "marker": "C5 unbacked",
                  "ok": any(f.startswith("C5 unbacked") for f in fails), "errors": fails})

    # Scoping the acceptance pass (added 2026-08-02, 059 S3). The plan builds one
    # skill at a time, so an unscoped --no-pending goes red the moment the next
    # skill starts — and a permanently red job is one nobody reads. Three arms,
    # because the middle one alone would pass even if the scope were ignored and the
    # gate had simply stopped checking.
    fails = gate("e2e-d", full, ["--no-pending", "xlsx"])
    cases.append({"case": "--no-pending SKILL still refuses that skill's debt",
                  "marker": "C5 pending",
                  "ok": any(f.startswith("C5 pending") for f in fails), "errors": fails})

    fails = gate("e2e-e", full, ["--no-pending", "docx"])
    cases.append({"case": "a scope naming a skill with no manifest asserts nothing "
                          "and is rejected",
                  "marker": "assert nothing",
                  "ok": any("assert nothing" in f for f in fails), "errors": fails})

    clean = {"skill": "xlsx", "capabilities": {i: ok_entry for i in xlsx_ids}}
    fails = gate("e2e-f", clean, ["--no-pending", "xlsx"])
    cases.append({"case": "CONTROL: a scoped pass goes green once that skill owes "
                          "nothing", "marker": None,
                  "ok": not any(f.startswith("C5") for f in fails), "errors": fails})
    return cases


def capability_selftest(as_json: bool) -> int:
    results = []
    with tempfile.TemporaryDirectory(prefix="cap-selftest-") as td:
        for i, (name, stub, expect, marker) in enumerate(SAMPLE_CASES):
            skill = Path(td) / f"case{i}" / "docx"
            (skill / "scripts").mkdir(parents=True)
            (skill / "scripts" / "gen.py").write_text(STUBS[stub], encoding="utf-8")
            (skill / "capabilities.json").write_text(json.dumps({
                "skill": "docx",
                "capabilities": {"W1": {
                    "entry": "scripts/gen.py",
                    "sample": {"args": ["--out", "{out}/o.docx"], "produces": ["o.docx"]},
                    "expect": expect}}}, ensure_ascii=False), encoding="utf-8")
            caps, errors, _ = probe_declared(skill, allow_missing=ALL_TIERS)
            ok = (not errors) if marker is None else any(marker in e for e in errors)
            results.append({"case": name, "marker": marker, "ok": ok, "errors": errors})

        # A capability that only declares `probe` must still pass, but must be
        # labelled startup-only — otherwise everyone declares `--help` and C4 is
        # back to where it started.
        skill = Path(td) / "weak" / "docx"
        (skill / "scripts").mkdir(parents=True)
        (skill / "scripts" / "gen.py").write_text("import sys\nsys.exit(0)\n",
                                                  encoding="utf-8")
        (skill / "capabilities.json").write_text(json.dumps({
            "skill": "docx",
            "capabilities": {"W1": {"entry": "scripts/gen.py", "probe": ["--help"]}}}),
            encoding="utf-8")
        caps, errors, _ = probe_declared(skill, allow_missing=ALL_TIERS)
        results.append({"case": "startup-only proof passes but is labelled weak",
                        "marker": "proof=startup-only",
                        "ok": not errors and caps["W1"]["proof"] == "startup-only",
                        "errors": errors})

        results += pending_cases(Path(td))

    failed = [r for r in results if not r["ok"]]
    if as_json:
        print(json.dumps({"results": results, "failed": len(failed)},
                         ensure_ascii=False, indent=2))
    else:
        for r in results:
            print(f"{'PASS' if r['ok'] else 'FAIL'}  "
                  f"[{'must reject: ' + r['marker'] if r['marker'] else 'must accept'}] "
                  f"{r['case']}")
            if not r["ok"]:
                for e in r["errors"][:4]:
                    print(f"      · {e}")
                if not r["errors"]:
                    print("      · no error was raised at all")
        print(f"\n[capabilities-selftest] {len(results) - len(failed)} passed, "
              f"{len(failed)} failed")
    return 1 if failed else 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    # append+default would ADD to the defaults instead of replacing them, so an
    # explicit --skill would silently still probe pdf/docx/xlsx as well.
    ap.add_argument("--skill", action="append", default=None,
                    help="skill dirs holding capabilities.json (default: pdf docx xlsx)")
    ap.add_argument("--builtin-root", default=None,
                    help="override skills/builtin (used by the negative-control tests)")
    ap.add_argument("--json", action="store_true")
    # Scopeable because the plan builds one skill at a time (059 §6). An unscoped
    # --no-pending is the handover gate and stays the default; naming skills lets CI
    # keep asserting "pdf owes nothing" while xlsx is still mid-slice, instead of the
    # only options being a permanently red job or no acceptance signal at all.
    ap.add_argument("--no-pending", nargs="*", default=None, metavar="SKILL",
                    help="acceptance mode: any capability still listed as pending "
                         "fails. Names limit it to those skills; bare = every skill")
    ap.add_argument("--allow-missing", nargs="+", default=[], metavar="TIER",
                    help="forwarded to the L2 gate for dev machines lacking soffice/xsd")
    ap.add_argument("--selftest", action="store_true",
                    help="prove the C4 sample path rejects what it should")
    args = ap.parse_args()
    if args.selftest:
        return capability_selftest(args.json)
    builtin = Path(args.builtin_root).resolve() if args.builtin_root else BUILTIN

    rows, columns = parse_matrix(DOC)
    if not rows:
        print("[error] parsed 0 capability rows — did the table format change?",
              file=sys.stderr)
        return 1

    ref_cols = [c for c in columns
                if c not in {"#", "能力", CURRENT_COL, TARGET_COL} and c.strip()]

    failures: list[str] = []

    # ---- C1: no regression against what we ship today ----
    regressions = [r["id"] for r in rows
                   if RANK[mark(r[TARGET_COL])] < RANK[mark(r[CURRENT_COL])]]
    if regressions:
        failures.append(f"C1 regression — target below current for: {', '.join(regressions)}")

    # ---- C2: out-cover every single reference ----
    target_yes = sum(1 for r in rows if mark(r[TARGET_COL]) == YES)
    ref_yes = {c: sum(1 for r in rows if mark(r[c]) == YES) for c in ref_cols}
    beaten = [c for c, n in ref_yes.items() if target_yes <= n]
    if beaten:
        failures.append(f"C2 coverage — target {target_yes} does not exceed: "
                        + ", ".join(f"{c}={ref_yes[c]}" for c in beaten))

    # ---- C3: genuinely novel capabilities ----
    beyond = [r["id"] for r in rows if r["beyond"]]
    if len(beyond) < MIN_BEYOND:
        failures.append(f"C3 novelty — only {len(beyond)} beyond-reference items, "
                        f"need {MIN_BEYOND}")

    # ---- C4: the matrix must be backed by something that runs ----
    skills = args.skill or list(SKILL_PREFIX)
    declared: dict[str, dict] = {}
    pending: dict[str, str] = {}
    probe_errors: list[str] = []
    undeclared: list[str] = []
    migrated: list[str] = []
    for name in skills:
        d = builtin / name
        if not d.is_dir():
            undeclared.append(f"{name} (skill dir does not exist yet)")
            continue
        caps, errs, pend = probe_declared(d, allow_missing=frozenset(args.allow_missing))
        pending.update({f"{name}:{k}": v for k, v in pend.items()})
        if not caps and not errs and not pend:
            undeclared.append(f"{name} (no capabilities.json)")
        else:
            migrated.append(name)
        declared.update({f"{name}:{k}": v for k, v in caps.items()})
        probe_errors += errs
    if probe_errors:
        failures.append(f"C4 probes — {len(probe_errors)} declared entry point(s) do not run")

    # C5: once a skill declares capabilities.json it has opted into the contract and
    # must back EVERY target-✅ row in its ID range. Without this, forgetting to declare
    # a capability reads as "OK" — exactly the "ticked on paper, absent in code" failure
    # this gate exists to prevent. Skills with no manifest are mid-migration, not failing.
    claimed_ids = {k.split(":", 1)[1] for k in declared}
    pending_ids = {k.split(":", 1)[1] for k in pending}
    unbacked: list[str] = []
    for name in migrated:
        prefix = SKILL_PREFIX[name]
        unbacked += [r["id"] for r in rows
                     if r["id"].startswith(prefix)
                     and mark(r[TARGET_COL]) == YES
                     and r["id"] not in claimed_ids and r["id"] not in pending_ids]
    if unbacked:
        failures.append(f"C5 unbacked — {len(unbacked)} target-✅ row(s) are neither "
                        f"implemented nor listed as pending in an already-migrated "
                        f"skill: {', '.join(unbacked)}")
    # Acknowledged debt is fine while building and unacceptable at handover: §5's
    # replacement criterion says every gate must be green before the new skills
    # replace anything. --no-pending is that final pass.
    if args.no_pending is not None:
        scope = set(args.no_pending)
        owed = sorted(k for k in pending if not scope or k.split(":")[0] in scope)
        # A scope naming a skill that declares nothing would pass in silence, which
        # is the same shape as the empty-manifest hole C5 exists to close.
        unknown = sorted(s for s in scope if s not in set(migrated))
        if unknown:
            failures.append(f"C5 pending — --no-pending names {', '.join(unknown)}, "
                            f"which declare(s) no capabilities.json; the scope would "
                            f"assert nothing")
        if owed:
            where = f" in {', '.join(sorted(scope))}" if scope else ""
            failures.append(f"C5 pending — {len(owed)} capability(ies) still owed{where} "
                            f"and --no-pending was requested: {', '.join(owed)}")

    if args.json:
        print(json.dumps({
            "rows": len(rows), "target_yes": target_yes, "reference_yes": ref_yes,
            "beyond": beyond, "declared": declared, "unbacked": unbacked,
            "pending": pending,
            "weak_proofs": sorted(k for k, v in declared.items()
                                  if v.get("proof") == "startup-only"),
            "undeclared_skills": undeclared, "probe_errors": probe_errors,
            "failures": failures,
        }, ensure_ascii=False, indent=2))
    else:
        print(f"[capabilities] matrix: {len(rows)} rows across {len(ref_cols)} references")
        print(f"  target ✅ = {target_yes}   references: "
              + ", ".join(f"{c}={n}" for c, n in ref_yes.items()))
        print(f"  beyond-reference items ({len(beyond)}): {', '.join(beyond)}")
        if undeclared:
            print(f"  not yet declaring capabilities.json: {', '.join(undeclared)}")
        if declared:
            ok = sum(1 for v in declared.values() if v["runs"])
            by_sample = [k for k, v in declared.items() if v.get("proof") == "sample"]
            weak = [k for k, v in declared.items() if v.get("proof") == "startup-only"]
            arts = sum(v.get("artifacts_checked", 0) for v in declared.values())
            print(f"  declared entry points: {ok}/{len(declared)} pass "
                  f"({len(by_sample)} proved by sample, {arts} artifact(s) verified by L2)")
            for note in sorted({n for v in declared.values() for n in v.get("notes", [])}):
                print(f"  NOTE {note}")
            if weak:
                # Never let this stay invisible: a capability proved only by `--help`
                # is a capability nobody has watched do anything.
                print(f"  ⚠ {len(weak)} capability(ies) proved by startup only "
                      f"(no artifact was produced or checked): {', '.join(sorted(weak))}")
        for e in probe_errors:
            print(f"    PROBE FAIL {e}")
        if pending:
            # Never a one-line count: debt that is not itemised stops being read.
            print(f"  {len(pending)} capability(ies) acknowledged as PENDING "
                  f"(not implemented, not hidden):")
            for k in sorted(pending):
                print(f"    · {k}" + (f" — {pending[k]}" if pending[k]
                                      else "  (no reason given)"))
        if unbacked:
            print(f"  {len(unbacked)} target-✅ rows neither implemented nor pending: "
                  f"{', '.join(unbacked)}")
        for f in failures:
            print(f"  FAIL {f}")
        print(f"[capabilities] {'OK' if not failures else str(len(failures)) + ' gate(s) failed'}")

    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
